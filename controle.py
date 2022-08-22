import sys
from novo.login import *
from novo.menu import *
from novo.obra import *
from novo.client import *
from novo.make_doc import *
from novo.pop_up import *
from PyQt6.QtWidgets import QMainWindow, QApplication, QMessageBox, QSizeGrip
from PyQt6 import QtWidgets, uic, QtCore, QtGui
import sqlite3
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date
from num2words import num2words
from pathlib import Path

multi = 0

today = date.today().strftime('%d-%m-%Y')
dia = date.today().strftime('%d')
mes = date.today().strftime('%m')
ano = date.today().strftime('%Y')
if mes == '01':
    mesescrito = 'janeiro'
elif mes == '02':
    mesescrito = 'fevereiro'
elif mes == '03':
    mesescrito = 'março'
elif mes == '04':
    mesescrito = 'abril'
elif mes == '05':
    mesescrito = 'maio'
elif mes == '06':
    mesescrito = 'junho'
elif mes == '07':
    mesescrito = 'julho'
elif mes == '08':
    mesescrito = 'agosto'
elif mes == '09':
    mesescrito = 'setembro'
elif mes == '10':
    mesescrito = 'outubro'
elif mes == '11':
    mesescrito = 'novembro'
elif mes == '12':
    mesescrito = 'dezembro'

class Menu(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        uic.loadUi("menu.ui", self)
        self.btn_cadobra.clicked.connect(self.abriobra)
        self.btn_cliente.clicked.connect(self.abricliente)
        self.btn_makedoc.clicked.connect(self.abrirdoc)
        self.minwin.clicked.connect(self.mini)
        self.maxwin.clicked.connect(self.max)
        self.closewin.clicked.connect(self.fecha)
        self.left.mousePressEvent=self.myfunction
        self.left.mouseMoveEvent = self.myfunc


        self.setWindowFlags(QtCore.Qt.WindowType.FramelessWindowHint)
        self.gripSize = 16
        self.grips = []
        for i in range(4):
            grip = QSizeGrip(self)
            grip.resize(self.gripSize, self.gripSize)
            self.grips.append(grip)


    def resizeEvent(self, event):
        QMainWindow.resizeEvent(self, event)
        rect = self.rect()
        # top left grip doesn't need to be moved...
        # top right
        self.grips[1].move(rect.right() - self.gripSize, 0)
        # bottom right
        self.grips[2].move(
        rect.right() - self.gripSize, rect.bottom() - self.gripSize)
        # bottom left
        self.grips[3].move(0, rect.bottom() - self.gripSize)

    def mini(self):
        widget.showMinimized()

    def max(self):
        global multi
        status = multi
        if status == 0:
            widget.showMaximized()
            self.frame.setStyleSheet("background-color: rgb(40, 40, 40);\n"
                                     "border-radius: 0px;\n"
                                     "color:rgb(200, 200, 255);")
            self.maxwin.setText("❐")
            multi = 1
        else:
            widget.showNormal()
            menu.resize(menu.width()+1, menu.height()+1)
            self.frame.setStyleSheet("background-color: rgb(40, 40, 40);\n"
                                     "border-radius: 10px;\n"
                                     "color:rgb(200, 200, 255);")
            self.maxwin.setText("☐")
            multi = 0


    def fecha(self):
        widget.close()


    def myfunction(self, event):
        widget.dragPos = event.globalPosition().toPoint()


    def myfunc(self, event):
        widget.move(widget.pos() + event.globalPosition().toPoint() - widget.dragPos )
        widget.dragPos = event.globalPosition().toPoint()
        event.accept()

    def abriobra(self):
        obra = Obra()
        widget.addWidget(obra)
        widget.setCurrentIndex(widget.currentIndex() + 1)

    def abricliente(self):
        cliente = Cliente()
        widget.addWidget(cliente)
        widget.setCurrentIndex(widget.currentIndex()+1)

    def abrirdoc(self):
        doc = Doc()
        widget.addWidget(doc)
        widget.setCurrentIndex(widget.currentIndex() + 1)

class Pop(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        uic.loadUi("pop_up.ui", self)
        self.close_popup.clicked.connect(self.fecha)
        self.lbl_popup.mousePressEvent = self.myfunction
        self.lbl_popup.mouseMoveEvent = self.myfunc

    def myfunction(self, event):
        up.dragPos = event.globalPosition().toPoint()

    def myfunc(self, event):
        up.move(up.pos() + event.globalPosition().toPoint() - up.dragPos)
        up.dragPos = event.globalPosition().toPoint()
        event.accept()

    def fecha(self):
        up.close()


class Obra(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        uic.loadUi("obra.ui", self)
        self.btn_return.clicked.connect(self.volta)
        self.btn_search.clicked.connect(self.search)
        self.btn_add.clicked.connect(self.add)
        #self.pushButton.clicked.connect(self.close)
        self.btn_copy.clicked.connect(self.copy)
        self.btn_delete.clicked.connect(self.delete)
        self.btn_edit.clicked.connect(self.edit)
        self.btn_copycli1.clicked.connect(self.copycli1)
        self.btn_copycli2.clicked.connect(self.copycli2)
        self.btn_copycli3.clicked.connect(self.copycli3)
        self.btn_copycli4.clicked.connect(self.copycli4)
        self.minwin.clicked.connect(self.mini)
        self.maxwin.clicked.connect(self.max)
        self.closewin.clicked.connect(self.fecha)
        self.framelogo.mousePressEvent = self.myfunction
        self.framelogo.mouseMoveEvent = self.myfunc

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        # consulta = 'SELECT * FROM cliente'
        cursor.execute('SELECT * FROM obra')
        dados_lidos = cursor.fetchall()
        self.tabWid_obra.setRowCount(len(dados_lidos))
        self.tabWid_obra.setColumnCount(20)

        for i in range(0, len(dados_lidos)):
            for j in range(0, 20):
                self.tabWid_obra.setItem(i, j, QtWidgets.QTableWidgetItem(str(dados_lidos[i][j])))
        banco.commit()
        banco.close()

        self.setWindowFlags(QtCore.Qt.WindowType.FramelessWindowHint)
        self.gripSize = 16
        self.grips = []
        for i in range(4):
            grip = QSizeGrip(self)
            grip.resize(self.gripSize, self.gripSize)
            self.grips.append(grip)

    def resizeEvent(self, event):
        QMainWindow.resizeEvent(self, event)
        rect = self.rect()
        # top left grip doesn't need to be moved...
        # top right
        self.grips[1].move(rect.right() - self.gripSize, 0)
        # bottom right
        self.grips[2].move(
            rect.right() - self.gripSize, rect.bottom() - self.gripSize)
        # bottom left
        self.grips[3].move(0, rect.bottom() - self.gripSize)

    def myfunction(self, event):
        widget.dragPos = event.globalPosition().toPoint()

    def myfunc(self, event):
        widget.move(widget.pos() + event.globalPosition().toPoint() - widget.dragPos)
        widget.dragPos = event.globalPosition().toPoint()
        event.accept()

    def mini(self):
        widget.showMinimized()

    def max(self):
        global multi
        status = multi
        if status == 0:
            widget.showMaximized()
            self.frame.setStyleSheet("background-color: rgb(40, 40, 40);\n"
                                     "border-radius: 0px;\n"
                                     "color:rgb(200, 200, 255);")
            self.maxwin.setText("❐")
            multi = 1
        else:
            widget.showNormal()
            menu.resize(menu.width() + 1, menu.height() + 1)
            self.frame.setStyleSheet("background-color: rgb(40, 40, 40);\n"
                                     "border-radius: 10px;\n"
                                     "color:rgb(200, 200, 255);")
            self.maxwin.setText("☐")
            multi = 0

    def fecha(self):
        widget.close()

    def volta(self):
        menu = Menu()
        widget.addWidget(menu)
        widget.setCurrentIndex(widget.currentIndex() + 1)

    def copy(self):
        idobra = self.txt_id.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'SELECT * FROM obra  WHERE id LIKE ?'
        cursor.execute(consulta, (idobra,))

        dados_lidos = cursor.fetchall()

        self.txt_obraend.setText(dados_lidos[0][1])
        self.txt_obrabairro.setText(dados_lidos[0][2])
        self.txt_obranumero.setText(dados_lidos[0][3])
        self.txt_obracidade.setText(dados_lidos[0][4])
        self.txt_obralote.setText(dados_lidos[0][5])
        self.txt_obraquadra.setText(dados_lidos[0][6])
        self.txt_obraquarteirao.setText(dados_lidos[0][7])
        self.txt_obratipo.setText(dados_lidos[0][8])
        self.txt_obraareaterreno.setText(dados_lidos[0][9])
        self.txt_obraart.setText(dados_lidos[0][10])
        self.txt_obravalorparcela.setText(dados_lidos[0][11])
        self.txt_obraquantparcela.setText(dados_lidos[0][12])
        self.txt_obradatacontrato.setText(dados_lidos[0][13])
        self.txt_obravalorvisita.setText(dados_lidos[0][14])
        self.txt_obravalorvisita_2.setText(dados_lidos[0][15])
        self.txt_idcli1.setText(dados_lidos[0][16])
        self.txt_idcli2.setText(dados_lidos[0][17])
        self.txt_idcli3.setText(dados_lidos[0][18])
        self.txt_idcli4.setText(dados_lidos[0][19])
        banco.commit()
        banco.close()

    def copycli1(self):
        idocli = self.txt_idcli1.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'SELECT * FROM cliente  WHERE id LIKE ?'
        cursor.execute(consulta, (idocli,))

        dados_lidos = cursor.fetchall()

        self.txt_obraend.setText(dados_lidos[0][4])
        self.txt_obrabairro.setText(dados_lidos[0][5])
        self.txt_obranumero.setText(dados_lidos[0][6])
        self.txt_obracidade.setText(dados_lidos[0][7])
        banco.commit()
        banco.close()

    def copycli2(self):
        idocli = self.txt_idcli2.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'SELECT * FROM cliente  WHERE id LIKE ?'
        cursor.execute(consulta, (idocli,))

        dados_lidos = cursor.fetchall()

        self.txt_obraend.setText(dados_lidos[0][4])
        self.txt_obrabairro.setText(dados_lidos[0][5])
        self.txt_obranumero.setText(dados_lidos[0][6])
        self.txt_obracidade.setText(dados_lidos[0][7])
        banco.commit()
        banco.close()

    def copycli3(self):
        idocli = self.txt_idcli3.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'SELECT * FROM cliente  WHERE id LIKE ?'
        cursor.execute(consulta, (idocli,))

        dados_lidos = cursor.fetchall()

        self.txt_obraend.setText(dados_lidos[0][4])
        self.txt_obrabairro.setText(dados_lidos[0][5])
        self.txt_obranumero.setText(dados_lidos[0][6])
        self.txt_obracidade.setText(dados_lidos[0][7])
        banco.commit()
        banco.close()

    def copycli4(self):
        idocli = self.txt_idcli4.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'SELECT * FROM cliente  WHERE id LIKE ?'
        cursor.execute(consulta, (idocli,))

        dados_lidos = cursor.fetchall()

        self.txt_obraend.setText(dados_lidos[0][4])
        self.txt_obrabairro.setText(dados_lidos[0][5])
        self.txt_obranumero.setText(dados_lidos[0][6])
        self.txt_obracidade.setText(dados_lidos[0][7])
        banco.commit()
        banco.close()

    def delete(self):
        idobra = self.txt_id.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'DELETE FROM obra WHERE id=?'
        cursor.execute(consulta, (idobra,))

        self.txt_id.setText(None)
        self.txt_obraend.setText(None)
        self.txt_obrabairro.setText(None)
        self.txt_obranumero.setText(None)
        self.txt_obracidade.setText(None)
        self.txt_obralote.setText(None)
        self.txt_obraquadra.setText(None)
        self.txt_obraquarteirao.setText(None)
        self.txt_obratipo.setText(None)
        self.txt_obraareaterreno.setText(None)
        self.txt_obraart.setText(None)
        self.txt_obravalorparcela.setText(None)
        self.txt_obraquantparcela.setText(None)
        self.txt_obradatacontrato.setText(None)
        self.txt_obravalorvisita.setText(None)
        self.txt_obravalorvisita_2.setText(None)
        self.txt_idcli1.setText(None)
        self.txt_idcli2.setText(None)
        self.txt_idcli3.setText(None)
        self.txt_idcli4.setText(None)

        cursor.execute('SELECT * FROM obra')
        dados_lidos = cursor.fetchall()
        self.tabWid_obra.setRowCount(len(dados_lidos))
        self.tabWid_obra.setColumnCount(19)

        for i in range(0, len(dados_lidos)):
            for j in range(0, 20):
                self.tabWid_obra.setItem(i, j, QtWidgets.QTableWidgetItem(str(dados_lidos[i][j])))
        banco.commit()
        banco.close()

    def edit(self):
        idobra = self.txt_id.text()
        obraend = self.txt_obraend.text()
        obrabairro = self.txt_obrabairro.text()
        obranumero = self.txt_obranumero.text()
        obracidade = self.txt_obracidade.text()
        obralote = self.txt_obralote.text()
        obraquadra = self.txt_obraquadra.text()
        obraquarteirao = self.txt_obraquarteirao.text()
        obratipo = self.txt_obratipo.text()
        obraarea = self.txt_obraareaterreno.text()
        obraart = self.txt_obraart.text()
        obravalorparc = self.txt_obravalorparcela.text()
        obraquantparc = self.txt_obraquantparcela.text()
        obradatacont = self.txt_obradatacontrato.text()
        obravalorvisit = self.txt_obravalorvisita.text()
        obrainscmob = self.txt_obravalorvisita_2.text()
        obraidcli1 = self.txt_idcli1.text()
        obraidcli2 = self.txt_idcli2.text()
        obraidcli3 = self.txt_idcli3.text()
        obraidcli4 = self.txt_idcli4.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'UPDATE OR IGNORE obra SET end=?, bairro=?, num=?, cidade=?, lote=?, quadra=?, quarteirao=?, tipo=?, area=?, art=?, valorparc=?, quantparc=?, datacontrato=?, valorvisita=?, inscimob=?, idcli1=?, idcli2=?, idcli3=?, idcli4=? WHERE id=?'
        cursor.execute(consulta, (obraend, obrabairro, obranumero, obracidade, obralote, obraquadra, obraquarteirao, obratipo, obraarea, obraart, obravalorparc, obraquantparc, obradatacont, obravalorvisit, obrainscmob, obraidcli1, obraidcli2, obraidcli3, obraidcli4, idobra))

        self.txt_id.setText(None)
        self.txt_obraend.setText(None)
        self.txt_obrabairro.setText(None)
        self.txt_obranumero.setText(None)
        self.txt_obracidade.setText(None)
        self.txt_obralote.setText(None)
        self.txt_obraquadra.setText(None)
        self.txt_obraquarteirao.setText(None)
        self.txt_obratipo.setText(None)
        self.txt_obraareaterreno.setText(None)
        self.txt_obraart.setText(None)
        self.txt_obravalorparcela.setText(None)
        self.txt_obraquantparcela.setText(None)
        self.txt_obradatacontrato.setText(None)
        self.txt_obravalorvisita.setText(None)
        self.txt_obravalorvisita_2.setText(None)
        self.txt_idcli1.setText(None)
        self.txt_idcli2.setText(None)
        self.txt_idcli3.setText(None)
        self.txt_idcli4.setText(None)

        cursor.execute('SELECT * FROM obra')
        dados_lidos = cursor.fetchall()
        self.tabWid_obra.setRowCount(len(dados_lidos))
        self.tabWid_obra.setColumnCount(20)

        for i in range(0, len(dados_lidos)):
            for j in range(0, 20):
                self.tabWid_obra.setItem(i, j, QtWidgets.QTableWidgetItem(str(dados_lidos[i][j])))
        banco.commit()
        banco.close()


    def add(self):
        obraend = self.txt_obraend.text()
        obrabairro = self.txt_obrabairro.text()
        obranumero = self.txt_obranumero.text()
        obracidade = self.txt_obracidade.text()
        obralote = self.txt_obralote.text()
        obraquadra = self.txt_obraquadra.text()
        obraquarteirao = self.txt_obraquarteirao.text()
        obratipo = self.txt_obratipo.text()
        obraarea = self.txt_obraareaterreno.text()
        obraart = self.txt_obraart.text()
        obravalorparc = self.txt_obravalorparcela.text()
        obraquantparc = self.txt_obraquantparcela.text()
        obradatacont = self.txt_obradatacontrato.text()
        obravalorvisit = self.txt_obravalorvisita.text()
        obrainscmob = self.txt_obravalorvisita_2.text()
        obraidcli1 = self.txt_idcli1.text()
        obraidcli2 = self.txt_idcli2.text()
        obraidcli3 = self.txt_idcli3.text()
        obraidcli4 = self.txt_idcli4.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'INSERT OR IGNORE INTO obra (end, bairro, num, cidade, lote, quadra, quarteirao, tipo, area, art, valorparc, quantparc, datacontrato, valorvisita, inscimob, idcli1, idcli2, idcli3, idcli4) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)'
        cursor.execute(consulta, (obraend, obrabairro, obranumero, obracidade, obralote, obraquadra, obraquarteirao, obratipo, obraarea, obraart, obravalorparc, obraquantparc, obradatacont, obravalorvisit, obrainscmob, obraidcli1, obraidcli2, obraidcli3, obraidcli4))

        self.txt_id.setText(None)
        self.txt_obraend.setText(None)
        self.txt_obrabairro.setText(None)
        self.txt_obranumero.setText(None)
        self.txt_obracidade.setText(None)
        self.txt_obralote.setText(None)
        self.txt_obraquadra.setText(None)
        self.txt_obraquarteirao.setText(None)
        self.txt_obratipo.setText(None)
        self.txt_obraareaterreno.setText(None)
        self.txt_obraart.setText(None)
        self.txt_obravalorparcela.setText(None)
        self.txt_obraquantparcela.setText(None)
        self.txt_obradatacontrato.setText(None)
        self.txt_obravalorvisita.setText(None)
        self.txt_obravalorvisita_2.setText(None)
        self.txt_idcli1.setText(None)
        self.txt_idcli2.setText(None)
        self.txt_idcli3.setText(None)
        self.txt_idcli4.setText(None)

        cursor.execute('SELECT * FROM obra')
        dados_lidos = cursor.fetchall()
        self.tabWid_obra.setRowCount(len(dados_lidos))
        self.tabWid_obra.setColumnCount(19)

        for i in range(0, len(dados_lidos)):
            for j in range(0, 20):
                self.tabWid_obra.setItem(i, j, QtWidgets.QTableWidgetItem(str(dados_lidos[i][j])))
        banco.commit()
        banco.close()

    def search(self):
        name = self.txt_rua.text().lower()
        for row in range(self.tabWid_obra.rowCount()):
            item = self.tabWid_obra.item(row, 1)
            # if the search is *not* in the item's text *do not hide* the row
            self.tabWid_obra.setRowHidden(row, name not in item.text().lower())


class Cliente(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        uic.loadUi("client.ui", self)
        self.btn_return.clicked.connect(self.volta)
        self.btn_search.clicked.connect(self.pesqui)
        self.btn_add.clicked.connect(self.add)
        self.btn_copy.clicked.connect(self.copy)
        self.btn_edit.clicked.connect(self.edit)
        self.btn_delete.clicked.connect(self.delete)
        self.minwin.clicked.connect(self.mini)
        self.maxwin.clicked.connect(self.max)
        self.closewin.clicked.connect(self.fecha)
        self.framelogo.mousePressEvent = self.myfunction
        self.framelogo.mouseMoveEvent = self.myfunc

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        # consulta = 'SELECT * FROM cliente'
        cursor.execute('SELECT * FROM cliente')
        dados_lidos = cursor.fetchall()
        self.tabWid_cli.setRowCount(len(dados_lidos))
        self.tabWid_cli.setColumnCount(15)

        for i in range(0, len(dados_lidos)):
            for j in range(0, 15):
                self.tabWid_cli.setItem(i, j, QtWidgets.QTableWidgetItem(str(dados_lidos[i][j])))
        banco.commit()
        banco.close()

        self.setWindowFlags(QtCore.Qt.WindowType.FramelessWindowHint)
        self.gripSize = 16
        self.grips = []
        for i in range(4):
            grip = QSizeGrip(self)
            grip.resize(self.gripSize, self.gripSize)
            self.grips.append(grip)

    def resizeEvent(self, event):
        QMainWindow.resizeEvent(self, event)
        rect = self.rect()
        # top left grip doesn't need to be moved...
        # top right
        self.grips[1].move(rect.right() - self.gripSize, 0)
        # bottom right
        self.grips[2].move(
            rect.right() - self.gripSize, rect.bottom() - self.gripSize)
        # bottom left
        self.grips[3].move(0, rect.bottom() - self.gripSize)

    def myfunction(self, event):
        widget.dragPos = event.globalPosition().toPoint()

    def myfunc(self, event):
        widget.move(widget.pos() + event.globalPosition().toPoint() - widget.dragPos)
        widget.dragPos = event.globalPosition().toPoint()
        event.accept()

    def mini(self):
        widget.showMinimized()

    def max(self):
        global multi
        status = multi
        if status == 0:
            widget.showMaximized()
            self.frame.setStyleSheet("background-color: rgb(40, 40, 40);\n"
                                     "border-radius: 0px;\n"
                                     "color:rgb(200, 200, 255);")
            self.maxwin.setText("❐")
            multi = 1
        else:
            widget.showNormal()
            menu.resize(menu.width() + 1, menu.height() + 1)
            self.frame.setStyleSheet("background-color: rgb(40, 40, 40);\n"
                                     "border-radius: 10px;\n"
                                     "color:rgb(200, 200, 255);")
            self.maxwin.setText("☐")
            multi = 0

    def fecha(self):
        widget.close()

    def volta(self):
        menu = Menu()
        widget.addWidget(menu)
        widget.setCurrentIndex(widget.currentIndex() + 1)

    def copy(self):
        idcli = self.txt_id.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'SELECT * FROM cliente  WHERE id LIKE ?'
        cursor.execute(consulta, (idcli,))

        dados_lidos = cursor.fetchall()

        self.txt_clinome.setText(dados_lidos[0][1])
        self.txt_clicpf.setText(dados_lidos[0][2])
        self.txt_clirg.setText(dados_lidos[0][3])
        self.txt_cliend.setText(dados_lidos[0][4])
        self.txt_clibairro.setText(dados_lidos[0][5])
        self.txt_clinumero.setText(dados_lidos[0][6])
        self.txt_clicidade.setText(dados_lidos[0][7])
        self.txt_cliestado.setText(dados_lidos[0][8])
        self.txt_clicep.setText(dados_lidos[0][9])
        self.txt_clinacionalidade.setText(dados_lidos[0][10])
        self.txt_cliprofissao.setText(dados_lidos[0][11])
        self.txt_cliestadocivil.setText(dados_lidos[0][12])
        self.txt_clicelular.setText(dados_lidos[0][13])
        self.txt_cliemail.setText(dados_lidos[0][14])
        banco.commit()
        banco.close()

    def delete(self):
        cliid = self.txt_id.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'DELETE FROM cliente WHERE id=?'
        cursor.execute(consulta, (cliid,))

        self.txt_id.setText(None)
        self.txt_clinome.setText(None)
        self.txt_clicpf.setText(None)
        self.txt_clirg.setText(None)
        self.txt_cliend.setText(None)
        self.txt_clibairro.setText(None)
        self.txt_clinumero.setText(None)
        self.txt_clicidade.setText(None)
        self.txt_cliestado.setText(None)
        self.txt_clicep.setText(None)
        self.txt_clinacionalidade.setText(None)
        self.txt_cliprofissao.setText(None)
        self.txt_cliestadocivil.setText(None)
        self.txt_clicelular.setText(None)
        self.txt_cliemail.setText(None)

        cursor.execute('SELECT * FROM cliente')
        dados_lidos = cursor.fetchall()
        self.tabWid_cli.setRowCount(len(dados_lidos))
        self.tabWid_cli.setColumnCount(15)

        for i in range(0, len(dados_lidos)):
            for j in range(0, 15):
                self.tabWid_cli.setItem(i, j, QtWidgets.QTableWidgetItem(str(dados_lidos[i][j])))
        banco.commit()
        banco.close()

    def edit(self):
        cliid = self.txt_id.text()
        clinome = self.txt_clinome.text()
        clicpf = self.txt_clicpf.text()
        clirg = self.txt_clirg.text()
        cliend = self.txt_cliend.text()
        clibairro = self.txt_clibairro.text()
        clinumero = self.txt_clinumero.text()
        clicidade = self.txt_clicidade.text()
        cliestado = self.txt_cliestado.text()
        clicep = self.txt_clicep.text()
        clinacionalidade = self.txt_clinacionalidade.text()
        cliprofissao = self.txt_cliprofissao.text()
        cliestadocivil = self.txt_cliestadocivil.text()
        clicelular = self.txt_clicelular.text()
        cliemail = self.txt_cliemail.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'UPDATE OR IGNORE cliente SET nome=?, cpf=?, rg=?, end=?, bairro=?, num=?, cidade=?, estado=?, cep=?, nacionalidade=?, profissao=?, estadocivil=?, celular=?, email=? WHERE id=?'
        cursor.execute(consulta, (clinome, clicpf, clirg, cliend, clibairro, clinumero, clicidade, cliestado, clicep, clinacionalidade, cliprofissao, cliestadocivil, clicelular, cliemail, cliid))

        self.txt_id.setText(None)
        self.txt_clinome.setText(None)
        self.txt_clicpf.setText(None)
        self.txt_clirg.setText(None)
        self.txt_cliend.setText(None)
        self.txt_clibairro.setText(None)
        self.txt_clinumero.setText(None)
        self.txt_clicidade.setText(None)
        self.txt_cliestado.setText(None)
        self.txt_clicep.setText(None)
        self.txt_clinacionalidade.setText(None)
        self.txt_cliprofissao.setText(None)
        self.txt_cliestadocivil.setText(None)
        self.txt_clicelular.setText(None)
        self.txt_cliemail.setText(None)

        cursor.execute('SELECT * FROM cliente')
        dados_lidos = cursor.fetchall()
        self.tabWid_cli.setRowCount(len(dados_lidos))
        self.tabWid_cli.setColumnCount(15)

        for i in range(0, len(dados_lidos)):
            for j in range(0, 15):
                self.tabWid_cli.setItem(i, j, QtWidgets.QTableWidgetItem(str(dados_lidos[i][j])))
        banco.commit()
        banco.close()

    def add(self):
        clinome = self.txt_clinome.text()
        clicpf = self.txt_clicpf.text()
        clirg = self.txt_clirg.text()
        cliend = self.txt_cliend.text()
        clibairro = self.txt_clibairro.text()
        clinumero = self.txt_clinumero.text()
        clicidade = self.txt_clicidade.text()
        cliestado = self.txt_cliestado.text()
        clicep = self.txt_clicep.text()
        clinacionalidade = self.txt_clinacionalidade.text()
        cliprofissao = self.txt_cliprofissao.text()
        cliestadocivil = self.txt_cliestadocivil.text()
        clicelular = self.txt_clicelular.text()
        cliemail = self.txt_cliemail.text()

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        consulta = 'INSERT OR IGNORE INTO cliente (nome, cpf, rg, end, bairro, num, cidade, estado, cep, nacionalidade, profissao, estadocivil, celular, email) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)'
        cursor.execute(consulta, (clinome, clicpf, clirg, cliend, clibairro, clinumero, clicidade, cliestado, clicep, clinacionalidade, cliprofissao, cliestadocivil, clicelular, cliemail))

        self.txt_id.setText(None)
        self.txt_clinome.setText(None)
        self.txt_clicpf.setText(None)
        self.txt_clirg.setText(None)
        self.txt_cliend.setText(None)
        self.txt_clibairro.setText(None)
        self.txt_clinumero.setText(None)
        self.txt_clicidade.setText(None)
        self.txt_cliestado.setText(None)
        self.txt_clicep.setText(None)
        self.txt_clinacionalidade.setText(None)
        self.txt_cliprofissao.setText(None)
        self.txt_cliestadocivil.setText(None)
        self.txt_clicelular.setText(None)
        self.txt_cliemail.setText(None)

        cursor.execute('SELECT * FROM cliente')
        dados_lidos = cursor.fetchall()
        self.tabWid_cli.setRowCount(len(dados_lidos))
        self.tabWid_cli.setColumnCount(15)

        for i in range(0, len(dados_lidos)):
            for j in range(0, 15):
                self.tabWid_cli.setItem(i, j, QtWidgets.QTableWidgetItem(str(dados_lidos[i][j])))
        banco.commit()
        banco.close()

    def pesqui(self):
        name = self.txt_nome.text().lower()
        for row in range(self.tabWid_cli.rowCount()):
            item = self.tabWid_cli.item(row, 1)
            # if the search is *not* in the item's text *do not hide* the row
            self.tabWid_cli.setRowHidden(row, name not in item.text().lower())

class Doc(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        uic.loadUi("doc.ui", self)
        self.btn_return.clicked.connect(self.volta)
        self.btn_search.clicked.connect(self.searchobra)
        self.btn_search_2.clicked.connect(self.searchcli)
        self.cbox_1cli.setChecked(True)
        self.txt_idcli1.show()
        self.txt_idcli2.hide()
        self.txt_idcli3.hide()
        self.txt_idcli4.hide()
        self.cbox_1cli.clicked.connect(self.checked1)
        self.cbox_2cli.clicked.connect(self.checked2)
        self.cbox_3cli.clicked.connect(self.checked3)
        self.cbox_4cli.clicked.connect(self.checked4)
        self.btn_make.clicked.connect(self.make)
        self.minwin.clicked.connect(self.mini)
        self.maxwin.clicked.connect(self.max)
        self.closewin.clicked.connect(self.fecha)
        self.framelogo.mousePressEvent = self.myfunction
        self.framelogo.mouseMoveEvent = self.myfunc

        banco = sqlite3.connect('./bd/banco.db')
        cursor = banco.cursor()
        # consulta = 'SELECT * FROM cliente'
        cursor.execute('SELECT * FROM obra')
        dados_lidos = cursor.fetchall()
        self.tabWid_obra.setRowCount(len(dados_lidos))
        self.tabWid_obra.setColumnCount(19)

        for i in range(0, len(dados_lidos)):
            for j in range(0, 19):
                self.tabWid_obra.setItem(i, j, QtWidgets.QTableWidgetItem(str(dados_lidos[i][j])))

        cursor.execute('SELECT * FROM cliente')
        dados_lidos = cursor.fetchall()
        self.tabWid_cli.setRowCount(len(dados_lidos))
        self.tabWid_cli.setColumnCount(15)

        for i in range(0, len(dados_lidos)):
            for j in range(0, 15):
                self.tabWid_cli.setItem(i, j, QtWidgets.QTableWidgetItem(str(dados_lidos[i][j])))

        banco.commit()
        banco.close()

        self.setWindowFlags(QtCore.Qt.WindowType.FramelessWindowHint)
        self.gripSize = 16
        self.grips = []
        for i in range(4):
            grip = QSizeGrip(self)
            grip.resize(self.gripSize, self.gripSize)
            self.grips.append(grip)

    def resizeEvent(self, event):
        QMainWindow.resizeEvent(self, event)
        rect = self.rect()
        # top left grip doesn't need to be moved...
        # top right
        self.grips[1].move(rect.right() - self.gripSize, 0)
        # bottom right
        self.grips[2].move(
            rect.right() - self.gripSize, rect.bottom() - self.gripSize)
        # bottom left
        self.grips[3].move(0, rect.bottom() - self.gripSize)

    def myfunction(self, event):
        widget.dragPos = event.globalPosition().toPoint()

    def myfunc(self, event):
        widget.move(widget.pos() + event.globalPosition().toPoint() - widget.dragPos)
        widget.dragPos = event.globalPosition().toPoint()
        event.accept()

    def mini(self):
        widget.showMinimized()

    def max(self):
        global multi
        status = multi
        if status == 0:
            widget.showMaximized()
            self.all.setStyleSheet("background-color: rgb(40, 40, 40);\n"
                                     "border-radius: 0px;\n"
                                     "color:rgb(200, 200, 255);")
            self.maxwin.setText("❐")
            multi = 1
        else:
            widget.showNormal()
            menu.resize(menu.width() + 1, menu.height() + 1)
            self.all.setStyleSheet("background-color: rgb(40, 40, 40);\n"
                                     "border-radius: 10px;\n"
                                     "color:rgb(200, 200, 255);")
            self.maxwin.setText("☐")
            multi = 0

    def fecha(self):
        widget.close()

    def volta(self):
        menu = Menu()
        widget.addWidget(menu)
        widget.setCurrentIndex(widget.currentIndex() + 1)

    def searchobra(self):
        name = self.txt_searchrua.text().lower()
        for row in range(self.tabWid_obra.rowCount()):
            item = self.tabWid_obra.item(row, 1)
            # if the search is *not* in the item's text *do not hide* the row
            self.tabWid_obra.setRowHidden(row, name not in item.text().lower())

    def searchcli(self):
        name = self.txt_searchnome.text().lower()
        for row in range(self.tabWid_cli.rowCount()):
            item = self.tabWid_cli.item(row, 1)
            # if the search is *not* in the item's text *do not hide* the row
            self.tabWid_cli.setRowHidden(row, name not in item.text().lower())


    def checked1(self):
        if self.cbox_1cli.isChecked() == True:
            self.cbox_2cli.setChecked(False)
            self.cbox_3cli.setChecked(False)
            self.cbox_4cli.setChecked(False)
            self.txt_idcli1.show()
            self.txt_idcli2.hide()
            self.txt_idcli3.hide()
            self.txt_idcli4.hide()

    def checked2(self):
        if self.cbox_2cli.isChecked() == True:
            self.cbox_1cli.setChecked(False)
            self.cbox_3cli.setChecked(False)
            self.cbox_4cli.setChecked(False)

            self.txt_idcli1.show()
            self.txt_idcli2.show()
            self.txt_idcli3.hide()
            self.txt_idcli4.hide()

    def checked3(self):
        if self.cbox_3cli.isChecked() == True:
            self.cbox_1cli.setChecked(False)
            self.cbox_2cli.setChecked(False)
            self.cbox_4cli.setChecked(False)

            self.txt_idcli1.show()
            self.txt_idcli2.show()
            self.txt_idcli3.show()
            self.txt_idcli4.hide()

    def checked4(self):
        if self.cbox_4cli.isChecked() == True:
            self.cbox_1cli.setChecked(False)
            self.cbox_2cli.setChecked(False)
            self.cbox_3cli.setChecked(False)

            self.txt_idcli1.show()
            self.txt_idcli2.show()
            self.txt_idcli3.show()
            self.txt_idcli4.show()

    def make(self):
        if self.txt_idobra.text() != "":
            banco = sqlite3.connect('./bd/banco.db')
            cursor = banco.cursor()
            idobra = self.txt_idobra.text()
            consulta = 'SELECT * FROM obra WHERE id=?'
            cursor.execute(consulta, (idobra,))

            dados_lidos = cursor.fetchall()

            endobra = dados_lidos[0][1]
            bairroobra = dados_lidos[0][2]
            numobra = dados_lidos[0][3]
            cidadeobra = dados_lidos[0][4]
            loteobra = dados_lidos[0][5]
            quadraobra = dados_lidos[0][6]
            quarteiraoobra = dados_lidos[0][7]
            tipoobra = dados_lidos[0][8]
            areaobra = dados_lidos[0][9]
            artobra = dados_lidos[0][10]
            valorparcobra = dados_lidos[0][11]
            quantparcobra = dados_lidos[0][12]
            datacontratoobra = dados_lidos[0][13]
            valorvisitaobra = dados_lidos[0][14]
            valorobra = str(float(valorparcobra.replace(",",".")) * float(quantparcobra.replace(",","."))).replace(",",".")
            inscimobobra = dados_lidos[0][15]


            if self.cbox_1cli.isChecked() == True:
                if self.txt_idcli1.text() == "":
                    up.show()
                    pop.lbl_popup.setText("DIGITE O ID DO CLIENTE!")
                    pop.frame_popup.setStyleSheet("background-color: rgb(255, 11, 15);\n"
                                                  "border-radius:5px;")
                    pop.lbl_popup.setStyleSheet("color: rgb(200, 200, 255)")

                elif self.txt_idcli1.text() != "":
                    idcli1 = self.txt_idcli1.text()

                    consultacli = 'SELECT * FROM cliente WHERE id=?'
                    cursor.execute(consultacli, (idcli1,))

                    dados_cli1 = cursor.fetchall()

                    nomecli1 = dados_cli1[0][1]
                    cpfcli1 = dados_cli1[0][2]
                    rgcli1 = dados_cli1[0][3]
                    endcli1 = dados_cli1[0][4]
                    bairrocli1 = dados_cli1[0][5]
                    numcli1 = dados_cli1[0][6]
                    cidadecli1 = dados_cli1[0][7]
                    estadocli1 = dados_cli1[0][8]
                    cepcli1 = dados_cli1[0][9]
                    nacionalidadecli1 = dados_cli1[0][10]
                    profissaocli1 = dados_cli1[0][11]
                    estadocivilcli1 = dados_cli1[0][12]
                    celularcli1 = dados_cli1[0][13]
                    emailcli1 = dados_cli1[0][14]

                    Path(
                        '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + '/' + tipoobra +'/' + ano +   '/Documentos').mkdir(
                        parents=True, exist_ok=True)

                    if self.cbox_contrato.isChecked() == False and self.cbox_recibo.isChecked() == False\
                            and self.cbox_procuracao.isChecked() == False \
                            and self.cbox_reqslei.isChecked() == False and self.cbox_reqclei.isChecked() == False \
                            and self.cbox_memorial.isChecked() == False and self.cbox_memorialcontrucao.isChecked() == False:
                        up.show()
                        pop.lbl_popup.setText("ESCOLHA UM TIPO DE DOCUMENTO!")
                        pop.frame_popup.setStyleSheet("background-color: rgb(255, 11, 15);\n"
                                                      "border-radius:5px;")
                        pop.lbl_popup.setStyleSheet("color: rgb(200, 200, 255)")
                    else:
    # ---------------------Contrato---------------------------------------------------------------------------------------------
                        if self.cbox_contrato.isChecked() == True:
                            document = Document()

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('CONTRATO DE PRESTAÇÃO DE SERVIÇOS TÉCNICOS')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.underline = True
                            font.color.rgb = RGBColor(0, 0, 255)
                            font.name = 'Book Antiqua'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                'Entre as partes abaixo identificadas e no final assinadas fica contratada os serviços de um imóvel mediante as seguintes clausulas e condições:')
                            paragraph.style = document.styles.add_style('style01', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.PARTES')
                            paragraph.style = document.styles.add_style('style0', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.1 CONTRATADO:')
                            paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run(' ROGÉRIO ROCHA SOARES, ').underline = True
                            paragraph.add_run(
                                'brasileiro, casado, Engenheiro civil, portador do RG n° 27.274.081-0 SSP/SP, inscrito no CPF n° 183.125.858-77 e CREA: 5070347192, residente e domiciliado na Rua Ricardo Mendes Horacy, n° 125, Jardim Nossa Senhora Auxiliadora na cidade de Hortolândia/SP. ').bold = False
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.2 CONTRATANTE: ')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run('' + nomecli1 + ', ').underline = True
                            paragraph.add_run(
                                '' + nacionalidadecli1 + ', ' + estadocivilcli1 + ', ' + profissaocli1 + ', '
                                                                                                                              'portador(a) do RG n° ' + rgcli1 + ' SSP/SP, inscrito(a) no CPF n° ' + cpfcli1 + ','
                                                                                                                                                                                                                               ' residente e domiciliado(a) na ' + endcli1 + ', '
                                                                                                                                                                                                                                                                                     'n° ' + numcli1 + ', ' + bairrocli1 + ' na cidade de ' + cidadecli1 + '/' + estadocli1 + '. ').bold = False
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('2.0 OBJETO E FINALIDADES')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('A confecção de projeto arquitetônico e simplificado '
                                                               'residencial em “AUTOCAD”, conforme características do imóvel do CONTRATANTE e '
                                                               'legislações pertinentes nos termos de Leis uso e ocupação de solo do município de '
                                                               + cidadeobra + ' e normas dos órgãos competentes e com aproveitamento para a aprovação de '
                                                                                  'projeto até a liberação do Alvará. Para o imóvel: lote ' + loteobra + ', quadra ' + quadraobra + '; do loteamento '
                                                                                                                                                                                            'denominado “' + bairroobra + '”, no município de ' + cidadeobra + '-SP.')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            font.italic = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '2.2 Atender e fornecer as informações necessárias para o bom andamento dos serviços;')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            paragraph = document.add_paragraph('a)	Documentos necessários;\nb)	Livre acesso ao local.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('3. VISITAS')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '3.1 Este contrato não prevê acompanhamento da execução da obra. \n\n'
                                '3.2 Este contrato prevê 2 (duas) visitas técnicas no local da obra.\n\n'
                                '3.3 As visitas devem ser agendadas com antecedência mínima de 2 (dois) dias e seguir a agenda de atividades do responsável técnico.\n\n'
                                '3.4 Caso houver interesse do CONTRATANTE ')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            paragraph.add_run('CONTRATANTE').bold = True
                            paragraph.add_run('de mais visitas técnicas na obra, será cobrado a parte.\n'
                                              'Valor para cada visita técnica é de R$ ' + valorvisitaobra + ',00 ('+num2words(valorvisitaobra.replace(",","."), lang='pt-br')+') hora.')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('4. PRAZO, PREÇO, FORMA DE PAGAMENTO')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '4.1 O prazo da presente prestação de serviço dá-se início após a assinatura deste contrato e se estendera conforme tramitação legal dos órgãos responsáveis para a liberação dos documentos essenciais com resultado final.\n\n'
                                '4.2 O valor deste contrato é de ')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            runner = paragraph.add_run("" + (valorobra) + ' (' + num2words((valorobra).replace(",","."), lang='pt-br') + ')')
                            runner.bold = True
                            runner.underline = True

                            paragraph.add_run(', que o ')
                            runner = paragraph.add_run('CONTRATANTE ')
                            runner.bold = True
                            paragraph.add_run('se obriga a pagar ao  ')
                            runner = paragraph.add_run('CONTRATADO ')
                            runner.bold = True
                            paragraph.add_run(
                                'em ' + valorparcobra + ',00 (' + num2words(valorparcobra.replace(",","."), lang='pt-br') +') vezes mensais, com vencimento '
                                                                                                   'todo o dia ' + (datacontratoobra[:2]) + ' de cada mês, com início em '
                                + datacontratoobra + ', constituindo-se nenhuma '
                                                       'tolerância de qualquer recebimento depois do '
                                                       'prazo estipulado.\n\n 4.3 Ao ')
                            runner = paragraph.add_run('CONTRATANTE ')
                            paragraph.add_run(
                                'fica ciente que, o pagamento das ou qualquer outro encargo feito através da emissão de cheque e este, por qualquer motivo que seja recusado pela entidade financeira, será considerado nulo, ficando sem efeito de quitação anterior.\n\n4.4 Fica estipulado ao ')
                            runner = paragraph.add_run('CONTRATANTE ')
                            runner.bold = True
                            paragraph.add_run(
                                ' o pagamento de taxas, emissões de documentos, impressões e plotagens   referente as tramitações do processo bem como a exigências dos órgãos envolvidos.')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('5. MULTAS')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('5.1')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            runner = paragraph.add_run('MULTA DE MORA:')
                            runner.bold = True
                            paragraph.add_run('Fica estipulada a multa de ')
                            runner = paragraph.add_run('10%')
                            runner.bold = True
                            paragraph.add_run(
                                '% (dez por cento) calculada sobre o valor do contrato devido à qual será devidamente cobrada juntamente com este pelo não pagamento no prazo previsto neste contrato. O')
                            runner = paragraph.add_run(' CONTRATANTE ')
                            runner.bold = True
                            paragraph.add_run(
                                'fica ciente, ainda que a multa se refere simplesmente a mora, não impedindo a cobrança de outras multas devidas pela inflação ou inexecução de quaisquer das mais clausulas deste contrato e a sua cobrança em meses posteriores. Além dessa multa serão devidos juros e correção monetária, ambos calculados pelo máximo e na forma permitidos pela lei vigente, sempre que houver atraso no pagamento do contrato ou dos encargos.')

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('6. FORO COMPETENTE, DESAPROPRIAÇÂO E AÇOES.')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '6.1 O presente contrato reger-se-á pelo Código Civil Brasileiro e Legislação suplementar, não estabelecendo, em hipóteses alguma, vínculo empregatício entre os contratantes. \n\n'
                                '6.2 Fica eleito o foro')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            runner = paragraph.add_run(' HORTOLÂNDIA – SP')
                            runner.bold = True
                            paragraph.add_run(
                                ', com renúncia de outro por mais privilegio que seja, para quaisquer dúvidas decorrentes deste contrato, sejam quais forem os domicílios dos contratos, mesmo para ações posteriores ao termino deste contrato, inclusive para ações de cobrança, ainda que o ')
                            runner = paragraph.add_run('CONTRATANTE ')
                            runner.bold = True
                            paragraph.add_run('tenham se mudado ou já residam em outra comarca.\n\n'
                                              '6.3 Nos termos do deciso no processo 85.232/88, da Serasa, com força de Provimento, publicado no diário Oficial de 22/6/1995, Cad. I Parte I, se o poder Judiciário informar a ação á Serasa, as providencias pela baixa do nome nesse órgão ficam por conta do devedor, uma vez que não foi a locadora ou sua representante legal que fez a informação.')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                'E assim por estarem justos e contratados assinam o presente em 02 (duas) vias de igual teor e valor.')
                            paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + '/SP, ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '_________________________________          _____________________________________')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.italic = True
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'CONTRATADO:                                            CONTRATANTE:')
                            paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            paragraph = document.add_paragraph('ROGÉRIO ROCHA SOARES                      ' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + '/' + tipoobra +'/' + ano +   '/Documentos/Contrato ' + nomecli1 + '.docx')
            # ---------------------Procuração---------------------------------------------------------------------------------------------
                        if self.cbox_procuracao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('PROCURAÇÃO ')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                'Constitui procurador para a prática de atos perante a Prefeitura Municipal de Hortolândia\n')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(11.5)

                            runner = paragraph.add_run('I - OUTORGANTE:')
                            runner.bold = True

                            paragraph.add_run('\nSr.(a) ' + nomecli1 + ' CPF: ' + cpfcli1 + '\n\n')

                            runner = paragraph.add_run('II – OUTORGADO: ')
                            runner.bold = True

                            paragraph.add_run(
                                '\nSr. Rogério Rocha Soares, portador do CPF nº 183.125.858-77- endereço comercial: Rua Geraldo Denadai, n° 74 1º andar sala 03, Jardim da Paineiras na cidade de Hortolândia/SP.\n\n')

                            runner = paragraph.add_run('III – EXTENSÃO DOS PODERES: ')
                            runner.bold = True

                            paragraph.add_run('\nO OUTORGANTE é proprietário e legítimo dono do imóvel, '
                                              'Lote' + loteobra + 'da Quadra ' + quadraobra + ', localizado no endereço: '
                                                                                                      '' + endobra + ' nº ' + numobra + ' Loteamento: ' + bairroobra + '.\n\n')

                            runner = paragraph.add_run('Específicos para a prática de Ato Determinado. ')
                            runner.bold = True

                            paragraph.add_run('\n(X) Retirar projeto aprovado e Alvará de construção, referente ao imóvel acima. '
                                              '\nPor este instrumento particular de mandato e na melhor forma de direito, o OUTORGANTE acima qualificado, nomeia e constitui o PROCURADOR acima qualificado, a quem confere plenos poderes de representação perante a ')

                            runner = paragraph.add_run('PREFEITURA MUNICIPAL DE' + cidadeobra + ',')
                            runner.bold = True

                            paragraph.add_run(
                                'especialmente para em seu nome e como se o próprio fosse praticar os atos especificados acima.')

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n_________________________________')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('OUTORGANTE:')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(9)
                            font.bold = True

                            paragraph = document.add_paragraph('' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli1 + '')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('OUTORGANTE:')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(9)
                            font.bold = True

                            paragraph = document.add_paragraph('Eng.º Civil: Rogério Rocha Soares')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(10)

                            paragraph = document.add_paragraph(
                                'CPF: 183.125.858-77')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(10)

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + '/' + tipoobra +'/' + ano +   '/Documentos/Procuração ' + nomecli1 + '.docx')
            # ---------------------RRC com lei---------------------------------------------------------------------------------------------
                        if self.cbox_reqclei.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(16)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('DE ACORDO COM A LEI Nª3.491/2018')
                            paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(11)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('\n\nEu')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            runner = paragraph.add_run(' ' + nomecli1 + ' ')
                            runner.bold = True
                            paragraph.add_run(
                                'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

                            paragraph = document.add_paragraph(
                                '                                                                         Nestes Termos,\n'
                                '                                                                         Pede Deferimento.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '_________________________________')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli1 + '')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Dados Complementares:')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Do Proprietário')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Nome:' + nomecli1 + '\n'
                                                                                            'Endereço: ' + endcli1 + ' N°' + numcli1 + '\n'
                                                                                                                                                     'Loteamento:' + bairrocli1 + '\n'
                                                                                                                                                                                    'CEP:' + cepcli1 + '\n'
                                                                                                                                                                                                               'Cidade/Estado:' + cidadecli1 + '-' + estadocli1 + '\n'
                                                                                                                                                                                                                                                                                  'Telefone: ' + celularcli1 + '')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Da Obra')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Endereço: ' + endobra + ' nº ' + numobra +
                                                               'LOTE N° ' + loteobra + '\n'
                                                                                            'Loteamento:' + bairroobra + '\n'
                                                                                                                           'Quadra:' + quadraobra + '')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Do Responsável Técnico')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                                               'CPF: 183.125.858-77\n'
                                                               'Celular: (19) 982009858\n'
                                                               'Inscrição SMPUGE: 1036/18\n'
                                                               'E-mail: rocha.soares@hotmail.com\n')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + '/' + tipoobra +'/' + ano +   '/Documentos/Requerimento sem Lei_' + nomecli1 + '.docx')

                        # ---------------------RRC sem lei---------------------------------------------------------------------------------------------
                        if self.cbox_reqslei.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(1)
                                section.bottom_margin = Cm(1)
                                section.left_margin = Cm(3)
                                section.right_margin = Cm(1)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(1.48), height=Cm(1.48))

                            paragraph = document.add_paragraph('Ao'
                                                               '\nExcelentíssimo Senhor Prefeito Municipal,')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '                    Venho respeitosamente à presença de Vossa Excelência requerer, por meio do representante legal que em conjunto este subscreve, que se digne em providenciar por meio do órgão competente o que segue:')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '(   ) PRED - Desdobro de lote 	(   ) PRED - Regularização de edificação')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            # add grid table

                            table = document.add_table(rows=3, cols=1, style='Table Grid')
                            table.left_margin = Cm(30.4)
                            row = table.rows[0]

                            tabela = 'Dados do requerente (titular do lote ou da edificação)'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(10)
                            tabela_formatada.bold = True

                            tabela = '\nRazão social/nome: '+nomecli1+''\
                                     '\nCNPJ/CPF nº: '+cpfcli1+'' \
                                     '\nE-mail*: '+emailcli1+'' \
                                     '\nTelefone para contato: '+celularcli1+''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(10)

                            tabela = '\n*as notificações sobre este processo serão enviadas por e-mail. Favor atentar-se a isso no momento do preenchimento.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(10)
                            tabela_formatada.bold = True

                            row = table.rows[1]

                            tabela = 'Dados do imóvel:'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(10)
                            tabela_formatada.bold = True

                            tabela = '\nLote/Gleba/Quinhão nº: '+loteobra+'' \
                                     '\nQuadra: '+quadraobra+'' \
                                     '\nLoteamento: '+bairroobra+''\
                                     '\nInscrição Imobiliária: '+inscimobobra+'' \
                                     '\nEndereço: '+endobra+'' \
                                     '\nCEP: '+cepcli1+''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(10)

                            row = table.rows[2]

                            tabela = 'Dados do Responsável Técnico pelo projeto'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(10)
                            tabela_formatada.bold = True

                            tabela = '\nNome completo: Rogério Rocha Soares' \
                                     '\nRegistro profissional: 5070374192 Órgão: CREA' \
                                     '\nEstá registrado no CPHO¹?  ( X ) sim     (   ) não' \
                                     '\nNº da Inscrição Mobiliária: 1036/18' \
                                     '\nE-mail²: rocha.soares@hotmail.com' \
                                     '\nTelefone para contato: (19)982009858' \
                                     '\n¹CPHO - Cadastro de Profissionais Habilitados junto aos órgãos da Prefeitura Municipal de Hortolândia.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(10)

                            tabela = '\n²as notificações sobre este processo serão enviadas por e-mail. Favor atentar-se a isso no momento do preenchimento.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(10)
                            tabela_formatada.bold = True

                            paragraph = document.add_paragraph(
                                '\n(  ) Declaro que os documentos, declarações e demais elementos submetidos na instrução deste requerimento são verdadeiros e que tenho ciência de que a falsidade de qualquer informação prestada acarreta automaticamente em crime de falsidade ideológica na forma do art. 299 do Código Penal Brasileiro.')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.line_spacing = Cm(0)
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph(
                                '(  ) Declaro ter ciência de que, caso meu pedido não seja instruído nos termos que determina a legislação vigente, deverei regularizá-lo no prazo de 30 (trinta) dias corridos, sob pena de arquivamento e indeferimento deste processo.')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.line_spacing = Cm(0)
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph(
                                '(  ) Declaro ter ciência do prazo de 180 (cento e oitenta) dias corridos, contados da entrega da planta aprovada, para o registro dos desdobros e das edificações junto ao Cartório de Registro de Imóveis competente.')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('        	Nestes termos,')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph('        	Peço Deferimento.')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('Hortolândia, '+dia+' de '+mesescrito+' de '+ano+'. ')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('\n\n____________________________________________ '
                                                               '\nProprietário'
                                                               '\n\n\n____________________________________________'
                                                               '\nResponsável técnico')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + '/' + tipoobra +'/' + ano +   '/Documentos/Requerimento com Lei_' + nomecli1 + '.docx')

            # ---------------------Declaração---------------------------------------------------------------------------------------------
                        if self.cbox_declaracao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(3.49)
                                section.bottom_margin = Cm(1.1)
                                section.left_margin = Cm(2)
                                section.right_margin = Cm(2)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo_branco.jpg", width=Cm(3.5), height=Cm(2.65))

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/Logo_Hortolandia.png", width=Cm(4.9), height=Cm(2.65))

                            paragraph = document.add_paragraph('ANEXO I')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('DECLARAÇÃO')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('Em conformidade com o disposto no artigo 4º '
                                                               'da Lei Municipal nº 2.529, de 04 de abril de 2011,'
                                                               ' que “Dispõe sobre controle ambiental para utilização'
                                                               ' de produtos e subprodutos de madeira de origem nativa'
                                                               ' em obras e serviços de Engenharia Civil no Município'
                                                               ' de Hortolândia”, eu,' + nomecli1 + ', (' + profissaocli1 + '), ')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run('Proprietário da obra ')
                            runner.bold = True

                            paragraph.add_run(
                                'localizada à ' + endobra + ', nº ' + numobra + ' Lote ' + loteobra + ', Quadra ' + quadraobra + ', Loteamento ' + bairroobra + ',' \
                                                                                                                                                                                  'cidade de Hortolândia-SP, DECLARO estar ciente das disposições ' \
                                                                                                                                                                                  'constantes da Instrução Normativa nº 112, de 21 de agosto de 2006, ' \
                                                                                                                                                                                  'do Instituto Brasileiro do Meio Ambiente e dos Recursos Naturais ' \
                                                                                                                                                                                  'Renováveis - IBAMA, e me comprometendo a adquirir, para qualquer ' \
                                                                                                                                                                                  'serviço a ser realizado no referido imóvel, produtos e subprodutos ' \
                                                                                                                                                                                  'de madeira de origem não nativa ou nativa que tenha procedência ' \
                                                                                                                                                                                  'legal, decorrentes de desmatamento autorizado ou de manejo florestal' \
                                                                                                                                                                                  ' aprovado por órgão ambiental competente, integrante do Sistema ' \
                                                                                                                                                                                  'Nacional do Meio Ambiente – SISNAMA, com autorização de transporte ' \
                                                                                                                                                                                  'reconhecida pelo órgão ambiental competente, exigindo no ato da ' \
                                                                                                                                                                                  'compra que as empresas que comercializem madeiras, forneçam o DOF' \
                                                                                                                                                                                  ' (Documento de Origem Florestal), acompanhado de nota fiscal. ')

                            runner = paragraph.add_run('DECLARO ')
                            runner.bold = True

                            paragraph.add_run(
                                'ainda, que quando da solicitação do Habite-se o proprietário deverá apresentar além dos documentos, declarações e comprovantes exigidos pelo poder Público, ')

                            runner = paragraph.add_run('um dos seguintes documentos: ')
                            runner.bold = True

                            paragraph = document.add_paragraph(
                                '       ●      Nota Fiscal constando o número do DOF, em caso de utilização de produtos ou subprodutos de madeira de ')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run('aaaaaaaaaaa')
                            font = runner.font
                            font.color.rgb = RGBColor(255, 255, 255)

                            paragraph.add_run('origem nativa;')

                            paragraph = document.add_paragraph(
                                '       ●      Declaração que fez a reutilização de madeira ou que utilizou madeira de reflorestamento;')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '       ●     Declaração que não utilizou madeira de origem nativa, por ter utilizado novas tecnologias ou produtos alternativos.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('PROPRIETÁRIO')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('\n\n\n\n\n\n\n\nEm conformidade com o disposto no artigo 4º da '
                                                               'Lei Municipal nº 2.529, de 04 de abril de 2011, '
                                                               'que “Dispõe sobre controle ambiental para utilização'
                                                               ' de produtos e subprodutos de madeira de origem nativa'
                                                               ' em obras e serviços de Engenharia Civil no Município '
                                                               'de Hortolândia”, eu, Rogério Rocha Soares, (Engenheiro Civil), '
                                                               'Autor do Projeto da obra localizada à '
                                                               'Rua ' + endobra + ', nº ' + numobra + ' Lote ' + loteobra + ', Quadra ' + quadraobra + ', Loteamento ' + bairroobra + ',cidade de Hortolândia-SP,'
                                                                                                                                                                                                        ' DECLARO estar ciente das disposições constantes da Instrução Normativa'
                                                                                                                                                                                                        ' nº 112, de 21 de agosto de 2006, do Instituto Brasileiro '
                                                                                                                                                                                                        'do Meio Ambiente e dos Recursos Naturais Renováveis - IBAMA, e ')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run(
                                'me comprometendo a instruir meu cliente a adquirir produtos e subprodutos de madeira de origem não nativa ou nativa que tenha procedência legal, ')
                            runner.underline = True

                            paragraph.add_run(
                                'decorrentes de desmatamento autorizado ou de manejo florestal aprovado por órgão ambiental competente, integrante do Sistema Nacional do Meio Ambiente – SISNAMA, com autorização de transporte reconhecida pelo órgão ambiental competente, exigindo no ato da compra que as empresas que comercializem madeiras, forneçam o DOF (Documento de Origem Florestal), acompanhado de nota fiscal.')

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('AUTOR DO PROJETO')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            # runner_word.size = Pt(10)

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + '/' + tipoobra +'/' + ano +   '/Documentos/Declaração ' + nomecli1 + '.docx')

            # ---------------------MEMORIAL DESCRITIVO PARA CONSTRUÇÃO---------------------------------------------------------------------------------------------
                        if self.cbox_memorialcontrucao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('MEMORIAL DESCRITIVO PARA CONSTRUÇÃO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('Obra: REGULARIZAÇÃO E CONSTRUÇÃO RESIDENCIAL MULTIFAMILIAR – R2')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('(DE ACORDO COM A LEI 3.491/2018 - ANISTIA)')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('Local: ' + endobra + '- N° ' + numobra + '')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Lote: ' + loteobra + ' Quadra: ' + quadraobra + '')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Loteamento: ' + loteobra + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Município: ' + cidadeobra + '/SP')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Proprietário: ' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Responsável Técnico: Eng. Civil ROGÉRIO ROCHA SOARES')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph(
                                '01 – Preparação do terreno: Raspagem com moto-niveladora e os resíduos será retirado e depositado em local devidamente autorizado e legalizado pela Prefeitura Municipal.')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '02 – Fundações: A fundação do alicerçada em estacas manuais, blocos e vigas baldrame em concreto armado.')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '03 – Impermeabilização: Com Vedacit no concreto e na argamassa de embasamento. Será aplicado neutrol no baldrame antes do reaterro.')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '04 – Estrutura: Os pilares, vergas, contra-vergas, vigas de apoio e vigas de respaldo em concreto armado batido na própria obra.')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '05 - Alvenaria: As paredes externas e internas executadas com bloco cerâmico (barro) de oito furos (15 cm), assentados com argamassa de areia e cimento. A altura do pé direito será de 2,80m em toda a casa.')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('06 - Forro: O forro de laje pré-moldada de concreto.')
                            paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '07- Cobertura: O telhado foi executado com telhas de barro estrutura de madeira com inclinação de 26% na residência.')
                            paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '08 - Instalações hidráulicas: Instalado conforme normas da SABESP e NBR 7229/93,')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        a) - Água fria: Abastecida pela rede pública e armazenada em dois reservatórios elevados, com capacidade de 1.000 litros d’agua na residência;')
                            paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        b) - Esgotos sanitários: canalizados com tubos de PVC com diâmetro 100 mm para a rede pública de afastamento de esgoto.')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        c) – Águas pluviais: coletada por meio de sistema de calhas galvanizadas e desce por condutores e são canalizadas até a sarjeta por meio de tubos de PVC, por sob a calçada.')
                            paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '09 - Instalações elétricas: Instalado conforme normas da CPFL, composto de um ponto de luz para cada compartimento e tomadas em número suficiente para atender a demanda.')
                            paragraph.style = document.styles.add_style('style21', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '10 – Revestimento: Os banheiros da residência e da dependência em azulejo até a altura do forro;')
                            paragraph.style = document.styles.add_style('style22', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '11 - Pisos: Nivelado com argamassa de areia e cimento e revestimento cerâmico em todos os compartimentos.')
                            paragraph.style = document.styles.add_style('style23', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '12 - Esquadrias: As Janelas e portas externas são em alumínio e as portas internas de madeira, conforme as medidas indicadas no projeto.')
                            paragraph.style = document.styles.add_style('style24', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '13 – Pintura: As paredes de alvenaria, internas pintadas com tinta látex PVA; as paredes externas pintadas com tinta látex, acrílica, as janelas e portas serão pintadas com esmalte sintético.')
                            paragraph.style = document.styles.add_style('style25', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '14 – Limpeza Geral da Obra: Após a conclusão da obra, foi feita a retirada dos restos de materiais e entulho.')
                            paragraph.style = document.styles.add_style('style26', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style27', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n_________________________________                    _____________________________________')
                            paragraph.style = document.styles.add_style('style28', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'ROGÉRIO ROCHA SOARES                                        Proprietário: ' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style29', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Engenheiro Civil                                                            CPF:' + cpfcli1 + '')
                            paragraph.style = document.styles.add_style('style30', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CREA: 5070347192-SP')
                            paragraph.style = document.styles.add_style('style31', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'SMPUGE: 1036/18')
                            paragraph.style = document.styles.add_style('style32', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + '/' + tipoobra +'/' + ano +   '/Documentos/Memorial Descritivo Para Construção ' + nomecli1 + '.docx')
# ---------------------MEMORIAL DESCRITIVO---------------------------------------------------------------------------------------------
                        if self.cbox_memorial.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(1.27)
                                section.left_margin = Cm(1.27)
                                section.right_margin = Cm(1.27)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('MEMORIAL DESCRITIVO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('Obra: Regularização Residencial Unifamiliar (R1)\n'
                                                               'Local: ' + endobra + ', nº ' + numobra + ' Lote: ' + loteobra + ' – Quadra: ' + quadraobra + '\n'
                                                                                                                                                                           'Loteamento: ' + bairroobra + ' -  ' + cidadeobra + ' - SP\n'
                                                                                                                                                                                                                                       'Proprietário: ')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)
                            runner = paragraph.add_run('' + nomecli1 + ' \n')
                            runner.bold = True
                            paragraph.add_run('                     CPF:')
                            runner = paragraph.add_run('' + cpfcli1 + ' \n')
                            runner.bold = True

                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('ÁREA DO TERRENO – ' + areaobra + ' m² ')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)

                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('Descrição')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                'MOVIMENTO DE TERRA: Será realizada uma limpeza da superfície do terreno, tal como remoção da camada vegetal.')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'TIPO DE FUNDAÇÃO: Serão executadas brocas de concreto armado, posteriormente, sapatas na cabeça das brocas, as quais apoiarão as vigas baldrames, que por sua vez embasam a alvenaria. Todos os componentes da fundação serão executados e concretados “in loco”.')
                            paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'IMPERMEABILIZAÇÃO: Para evitar os fenômenos de capilaridade e percolação (umidade nas partes inferiores da alvenaria), todas as vigas baldrames e primeiras fiadas da alvenaria receberão tratamento contra a umidade proveniente do solo.')
                            paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ESTRUTURA: Será utilizadas estrutura convencional de concreto armado: pilares e vigas. As peças serão devidamente moldadas e concretadas “in loco”.')
                            paragraph.style = document.styles.add_style('style5.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ALVENARIA: Será executada alvenaria com blocos cerâmicos assentados com argamassa mista de cimento cal e areia. Serão executadas também vergas e contravergas nos vãos de janelas e portas.')
                            paragraph.style = document.styles.add_style('style5.5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'FORRO: Será executado em laje pré-fabricada de concreto armado em todas as dependências da edificação, com uso de vigas pré-fabricadas e lajotas cerâmicas.')
                            paragraph.style = document.styles.add_style('style5.6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'COBERTURA: A cobertura será composta por telhas Metálicas com estrutura de aço; o telhado apresentará inclinação média de 11%.')
                            paragraph.style = document.styles.add_style('style5.7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'REVESTIMENTO PAREDES: O revestimento das paredes de dependências consideradas molhadas será executado com revestimento cerâmico até o teto. Nas demais paredes, o revestimento será constituído de chapisco grosso e emboço com argamassa mista de cimento e areia.')
                            paragraph.style = document.styles.add_style('style5.8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PISOS E RODAPÉS: O revestimento do piso será de cerâmica em todas as dependências internas; externamente o piso será revestido de cimentado desempenado.')
                            paragraph.style = document.styles.add_style('style5.9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PORTAS: As portas, tanto internas quanto externas, serão de madeira e alumínio, todas com as devidas ferragens e em bom estado de conservação.')
                            paragraph.style = document.styles.add_style('style5.10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'JANELAS: Todas as janelas serão de alumínio, de correr ou de abrir, estas também, devidamente tratadas e envernizadas.')
                            paragraph.style = document.styles.add_style('style5.11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ÁGUAS PLUVIAIS: Serão coletadas e conduzidas à via pública. A captação será feita pelo telhado que por sua vez conduzirá estas águas até as calhas de onde irão para a via pública, por meio de tubos de PVC, as águas lançadas no quintal irão para via pública.')
                            paragraph.style = document.styles.add_style('style5.12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'INSTALAÇÕES ELÉTRICAS: Serão executadas de acordo com as Normas Técnicas. Serão utilizados fios e cabos de cobre, cujas secções atenderão as necessidades a que serão submetidos. Todas a instalações serão instaladas em perfeitas condições de uso. O medidor de entrada de energia será executado e instalado de acordo com as necessidades e exigências da concessionária fornecedora deste serviço.')
                            paragraph.style = document.styles.add_style('style5.13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'INSTALAÇÕES HIDRO-SANITÁRIAS: Estas instalações serão dimensionadas e executadas de acordo com as Normas Técnicas. Sendo que todos os aparelhos hidrossanitários estarão em funcionamento correto. Serão utilizadas e executadas caixas de inspeção e de gordura nos pontos necessários para boa manutenção e funcionamento destas instalações. Para condução de água (fria) potável, serão utilizados tubos de PVC marrom soldável. O sistema de esgoto e águas servidas tem seu lançamento para rede publica de esgoto.')
                            paragraph.style = document.styles.add_style('style5.14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PINTURA: A pintura será executada em látex em todas as paredes, tanto internas quando externas. As esquadrias e caixilharias em geral, como já descrito, serão devidamente impermeabilizadas e protegidas contra as intempéries.')
                            paragraph.style = document.styles.add_style('style5.15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'LIMPEZA: A obra será totalmente limpa de entulhos. A edificação será apta a ser habitada, obedecendo às condições mínimas de conforto, segurança e habitabilidade.')
                            paragraph.style = document.styles.add_style('style5.16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________                                    _____________________________________')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'ROGÉRIO ROCHA SOARES                                                       Proprietário:' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Engenheiro Civil                                                                            CPF:' + cpfcli1 + '')
                            paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CREA: 5070347192-SP')
                            paragraph.style = document.styles.add_style('style21', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'ART' + artobra + '')
                            paragraph.style = document.styles.add_style('style22', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + '/' + tipoobra +'/' + ano +   '/Documentos/Memorial Descritivo ' + nomecli1 + '.docx')

                            # ---------------------Recibo---------------------------------------------------------------------------------------------
                        if self.cbox_recibo.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('RECIBO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            font.size = Pt(16)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('Eu ROGÉRIO ROCHA SOARES engenheiro civil CREA: '
                                                               '5070374192, recebi de ' + nomecli1 + ', '
                                                                                                             'parte do pagamento para aprovação de projeto '
                                                                                                             'arquitetônico a quantia de R$ ' + valorparcobra + ',00 (' + num2words(valorparcobra.replace(",","."), lang='pt-br') +
                                                               '), de um total de '
                                                               'R$ ' + valorobra + ' (' + num2words((valorobra).replace(",",".")) + ').')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            font.size = Pt(12)

                            paragraph = document.add_paragraph(
                                '\n\n' + cidadeobra + ' ' + dia + ' / ' + mesescrito + ' / ' + ano + '.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            font.size = Pt(12)

                            # footer section
                            footer_section = document.sections[0]
                            footer = footer_section.footer

                            # footer text
                            footer_text = footer.paragraphs[0]
                            footer_text.text = "_______________________________________________________________________________________________" \
                                               "\nROGÉRIO ROCHA SOARES   TEL. : (19) 982009858" \
                                               "\nCREA: 5070347192" \
                                               "\nE-MAIL: rocha.soares@hotmail.com"
                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + '/' + tipoobra +'/' + ano +   '/Documentos/Recibo ' + nomecli1 + '.docx')

                        up.show()
                        pop.lbl_popup.setText("DOCUMENTOS CRIADOS")
                        pop.frame_popup.setStyleSheet("background-color: rgb(57, 173, 84);\n"
                                                      "border-radius:5px;")
                        pop.lbl_popup.setStyleSheet("color: rgb(35, 35, 35)")

                        self.txt_idobra.setText(None)
                        self.txt_idcli1.setText(None)

            elif self.cbox_2cli.isChecked() == True:
                if self.txt_idcli1.text() == "" or self.txt_idcli2.text() == "":
                    up.show()
                    pop.lbl_popup.setText("DIGITE O ID DO CLIENTE!")
                    pop.frame_popup.setStyleSheet("background-color: rgb(255, 11, 15);\n"
                                                  "border-radius:5px;")
                    pop.lbl_popup.setStyleSheet("color: rgb(200, 200, 255)")
                else:
                    idcli1 = self.txt_idcli1.text()

                    consultacli = 'SELECT * FROM cliente WHERE id=?'
                    cursor.execute(consultacli, (idcli1,))

                    dados_cli1 = cursor.fetchall()

                    nomecli1 = dados_cli1[0][1]
                    cpfcli1 = dados_cli1[0][2]
                    rgcli1 = dados_cli1[0][3]
                    endcli1 = dados_cli1[0][4]
                    bairrocli1 = dados_cli1[0][5]
                    numcli1 = dados_cli1[0][6]
                    cidadecli1 = dados_cli1[0][7]
                    estadocli1 = dados_cli1[0][8]
                    cepcli1 = dados_cli1[0][9]
                    nacionalidadecli1 = dados_cli1[0][10]
                    profissaocli1 = dados_cli1[0][11]
                    estadocivilcli1 = dados_cli1[0][12]
                    celularcli1 = dados_cli1[0][13]
                    emailcli1 = dados_cli1[0][14]

                    idcli2 = self.txt_idcli2.text()
                    cursor.execute(consultacli, (idcli2,))

                    dados_cli2 = cursor.fetchall()

                    nomecli2 = dados_cli2[0][1]
                    cpfcli2 = dados_cli2[0][2]
                    rgcli2 = dados_cli2[0][3]
                    endcli2 = dados_cli2[0][4]
                    bairrocli2 = dados_cli2[0][5]
                    numcli2 = dados_cli2[0][6]
                    cidadecli2 = dados_cli2[0][7]
                    estadocli2 = dados_cli2[0][8]
                    cepcli2 = dados_cli2[0][9]
                    nacionalidadecli2 = dados_cli2[0][10]
                    profissaocli2 = dados_cli2[0][11]
                    estadocivilcli2 = dados_cli2[0][12]
                    celularcli2 = dados_cli2[0][13]
                    emailcli2 = dados_cli2[0][14]

                    Path(
                        '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ' e ' + nomecli2 + '/' + tipoobra +  '/' + ano + '/Documentos''').mkdir(
                        parents=True, exist_ok=True)

                    if self.cbox_contrato.isChecked() == False and self.cbox_recibo.isChecked() == False \
                            and self.cbox_procuracao.isChecked() == False \
                            and self.cbox_reqslei.isChecked() == False and self.cbox_reqclei.isChecked() == False \
                            and self.cbox_memorial.isChecked() == False and self.cbox_memorialcontrucao.isChecked() == False:
                            up.show()
                            pop.lbl_popup.setText("ESCOLHA UM TIPO DE DOCUMENTO!")
                            pop.frame_popup.setStyleSheet("background-color: rgb(255, 11, 15);\n"
                                                      "border-radius:5px;")
                            pop.lbl_popup.setStyleSheet("color: rgb(200, 200, 255)")
                    else:
    # ---------------------Contrato---------------------------------------------------------------------------------------------
                        if self.cbox_contrato.isChecked() == True:
                            document = Document()

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('CONTRATO DE PRESTAÇÃO DE SERVIÇOS TÉCNICOS')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.underline = True
                            font.color.rgb = RGBColor(0, 0, 255)
                            font.name = 'Book Antiqua'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                'Entre as partes abaixo identificadas e no final assinadas fica contratada os serviços de um imóvel mediante as seguintes clausulas e condições:')
                            paragraph.style = document.styles.add_style('style01', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.PARTES')
                            paragraph.style = document.styles.add_style('style0', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.1 CONTRATADO:')
                            paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run(' ROGÉRIO ROCHA SOARES, ').underline = True
                            paragraph.add_run(
                                'brasileiro, casado, Engenheiro civil, portador do RG n° 27.274.081-0 SSP/SP, inscrito no CPF n° 183.125.858-77 e CREA: 5070347192, residente e domiciliado na Rua Ricardo Mendes Horacy, n° 125, Jardim Nossa Senhora Auxiliadora na cidade de Hortolândia/SP. ').bold = False
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.1 CONTRATANTE: ')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run('' + nomecli1 + ', ').underline = True
                            paragraph.add_run(
                                '' + nacionalidadecli1 + ', ' + estadocivilcli1 + ', ' + profissaocli1 + ', '
                                                                                                                              'portador(a) do RG n° ' + rgcli1 + ' SSP/SP, inscrito(a) no CPF n° ' + cpfcli1 + ''
                                                                                                                                                                                                                               ', residente e domiciliado(a) na ' + endcli1 + ', '
                                                                                                                                                                                                                                                                                      'n° ' + numcli1 + ', ' + bairrocli1 + ' na cidade de ' + cidadecli1 + '/' + estadocli1 + '. ').bold = False

                            paragraph = document.add_paragraph('    CONTRATANTE: ')
                            paragraph.style = document.styles.add_style('style2.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run('' + nomecli2 + ', ').underline = True
                            paragraph.add_run(
                                '' + nacionalidadecli2 + ', ' + estadocivilcli2 + ', ' + profissaocli2 + ', '
                                                                                                                              'portador(a) do RG n° ' + rgcli2 + ' SSP/SP, inscrito(a) no CPF n° ' + cpfcli2 + ''
                                                                                                                                                                                                                               ', residente e domiciliado(a) na ' + endcli2 + ', '
                                                                                                                                                                                                                                                                                      'n° ' + numcli2 + ', ' + bairrocli2 + ' na cidade de ' + cidadecli2 + '/' + estadocli2 + '. ').bold = False

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('2.0 OBJETO E FINALIDADES')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('A confecção de projeto arquitetônico e simplificado '
                                                               'residencial em “AUTOCAD”, conforme características do imóvel dos CONTRATANTES e '
                                                               'legislações pertinentes nos termos de Leis uso e ocupação de solo do município de '
                                                               + cidadeobra + ' e normas dos órgãos competentes e com aproveitamento para a aprovação de '
                                                                                  'projeto até a liberação do Alvará. Para o imóvel: lote ' + loteobra + ', quadra ' + quadraobra + '; do loteamento '
                                                                                                                                                                                            'denominado “' + bairroobra + '”, no município de ' + cidadeobra + '-SP.')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            font.italic = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '2.2 Atender e fornecer as informações necessárias para o bom andamento dos serviços;')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            paragraph = document.add_paragraph('a)	Documentos necessários;\nb)	Livre acesso ao local.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('3. VISITAS')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '3.1 Este contrato não prevê acompanhamento da execução da obra. \n\n'
                                '3.2 Este contrato prevê 2 (duas) visitas técnicas no local da obra.\n\n'
                                '3.3 As visitas devem ser agendadas com antecedência mínima de 2 (dois) dias e seguir a agenda de atividades do responsável técnico.\n\n'
                                '3.4 Caso houver interesse dos ')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            paragraph.add_run('CONTRATANTES').bold = True
                            paragraph.add_run('de mais visitas técnicas na obra, será cobrado a parte.\n'
                                              'Valor para cada visita técnica é de R$ ' + valorvisitaobra + ',00 (' + num2words(valorvisitaobra.replace(",","."), lang='pt-br') + ') hora.')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('4. PRAZO, PREÇO, FORMA DE PAGAMENTO')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '4.1 O prazo da presente prestação de serviço dá-se início após a assinatura deste contrato e se estendera conforme tramitação legal dos órgãos responsáveis para a liberação dos documentos essenciais com resultado final.\n\n'
                                '4.2 O valor deste contrato é de ')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            runner = paragraph.add_run("" + valorobra + ' (' + num2words((valorobra).replace(",","."), lang='pt-br') + ')')
                            runner.bold = True
                            runner.underline = True

                            paragraph.add_run(', que o ')
                            runner = paragraph.add_run('CONTRATANTE ')
                            runner.bold = True
                            paragraph.add_run('se obriga a pagar ao  ')
                            runner = paragraph.add_run('CONTRATADO ')
                            runner.bold = True
                            paragraph.add_run(
                                'em ' + valorparcobra + ',00 (' + num2words(valorparcobra.replace(",","."), lang='pt-br') + ') vezes mensais, com vencimento '
                                                                                                   'todo o dia ' + (datacontratoobra[:2]) + ' de cada mês, com início em '
                                + datacontratoobra + ', constituindo-se nenhuma '
                                                       'tolerância de qualquer recebimento depois do '
                                                       'prazo estipulado.\n\n 4.3 Ao ')
                            runner = paragraph.add_run('CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run(
                                'fica ciente que, o pagamento das ou qualquer outro encargo feito através da emissão de cheque e este, por qualquer motivo que seja recusado pela entidade financeira, será considerado nulo, ficando sem efeito de quitação anterior.\n\n4.4 Fica estipulado aos ')
                            runner = paragraph.add_run('CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run(
                                ' o pagamento de taxas, emissões de documentos, impressões e plotagens   referente as tramitações do processo bem como a exigências dos órgãos envolvidos.')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('5. MULTAS')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('5.1 ')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            runner = paragraph.add_run('MULTA DE MORA: ')
                            runner.bold = True
                            paragraph.add_run('Fica estipulada a multa de ')
                            runner = paragraph.add_run('10%')
                            runner.bold = True
                            paragraph.add_run(
                                '% (dez por cento) calculada sobre o valor do contrato devido à qual será devidamente cobrada juntamente com este pelo não pagamento no prazo previsto neste contrato. Os')
                            runner = paragraph.add_run(' CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run(
                                'fica ciente, ainda que a multa se refere simplesmente a mora, não impedindo a cobrança de outras multas devidas pela inflação ou inexecução de quaisquer das mais clausulas deste contrato e a sua cobrança em meses posteriores. Além dessa multa serão devidos juros e correção monetária, ambos calculados pelo máximo e na forma permitidos pela lei vigente, sempre que houver atraso no pagamento do contrato ou dos encargos.')

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('6. FORO COMPETENTE, DESAPROPRIAÇÂO E AÇOES.')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '6.1 O presente contrato reger-se-á pelo Código Civil Brasileiro e Legislação suplementar, não estabelecendo, em hipóteses alguma, vínculo empregatício entre os contratantes. \n\n'
                                '6.2 Fica eleito o foro')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            runner = paragraph.add_run(' HORTOLÂNDIA – SP')
                            runner.bold = True
                            paragraph.add_run(
                                ', com renúncia de outro por mais privilegio que seja, para quaisquer dúvidas decorrentes deste contrato, sejam quais forem os domicílios dos contratos, mesmo para ações posteriores ao termino deste contrato, inclusive para ações de cobrança, ainda que os ')
                            runner = paragraph.add_run('CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run('tenham se mudado ou já residam em outra comarca.\n\n'
                                              '6.3 Nos termos do deciso no processo 85.232/88, da Serasa, com força de Provimento, publicado no diário Oficial de 22/6/1995, Cad. I Parte I, se o poder Judiciário informar a ação á Serasa, as providencias pela baixa do nome nesse órgão ficam por conta do devedor, uma vez que não foi a locadora ou sua representante legal que fez a informação.')

                            paragraph = document.add_paragraph(
                                'E assim por estarem justos e contratados assinam o presente em 02 (duas) vias de igual teor e valor.')
                            paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ' / SP, ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '_________________________________          _____________________________________')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.italic = True
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'CONTRATADO:                                            CONTRATANTE:')
                            paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'ROGÉRIO ROCHA SOARES                      ' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '_________________________________')
                            paragraph.style = document.styles.add_style('style17.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.italic = True
                            font.bold = True

                            paragraph = document.add_paragraph('CONTRATANTE:')
                            paragraph.style = document.styles.add_style('style18.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            paragraph = document.add_paragraph('' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style19.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ' e ' + nomecli2 + '/' + tipoobra +  '/' + ano + '/Documentos/Contrato ' + nomecli1 + ' e ' + nomecli2 + '.docx')

                        if self.cbox_memorial.isChecked() == True:
                            # Memorial Descritivo
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(1.27)
                                section.left_margin = Cm(1.27)
                                section.right_margin = Cm(1.27)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('MEMORIAL DESCRITIVO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('Obra: Regularização Residencial Unifamiliar (R1)\n'
                                                               'Local: ' + endobra + ', nº ' + numobra + ' Lote: ' + loteobra + ' – Quadra: ' + quadraobra + '\n'
                                                                                                                                                                           'Loteamento: ' + bairroobra + ' -  ' + cidadeobra + ' - SP\n'
                                                                                                                                                                                                                                       'Proprietário: ')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)
                            runner = paragraph.add_run('' + nomecli1 + ' \n')
                            runner.bold = True
                            paragraph.add_run('                     CPF:')
                            runner = paragraph.add_run('' + cpfcli1 + ' \n')
                            runner.bold = True

                            runner = paragraph.add_run('                     ' + nomecli2 + ' \n')
                            runner.bold = True
                            paragraph.add_run('                     CPF:')
                            runner = paragraph.add_run('' + cpfcli2 + ' \n')
                            runner.bold = True

                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('ÁREA DO TERRENO – ' + areaobra + ' m² ')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)

                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('Descrição')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                'MOVIMENTO DE TERRA: Será realizada uma limpeza da superfície do terreno, tal como remoção da camada vegetal.')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'TIPO DE FUNDAÇÃO: Serão executadas brocas de concreto armado, posteriormente, sapatas na cabeça das brocas, as quais apoiarão as vigas baldrames, que por sua vez embasam a alvenaria. Todos os componentes da fundação serão executados e concretados “in loco”.')
                            paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'IMPERMEABILIZAÇÃO: Para evitar os fenômenos de capilaridade e percolação (umidade nas partes inferiores da alvenaria), todas as vigas baldrames e primeiras fiadas da alvenaria receberão tratamento contra a umidade proveniente do solo.')
                            paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ESTRUTURA: Será utilizadas estrutura convencional de concreto armado: pilares e vigas. As peças serão devidamente moldadas e concretadas “in loco”.')
                            paragraph.style = document.styles.add_style('style5.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ALVENARIA: Será executada alvenaria com blocos cerâmicos assentados com argamassa mista de cimento cal e areia. Serão executadas também vergas e contravergas nos vãos de janelas e portas.')
                            paragraph.style = document.styles.add_style('style5.5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'FORRO: Será executado em laje pré-fabricada de concreto armado em todas as dependências da edificação, com uso de vigas pré-fabricadas e lajotas cerâmicas.')
                            paragraph.style = document.styles.add_style('style5.6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'COBERTURA: A cobertura será composta por telhas Metálicas com estrutura de aço; o telhado apresentará inclinação média de 11%.')
                            paragraph.style = document.styles.add_style('style5.7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'REVESTIMENTO PAREDES: O revestimento das paredes de dependências consideradas molhadas será executado com revestimento cerâmico até o teto. Nas demais paredes, o revestimento será constituído de chapisco grosso e emboço com argamassa mista de cimento e areia.')
                            paragraph.style = document.styles.add_style('style5.8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PISOS E RODAPÉS: O revestimento do piso será de cerâmica em todas as dependências internas; externamente o piso será revestido de cimentado desempenado.')
                            paragraph.style = document.styles.add_style('style5.9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PORTAS: As portas, tanto internas quanto externas, serão de madeira e alumínio, todas com as devidas ferragens e em bom estado de conservação.')
                            paragraph.style = document.styles.add_style('style5.10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'JANELAS: Todas as janelas serão de alumínio, de correr ou de abrir, estas também, devidamente tratadas e envernizadas.')
                            paragraph.style = document.styles.add_style('style5.11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ÁGUAS PLUVIAIS: Serão coletadas e conduzidas à via pública. A captação será feita pelo telhado que por sua vez conduzirá estas águas até as calhas de onde irão para a via pública, por meio de tubos de PVC, as águas lançadas no quintal irão para via pública.')
                            paragraph.style = document.styles.add_style('style5.12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'INSTALAÇÕES ELÉTRICAS: Serão executadas de acordo com as Normas Técnicas. Serão utilizados fios e cabos de cobre, cujas secções atenderão as necessidades a que serão submetidos. Todas a instalações serão instaladas em perfeitas condições de uso. O medidor de entrada de energia será executado e instalado de acordo com as necessidades e exigências da concessionária fornecedora deste serviço.')
                            paragraph.style = document.styles.add_style('style5.13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'INSTALAÇÕES HIDRO-SANITÁRIAS: Estas instalações serão dimensionadas e executadas de acordo com as Normas Técnicas. Sendo que todos os aparelhos hidrossanitários estarão em funcionamento correto. Serão utilizadas e executadas caixas de inspeção e de gordura nos pontos necessários para boa manutenção e funcionamento destas instalações. Para condução de água (fria) potável, serão utilizados tubos de PVC marrom soldável. O sistema de esgoto e águas servidas tem seu lançamento para rede publica de esgoto.')
                            paragraph.style = document.styles.add_style('style5.14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PINTURA: A pintura será executada em látex em todas as paredes, tanto internas quando externas. As esquadrias e caixilharias em geral, como já descrito, serão devidamente impermeabilizadas e protegidas contra as intempéries.')
                            paragraph.style = document.styles.add_style('style5.15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'LIMPEZA: A obra será totalmente limpa de entulhos. A edificação será apta a ser habitada, obedecendo às condições mínimas de conforto, segurança e habitabilidade.')
                            paragraph.style = document.styles.add_style('style5.16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________                                    _____________________________________')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário:' + nomecli1 + '                                                Proprietário:' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli1 + '                                                                        CPF:' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style17.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'ROGÉRIO ROCHA SOARES')
                            paragraph.style = document.styles.add_style('style19.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Engenheiro Civil')
                            paragraph.style = document.styles.add_style('style20.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CREA: 5070347192-SP')
                            paragraph.style = document.styles.add_style('style21.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'ART' + artobra + '')
                            paragraph.style = document.styles.add_style('style22.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ' e ' + nomecli2 + '/' + tipoobra +  '/' + ano + '/Documentos/Memorial Descritivo' + nomecli1 + ' e ' + nomecli2 + '.docx')

                            # RRC sem lei
                        if self.cbox_reqslei.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(1)
                                section.bottom_margin = Cm(1)
                                section.left_margin = Cm(3)
                                section.right_margin = Cm(1)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(1.48), height=Cm(1.48))

                            paragraph = document.add_paragraph('Ao'
                                                               '\nExcelentíssimo Senhor Prefeito Municipal,')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '                    Venho respeitosamente à presença de Vossa Excelência requerer, por meio do representante legal que em conjunto este subscreve, que se digne em providenciar por meio do órgão competente o que segue:')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '(   ) PRED - Desdobro de lote 	(   ) PRED - Regularização de edificação')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            # add grid table

                            table = document.add_table(rows=3, cols=1, style='Table Grid')
                            table.left_margin = Cm(30.4)
                            row = table.rows[0]

                            tabela = 'Dados dos requerentes (titulares do lote ou da edificação)'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            tabela = '\nRazão social/nome: ' + nomecli1 + '' \
                                     '\nCNPJ/CPF nº: ' + cpfcli1 + '' \
                                     '\nE-mail*: ' + emailcli1 + '' \
                                     '\nTelefone para contato: ' + celularcli1 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\naaaaaaaaaaaaaaaaaa'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(3)
                            font = tabela_formatada.font
                            font.color.rgb = RGBColor(255, 255, 255)

                            tabela = '\nRazão social/nome: ' + nomecli2 + '' \
                                     '\nCNPJ/CPF nº: ' + cpfcli2 + '' \
                                     '\nE-mail*: ' + emailcli2 + '' \
                                     '\nTelefone para contato: ' + celularcli2 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\n*as notificações sobre este processo serão enviadas por e-mail. Favor atentar-se a isso no momento do preenchimento.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            row = table.rows[1]

                            tabela = 'Dados do imóvel:'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            tabela = '\nLote/Gleba/Quinhão nº: ' + loteobra + '' \
                                                                              '\nQuadra: ' + quadraobra + '' \
                                                                                                          '\nLoteamento: ' + bairroobra + '' \
                                                                                                                                          '\nInscrição Imobiliária: ' + inscimobobra + '' \
                                                                                                                                                                                       '\nEndereço: ' + endobra + '' \
                                                                                                                                                                                                                  '\nCEP: ' + cepcli1 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            row = table.rows[2]

                            tabela = 'Dados do Responsável Técnico pelo projeto'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            tabela = '\nNome completo: Rogério Rocha Soares' \
                                     '\nRegistro profissional: 5070374192 Órgão: CREA' \
                                     '\nEstá registrado no CPHO¹?  ( X ) sim     (   ) não' \
                                     '\nNº da Inscrição Mobiliária: 1036/18' \
                                     '\nE-mail²: rocha.soares@hotmail.com' \
                                     '\nTelefone para contato: (19)982009858' \
                                     '\n¹CPHO - Cadastro de Profissionais Habilitados junto aos órgãos da Prefeitura Municipal de Hortolândia.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\n²as notificações sobre este processo serão enviadas por e-mail. Favor atentar-se a isso no momento do preenchimento.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            paragraph = document.add_paragraph(
                                '\n(X) Declaro que os documentos, declarações e demais elementos submetidos na instrução deste requerimento são verdadeiros e que tenho ciência de que a falsidade de qualquer informação prestada acarreta automaticamente em crime de falsidade ideológica na forma do art. 299 do Código Penal Brasileiro.')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.line_spacing = Cm(0)
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph(
                                '(X) Declaro ter ciência de que, caso meu pedido não seja instruído nos termos que determina a legislação vigente, deverei regularizá-lo no prazo de 30 (trinta) dias corridos, sob pena de arquivamento e indeferimento deste processo.')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.line_spacing = Cm(0)
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph(
                                '(X) Declaro ter ciência do prazo de 180 (cento e oitenta) dias corridos, contados da entrega da planta aprovada, para o registro dos desdobros e das edificações junto ao Cartório de Registro de Imóveis competente.')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('        	Nestes termos,')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph('        	Peço Deferimento.')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'Hortolândia, ' + dia + ' de ' + mesescrito + ' de ' + ano + '. ')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('\n\n____________________________________________    ____________________________________________ '
                                                               '\nProprietário 1                                                                    Proprietário 2'
                                                               '\n\n\n____________________________________________'
                                                               '\nResponsável técnico')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ' e ' + nomecli2 + '/' + tipoobra +  '/' + ano + '/Documentos/Requerimento sem Lei ' + nomecli1 + ' e ' + nomecli2 + '.docx')

                            # RRC com lei
                        if self.cbox_reqclei.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(16)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('DE ACORDO COM A LEI Nª3.491/2018')
                            paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(11)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('\n\nNós,')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            runner = paragraph.add_run(' ' + nomecli1 + ' e ' + nomecli2 + ' ')
                            runner.bold = True
                            paragraph.add_run(
                                'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

                            paragraph = document.add_paragraph(
                                '                                                                         Nestes Termos,\n'
                                '                                                                         Pede Deferimento.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n_________________________________                            _________________________________')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                '' + nomecli1 + '                                                           ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: ' + cpfcli1 + '                                                              CPF: ' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Dados Complementares:')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Dos Proprietários')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Nome:' + nomecli1 + '\n'
                                                                                            'Endereço: ' + endcli1 + ' N°' + numcli1 + '\n'
                                                                                                                                                     'Loteamento:' + bairrocli1 + '\n'
                                                                                                                                                                                    'CEP:' + cepcli1 + '\n'
                                                                                                                                                                                                               'Cidade/Estado:' + cidadecli1 + '-' + estadocli1 + '\n'
                                                                                                                                                                                                                                                                                  'Telefone: ' + celularcli1 + '')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Nome:' + nomecli2 + '\n'
                                                                                            'Endereço: ' + endcli2 + ' N°' + numcli2 + '\n'
                                                                                                                                                     'Loteamento:' + bairrocli2 + '\n'
                                                                                                                                                                                    'CEP:' + cepcli2 + '\n'
                                                                                                                                                                                                               'Cidade/Estado:' + cidadecli2 + '-' + estadocli2 + '\n'
                                                                                                                                                                                                                                                                                  'Telefone: ' + celularcli2 + '')
                            paragraph.style = document.styles.add_style('style10.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Da Obra')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Endereço: ' + endobra + ' nº ' + numobra +
                                                               '\nLOTE N° ' + loteobra + 'Loteamento:' + bairroobra + '\n'
                                                                                                                           'Quadra:' + quadraobra + '')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('\n\nDo Responsável Técnico')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                                               'CPF: 183.125.858-77\n'
                                                               'Celular: (19) 982009858\n'
                                                               'Inscrição SMPUGE: 1036/18\n'
                                                               'E-mail: rocha.soares@hotmail.com\n')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ' e ' + nomecli2 + '/' + tipoobra +  '/' + ano + '/Documentos/Requerimento com Lei ' + nomecli1 + ' e ' + nomecli2 + '.docx')

                            # Procuração
                        if self.cbox_procuracao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('PROCURAÇÃO ')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                'Constitui procurador para a prática de atos perante a Prefeitura Municipal de Hortolândia\n')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(11.5)

                            runner = paragraph.add_run('I - OUTORGANTES:')
                            runner.bold = True

                            paragraph.add_run('\nSr.(a) ' + nomecli1 + ' CPF: ' + cpfcli1 + '')

                            paragraph.add_run('\nSr.(a) ' + nomecli2 + ' CPF: ' + cpfcli2 + '\n\n')

                            runner = paragraph.add_run('II – OUTORGADO: ')
                            runner.bold = True

                            paragraph.add_run(
                                '\nSr. Rogério Rocha Soares, portador do CPF nº 183.125.858-77- endereço comercial: Rua Geraldo Denadai, n° 74 1º andar sala 03, Jardim da Paineiras na cidade de Hortolândia/SP.\n\n')

                            runner = paragraph.add_run('III – EXTENSÃO DOS PODERES: ')
                            runner.bold = True

                            paragraph.add_run('\nO OUTORGANTE é proprietário e legítimo dono do imóvel, '
                                              'Lote' + loteobra + 'da Quadra ' + quadraobra + ', localizado no endereço: '
                                                                                                      '' + endobra + ' nº ' + numobra + ' Loteamento: ' + bairroobra + '.\n\n')

                            runner = paragraph.add_run('Específicos para a prática de Ato Determinado. ')
                            runner.bold = True

                            paragraph.add_run(
                                '\n(X) Retirar projeto aprovado e Alvará de construção, referente ao imóvel acima. '
                                '\nPor este instrumento particular de mandato e na melhor forma de direito, o OUTORGANTE acima qualificado, nomeia e constitui o PROCURADOR acima qualificado, a quem confere plenos poderes de representação perante a ')

                            runner = paragraph.add_run('PREFEITURA MUNICIPAL DE ' + cidadeobra + ',')
                            runner.bold = True

                            paragraph.add_run(
                                'especialmente para em seu nome e como se o próprio fosse praticar os atos especificados acima.')

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n_________________________________                            _________________________________')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'OUTORGANTE:                                                                                                OUTORGANTE:')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(9)
                            font.bold = True

                            paragraph = document.add_paragraph(
                                '' + nomecli1 + '                                                           ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: ' + cpfcli1 + '                                                              CPF: ' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('OUTORGADO:')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(9)
                            font.bold = True

                            paragraph = document.add_paragraph('Eng.º Civil: Rogério Rocha Soares')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(10)

                            paragraph = document.add_paragraph(
                                'CPF: 183.125.858-77')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(10)

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ' e ' + nomecli2 + '/' + tipoobra +  '/' + ano + '/Documentos/Procuração ' + nomecli1 + ' e ' + nomecli2 + '.docx')

                            # Declaração
                        if self.cbox_declaracao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(3.49)
                                section.bottom_margin = Cm(1.1)
                                section.left_margin = Cm(2)
                                section.right_margin = Cm(2)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo_branco.jpg", width=Cm(3.5), height=Cm(2.65))

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/Logo_Hortolandia.png", width=Cm(4.9), height=Cm(2.65))

                            paragraph = document.add_paragraph('ANEXO I')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('DECLARAÇÃO')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('Em conformidade com o disposto no artigo 4º '
                                                               'da Lei Municipal nº 2.529, de 04 de abril de 2011,'
                                                               ' que “Dispõe sobre controle ambiental para utilização'
                                                               ' de produtos e subprodutos de madeira de origem nativa'
                                                               ' em obras e serviços de Engenharia Civil no Município'
                                                               ' de Hortolândia”, nós, ' + nomecli1 + ', (' + profissaocli1 + ') e ' + nomecli2 + ', (' + profissaocli2 + '),')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run(' Proprietários da obra ')
                            runner.bold = True

                            paragraph.add_run(
                                'localizada à ' + endobra + ', nº ' + numobra + ' Lote ' + loteobra + ', Quadra ' + quadraobra + ', Loteamento ' + bairroobra + ',' \
                                                                                                                                                                                  'cidade de Hortolândia-SP, DECLARAMOS estar ciente das disposições ' \
                                                                                                                                                                                  'constantes da Instrução Normativa nº 112, de 21 de agosto de 2006, ' \
                                                                                                                                                                                  'do Instituto Brasileiro do Meio Ambiente e dos Recursos Naturais ' \
                                                                                                                                                                                  'Renováveis - IBAMA, e me comprometendo a adquirir, para qualquer ' \
                                                                                                                                                                                  'serviço a ser realizado no referido imóvel, produtos e subprodutos ' \
                                                                                                                                                                                  'de madeira de origem não nativa ou nativa que tenha procedência ' \
                                                                                                                                                                                  'legal, decorrentes de desmatamento autorizado ou de manejo florestal' \
                                                                                                                                                                                  ' aprovado por órgão ambiental competente, integrante do Sistema ' \
                                                                                                                                                                                  'Nacional do Meio Ambiente – SISNAMA, com autorização de transporte ' \
                                                                                                                                                                                  'reconhecida pelo órgão ambiental competente, exigindo no ato da ' \
                                                                                                                                                                                  'compra que as empresas que comercializem madeiras, forneçam o DOF' \
                                                                                                                                                                                  ' (Documento de Origem Florestal), acompanhado de nota fiscal. ')

                            runner = paragraph.add_run('DECLARAMOS ')
                            runner.bold = True

                            paragraph.add_run(
                                'ainda, que quando da solicitação do Habite-se o proprietário deverá apresentar além dos documentos, declarações e comprovantes exigidos pelo poder Público, ')

                            runner = paragraph.add_run('um dos seguintes documentos: ')
                            runner.bold = True

                            paragraph = document.add_paragraph(
                                '       ●      Nota Fiscal constando o número do DOF, em caso de utilização de produtos ou subprodutos de madeira de ')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run('aaaaaaaaaaa')
                            font = runner.font
                            font.color.rgb = RGBColor(255, 255, 255)

                            paragraph.add_run('origem nativa;')

                            paragraph = document.add_paragraph(
                                '       ●      Declaração que fez a reutilização de madeira ou que utilizou madeira de reflorestamento;')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '       ●     Declaração que não utilizou madeira de origem nativa, por ter utilizado novas tecnologias ou produtos alternativos.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_____________________________________')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('PROPRIETÁRIO ' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_____________________________________')
                            paragraph.style = document.styles.add_style('style8.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('PROPRIETÁRIO ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style9.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('\n\n\n\n\n\n\n\nEm conformidade com o disposto no artigo 4º da '
                                                               'Lei Municipal nº 2.529, de 04 de abril de 2011, '
                                                               'que “Dispõe sobre controle ambiental para utilização'
                                                               ' de produtos e subprodutos de madeira de origem nativa'
                                                               ' em obras e serviços de Engenharia Civil no Município '
                                                               'de Hortolândia”, eu, Rogério Rocha Soares, (Engenheiro Civil), '
                                                               'Autor do Projeto da obra localizada à '
                                                               'Rua ' + endobra + ', nº ' + numobra + ' Lote ' + loteobra + ', Quadra ' + quadraobra + ', Loteamento ' + bairroobra + ',cidade de Hortolândia-SP,'
                                                                                                                                                                                                        ' DECLARO estar ciente das disposições constantes da Instrução Normativa'
                                                                                                                                                                                                        ' nº 112, de 21 de agosto de 2006, do Instituto Brasileiro '
                                                                                                                                                                                                        'do Meio Ambiente e dos Recursos Naturais Renováveis - IBAMA, e ')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run(
                                'me comprometendo a instruir meu cliente a adquirir produtos e subprodutos de madeira de origem não nativa ou nativa que tenha procedência legal, ')
                            runner.underline = True

                            paragraph.add_run(
                                'decorrentes de desmatamento autorizado ou de manejo florestal aprovado por órgão ambiental competente, integrante do Sistema Nacional do Meio Ambiente – SISNAMA, com autorização de transporte reconhecida pelo órgão ambiental competente, exigindo no ato da compra que as empresas que comercializem madeiras, forneçam o DOF (Documento de Origem Florestal), acompanhado de nota fiscal.')

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('AUTOR DO PROJETO')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            # runner_word.size = Pt(10)

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ' e ' + nomecli2 + '/' + tipoobra +  '/' + ano + '/Documentos/Declaração ' + nomecli1 + ' e ' + nomecli2 + '.docx')

                            # ---------------------MEMORIAL DESCRITIVO PARA CONSTRUÇÃO---------------------------------------------------------------------------------------------
                        if self.cbox_memorialcontrucao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('MEMORIAL DESCRITIVO PARA CONSTRUÇÃO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                'Obra: REGULARIZAÇÃO E CONSTRUÇÃO RESIDENCIAL MULTIFAMILIAR – R2')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('(DE ACORDO COM A LEI 3.491/2018 - ANISTIA)')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('Local: ' + endobra + '- N° ' + numobra + '')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Lote: ' + loteobra + ' Quadra: ' + quadraobra + '')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Loteamento: ' + bairroobra + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Município: ' + cidadeobra + '/SP')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph(
                                'Proprietário(s): ' + nomecli1 + ' e ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Responsável Técnico: Eng. Civil ROGÉRIO ROCHA SOARES')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph(
                                '01 – Preparação do terreno: Raspagem com moto-niveladora e os resíduos será retirado e depositado em local devidamente autorizado e legalizado pela Prefeitura Municipal.')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '02 – Fundações: A fundação do alicerçada em estacas manuais, blocos e vigas baldrame em concreto armado.')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '03 – Impermeabilização: Com Vedacit no concreto e na argamassa de embasamento. Será aplicado neutrol no baldrame antes do reaterro.')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '04 – Estrutura: Os pilares, vergas, contra-vergas, vigas de apoio e vigas de respaldo em concreto armado batido na própria obra.')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '05 - Alvenaria: As paredes externas e internas executadas com bloco cerâmico (barro) de oito furos (15 cm), assentados com argamassa de areia e cimento. A altura do pé direito será de 2,80m em toda a casa.')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('06 - Forro: O forro de laje pré-moldada de concreto.')
                            paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '07- Cobertura: O telhado foi executado com telhas de barro estrutura de madeira com inclinação de 26% na residência.')
                            paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '08 - Instalações hidráulicas: Instalado conforme normas da SABESP e NBR 7229/93,')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        a) - Água fria: Abastecida pela rede pública e armazenada em dois reservatórios elevados, com capacidade de 1.000 litros d’agua na residência;')
                            paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        b) - Esgotos sanitários: canalizados com tubos de PVC com diâmetro 100 mm para a rede pública de afastamento de esgoto.')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        c) – Águas pluviais: coletada por meio de sistema de calhas galvanizadas e desce por condutores e são canalizadas até a sarjeta por meio de tubos de PVC, por sob a calçada.')
                            paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '09 - Instalações elétricas: Instalado conforme normas da CPFL, composto de um ponto de luz para cada compartimento e tomadas em número suficiente para atender a demanda.')
                            paragraph.style = document.styles.add_style('style21', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '10 – Revestimento: Os banheiros da residência e da dependência em azulejo até a altura do forro;')
                            paragraph.style = document.styles.add_style('style22', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '11 - Pisos: Nivelado com argamassa de areia e cimento e revestimento cerâmico em todos os compartimentos.')
                            paragraph.style = document.styles.add_style('style23', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '12 - Esquadrias: As Janelas e portas externas são em alumínio e as portas internas de madeira, conforme as medidas indicadas no projeto.')
                            paragraph.style = document.styles.add_style('style24', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '13 – Pintura: As paredes de alvenaria, internas pintadas com tinta látex PVA; as paredes externas pintadas com tinta látex, acrílica, as janelas e portas serão pintadas com esmalte sintético.')
                            paragraph.style = document.styles.add_style('style25', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '14 – Limpeza Geral da Obra: Após a conclusão da obra, foi feita a retirada dos restos de materiais e entulho.')
                            paragraph.style = document.styles.add_style('style26', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style27', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n_________________________________                    _____________________________________')
                            paragraph.style = document.styles.add_style('style28', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário: ' + nomecli1 + '                               Proprietário: ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style29', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli1 + '                                                        CPF:' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style30', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style28.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'ROGÉRIO ROCHA SOARES')
                            paragraph.style = document.styles.add_style('style29.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Engenheiro Civil')
                            paragraph.style = document.styles.add_style('style30.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CREA: 5070347192-SP')
                            paragraph.style = document.styles.add_style('style31.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'SMPUGE: 1036/18')
                            paragraph.style = document.styles.add_style('style32.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ' e ' + nomecli2 + '/' + tipoobra +  '/' + ano + '/Documentos/Memorial Descritivo Para Construção ' + nomecli1 + ' e ' + nomecli2 + '.docx')

                            # ---------------------Recibo---------------------------------------------------------------------------------------------
                        if self.cbox_recibo.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('RECIBO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            font.size = Pt(16)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('Eu ROGÉRIO ROCHA SOARES engenheiro civil CREA: '
                                                               '5070374192, recebi de ' + nomecli1 + ' e ' + nomecli2 + ', '
                                                                                                                                        'parte do pagamento para aprovação de projeto '
                                                                                                                                        'arquitetônico a quantia de R$ ' + valorparcobra + ',00 (' + num2words(valorparcobra.replace(",","."), lang='pt-br') +
                                                               '), de um total de '
                                                               'R$ ' + valorobra + ',00 (' + num2words((valorobra).replace(",","."), lang='pt-br')  + ').')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            font.size = Pt(12)

                            paragraph = document.add_paragraph(
                                '\n\n' + cidadeobra + ', ' + dia + ' / ' + mesescrito + ' / ' + ano + '.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            font.size = Pt(12)

                            # footer section
                            footer_section = document.sections[0]
                            footer = footer_section.footer

                            # footer text
                            footer_text = footer.paragraphs[0]
                            footer_text.text = "_______________________________________________________________________________________________" \
                                               "\nROGÉRIO ROCHA SOARES   TEL. : (19) 982009858" \
                                               "\nCREA: 5070347192" \
                                               "\nE-MAIL: rocha.soares@hotmail.com"
                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ' e ' + nomecli2 + '/' + tipoobra +  '/' + ano + '/Documentos/Recibo ' + nomecli1 + ' e ' + nomecli2 + '.docx')
                        up.show()
                        pop.lbl_popup.setText("DOCUMENTOS CRIADOS")
                        pop.frame_popup.setStyleSheet("background-color: rgb(57, 173, 84);\n"
                                                      "border-radius:5px;")
                        pop.lbl_popup.setStyleSheet("color: rgb(35, 35, 35)")

                        self.txt_idobra.setText(None)
                        self.txt_idcli1.setText(None)
                        self.txt_idcli2.setText(None)

            elif self.cbox_3cli.isChecked() == True:
                if self.txt_idcli1.text() == "" or self.txt_idcli2.text() == "" or self.txt_idcli3.text() == "":
                    up.show()
                    pop.lbl_popup.setText("DIGITE O ID DO CLIENTE!")
                    pop.frame_popup.setStyleSheet("background-color: rgb(255, 11, 15);\n"
                                                  "border-radius:5px;")
                    pop.lbl_popup.setStyleSheet("color: rgb(200, 200, 255)")
                else:
                    idcli1 = self.txt_idcli1.text()

                    consultacli = 'SELECT * FROM cliente WHERE id=?'
                    cursor.execute(consultacli, (idcli1,))

                    dados_cli1 = cursor.fetchall()

                    nomecli1 = dados_cli1[0][1]
                    cpfcli1 = dados_cli1[0][2]
                    rgcli1 = dados_cli1[0][3]
                    endcli1 = dados_cli1[0][4]
                    bairrocli1 = dados_cli1[0][5]
                    numcli1 = dados_cli1[0][6]
                    cidadecli1 = dados_cli1[0][7]
                    estadocli1 = dados_cli1[0][8]
                    cepcli1 = dados_cli1[0][9]
                    nacionalidadecli1 = dados_cli1[0][10]
                    profissaocli1 = dados_cli1[0][11]
                    estadocivilcli1 = dados_cli1[0][12]
                    celularcli1 = dados_cli1[0][13]
                    emailcli1 = dados_cli1[0][14]

                    idcli2 = self.txt_idcli2.text()
                    cursor.execute(consultacli, (idcli2,))

                    dados_cli2 = cursor.fetchall()

                    nomecli2 = dados_cli2[0][1]
                    cpfcli2 = dados_cli2[0][2]
                    rgcli2 = dados_cli2[0][3]
                    endcli2 = dados_cli2[0][4]
                    bairrocli2 = dados_cli2[0][5]
                    numcli2 = dados_cli2[0][6]
                    cidadecli2 = dados_cli2[0][7]
                    estadocli2 = dados_cli2[0][8]
                    cepcli2 = dados_cli2[0][9]
                    nacionalidadecli2 = dados_cli2[0][10]
                    profissaocli2 = dados_cli2[0][11]
                    estadocivilcli2 = dados_cli2[0][12]
                    celularcli2 = dados_cli2[0][13]
                    emailcli2 = dados_cli2[0][14]

                    idcli3 = self.txt_idcli3.text()
                    cursor.execute(consultacli, (idcli3,))

                    dados_cli3 = cursor.fetchall()

                    nomecli3 = dados_cli3[0][1]
                    cpfcli3 = dados_cli3[0][2]
                    rgcli3 = dados_cli3[0][3]
                    endcli3 = dados_cli3[0][4]
                    bairrocli3 = dados_cli3[0][5]
                    numcli3 = dados_cli3[0][6]
                    cidadecli3 = dados_cli3[0][7]
                    estadocli3 = dados_cli3[0][8]
                    cepcli3 = dados_cli3[0][9]
                    nacionalidadecli3 = dados_cli3[0][10]
                    profissaocli3 = dados_cli3[0][11]
                    estadocivilcli3 = dados_cli3[0][12]
                    celularcli3 = dados_cli3[0][13]
                    emailcli3 = dados_cli3[0][14]

                    Path(
                        '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', ' + nomecli2 + ' e '+ nomecli3 +'/' + tipoobra + '/' + ano + '/Documentos''').mkdir(
                        parents=True, exist_ok=True)

                    if self.cbox_contrato.isChecked() == False and self.cbox_recibo.isChecked() == False \
                            and self.cbox_procuracao.isChecked() == False \
                            and self.cbox_reqslei.isChecked() == False and self.cbox_reqclei.isChecked() == False \
                            and self.cbox_memorial.isChecked() == False and self.cbox_memorialcontrucao.isChecked() == False:
                        up.show()
                        pop.lbl_popup.setText("ESCOLHA UM TIPO DE DOCUMENTO!")
                        pop.frame_popup.setStyleSheet("background-color: rgb(255, 11, 15);\n"
                                                      "border-radius:5px;")
                        pop.lbl_popup.setStyleSheet("color: rgb(200, 200, 255)")
                    else:
                        # ---------------------Contrato---------------------------------------------------------------------------------------------
                        if self.cbox_contrato.isChecked() == True:
                            document = Document()

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('CONTRATO DE PRESTAÇÃO DE SERVIÇOS TÉCNICOS')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.underline = True
                            font.color.rgb = RGBColor(0, 0, 255)
                            font.name = 'Book Antiqua'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                'Entre as partes abaixo identificadas e no final assinadas fica contratada os serviços de um imóvel mediante as seguintes clausulas e condições:')
                            paragraph.style = document.styles.add_style('style01', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.PARTES')
                            paragraph.style = document.styles.add_style('style0', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.1 CONTRATADO:')
                            paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run(' ROGÉRIO ROCHA SOARES, ').underline = True
                            paragraph.add_run(
                                'brasileiro, casado, Engenheiro civil, portador do RG n° 27.274.081-0 SSP/SP, inscrito no CPF n° 183.125.858-77 e CREA: 5070347192, residente e domiciliado na Rua Ricardo Mendes Horacy, n° 125, Jardim Nossa Senhora Auxiliadora na cidade de Hortolândia/SP. ').bold = False
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.1 CONTRATANTE: ')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run('' + nomecli1 + ', ').underline = True
                            paragraph.add_run(
                                '' + nacionalidadecli1 + ', ' + estadocivilcli1 + ', ' + profissaocli1 + ', '
                                                                                                         'portador(a) do RG n° ' + rgcli1 + ' SSP/SP, inscrito(a) no CPF n° ' + cpfcli1 + ''
                                                                                                                                                                                          ', residente e domiciliado(a) na ' + endcli1 + ', '
                                                                                                                                                                                                                                         'n° ' + numcli1 + ', ' + bairrocli1 + ' na cidade de ' + cidadecli1 + '/' + estadocli1 + '. ').bold = False

                            paragraph = document.add_paragraph('    CONTRATANTE: ')
                            paragraph.style = document.styles.add_style('style2.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run('' + nomecli2 + ', ').underline = True
                            paragraph.add_run(
                                '' + nacionalidadecli2 + ', ' + estadocivilcli2 + ', ' + profissaocli2 + ', '
                                                                                                         'portador(a) do RG n° ' + rgcli2 + ' SSP/SP, inscrito(a) no CPF n° ' + cpfcli2 + ''
                                                                                                                                                                                          ', residente e domiciliado(a) na ' + endcli2 + ', '
                                                                                                                                                                                                                                         'n° ' + numcli2 + ', ' + bairrocli2 + ' na cidade de ' + cidadecli2 + '/' + estadocli2 + '. ').bold = False

                            paragraph = document.add_paragraph('    CONTRATANTE: ')
                            paragraph.style = document.styles.add_style('style2.22', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run('' + nomecli3 + ', ').underline = True
                            paragraph.add_run(
                                '' + nacionalidadecli3 + ', ' + estadocivilcli3 + ', ' + profissaocli3 + ', '
                                                                                                         'portador(a) do RG n° ' + rgcli3 + ' SSP/SP, inscrito(a) no CPF n° ' + cpfcli3 + ''
                                                                                                                                                                                          ', residente e domiciliado(a) na ' + endcli3 + ', '
                                                                                                                                                                                                                                         'n° ' + numcli3 + ', ' + bairrocli3 + ' na cidade de ' + cidadecli3 + '/' + estadocli3 + '. ').bold = False

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('2.0 OBJETO E FINALIDADES')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('A confecção de projeto arquitetônico e simplificado '
                                                               'residencial em “AUTOCAD”, conforme características do imóvel dos CONTRATANTES e '
                                                               'legislações pertinentes nos termos de Leis uso e ocupação de solo do município de '
                                                               + cidadeobra + ' e normas dos órgãos competentes e com aproveitamento para a aprovação de '
                                                                              'projeto até a liberação do Alvará. Para o imóvel: lote ' + loteobra + ', quadra ' + quadraobra + '; do loteamento '
                                                                                                                                                                                'denominado “' + bairroobra + '”, no município de ' + cidadeobra + '-SP.')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            font.italic = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '2.2 Atender e fornecer as informações necessárias para o bom andamento dos serviços;')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            paragraph = document.add_paragraph(
                                'a)	Documentos necessários;\nb)	Livre acesso ao local.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('3. VISITAS')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '3.1 Este contrato não prevê acompanhamento da execução da obra. \n\n'
                                '3.2 Este contrato prevê 2 (duas) visitas técnicas no local da obra.\n\n'
                                '3.3 As visitas devem ser agendadas com antecedência mínima de 2 (dois) dias e seguir a agenda de atividades do responsável técnico.\n\n'
                                '3.4 Caso houver interesse dos ')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            paragraph.add_run('CONTRATANTES').bold = True
                            paragraph.add_run('de mais visitas técnicas na obra, será cobrado a parte.\n'
                                              'Valor para cada visita técnica é de R$ ' + valorvisitaobra + ',00 (' + num2words(
                                valorvisitaobra.replace(",", "."), lang='pt-br') + ') hora.')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('4. PRAZO, PREÇO, FORMA DE PAGAMENTO')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '4.1 O prazo da presente prestação de serviço dá-se início após a assinatura deste contrato e se estendera conforme tramitação legal dos órgãos responsáveis para a liberação dos documentos essenciais com resultado final.\n\n'
                                '4.2 O valor deste contrato é de ')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            runner = paragraph.add_run(
                                "" + valorobra + ' (' + num2words((valorobra).replace(",", "."), lang='pt-br') + ')')
                            runner.bold = True
                            runner.underline = True

                            paragraph.add_run(', que o ')
                            runner = paragraph.add_run('CONTRATANTE ')
                            runner.bold = True
                            paragraph.add_run('se obriga a pagar ao  ')
                            runner = paragraph.add_run('CONTRATADO ')
                            runner.bold = True
                            paragraph.add_run(
                                'em ' + valorparcobra + ',00 (' + num2words(valorparcobra.replace(",", "."),
                                                                            lang='pt-br') + ') vezes mensais, com vencimento '
                                                                                            'todo o dia ' + (
                                datacontratoobra[:2]) + ' de cada mês, com início em '
                                + datacontratoobra + ', constituindo-se nenhuma '
                                                     'tolerância de qualquer recebimento depois do '
                                                     'prazo estipulado.\n\n 4.3 Ao ')
                            runner = paragraph.add_run('CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run(
                                'fica ciente que, o pagamento das ou qualquer outro encargo feito através da emissão de cheque e este, por qualquer motivo que seja recusado pela entidade financeira, será considerado nulo, ficando sem efeito de quitação anterior.\n\n4.4 Fica estipulado aos ')
                            runner = paragraph.add_run('CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run(
                                ' o pagamento de taxas, emissões de documentos, impressões e plotagens   referente as tramitações do processo bem como a exigências dos órgãos envolvidos.')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('5. MULTAS')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('5.1 ')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            runner = paragraph.add_run('MULTA DE MORA: ')
                            runner.bold = True
                            paragraph.add_run('Fica estipulada a multa de ')
                            runner = paragraph.add_run('10%')
                            runner.bold = True
                            paragraph.add_run(
                                '% (dez por cento) calculada sobre o valor do contrato devido à qual será devidamente cobrada juntamente com este pelo não pagamento no prazo previsto neste contrato. Os')
                            runner = paragraph.add_run(' CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run(
                                'fica ciente, ainda que a multa se refere simplesmente a mora, não impedindo a cobrança de outras multas devidas pela inflação ou inexecução de quaisquer das mais clausulas deste contrato e a sua cobrança em meses posteriores. Além dessa multa serão devidos juros e correção monetária, ambos calculados pelo máximo e na forma permitidos pela lei vigente, sempre que houver atraso no pagamento do contrato ou dos encargos.')

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('6. FORO COMPETENTE, DESAPROPRIAÇÂO E AÇOES.')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '6.1 O presente contrato reger-se-á pelo Código Civil Brasileiro e Legislação suplementar, não estabelecendo, em hipóteses alguma, vínculo empregatício entre os contratantes. \n\n'
                                '6.2 Fica eleito o foro')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            runner = paragraph.add_run(' HORTOLÂNDIA – SP')
                            runner.bold = True
                            paragraph.add_run(
                                ', com renúncia de outro por mais privilegio que seja, para quaisquer dúvidas decorrentes deste contrato, sejam quais forem os domicílios dos contratos, mesmo para ações posteriores ao termino deste contrato, inclusive para ações de cobrança, ainda que os ')
                            runner = paragraph.add_run('CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run('tenham se mudado ou já residam em outra comarca.\n\n'
                                              '6.3 Nos termos do deciso no processo 85.232/88, da Serasa, com força de Provimento, publicado no diário Oficial de 22/6/1995, Cad. I Parte I, se o poder Judiciário informar a ação á Serasa, as providencias pela baixa do nome nesse órgão ficam por conta do devedor, uma vez que não foi a locadora ou sua representante legal que fez a informação.')

                            paragraph = document.add_paragraph(
                                'E assim por estarem justos e contratados assinam o presente em 02 (duas) vias de igual teor e valor.')
                            paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ' / SP, ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '_________________________________          _____________________________________')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.italic = True
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'CONTRATADO:                                            CONTRATANTE:')
                            paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'ROGÉRIO ROCHA SOARES                      ' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '_________________________________          _____________________________________')
                            paragraph.style = document.styles.add_style('style17.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.italic = True
                            font.bold = True

                            paragraph = document.add_paragraph('CONTRATANTE:                                            CONTRATANTE:')
                            paragraph.style = document.styles.add_style('style18.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            paragraph = document.add_paragraph('' + nomecli2 + '                      ' + nomecli3 + '')
                            paragraph.style = document.styles.add_style('style19.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+ nomecli2 +' e ' + nomecli3 + '/' + tipoobra + '/' + ano + '/Documentos/Contrato ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '.docx')

                        if self.cbox_memorial.isChecked() == True:
                            # Memorial Descritivo
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(1.27)
                                section.left_margin = Cm(1.27)
                                section.right_margin = Cm(1.27)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('MEMORIAL DESCRITIVO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('Obra: Regularização Residencial Unifamiliar (R1)\n'
                                                               'Local: ' + endobra + ', nº ' + numobra + ' Lote: ' + loteobra + ' – Quadra: ' + quadraobra + '\n'
                                                                                                                                                             'Loteamento: ' + bairroobra + ' -  ' + cidadeobra + ' - SP\n'
                                                                                                                                                                                                                 'Proprietário: ')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)
                            runner = paragraph.add_run('' + nomecli1 + ' \n')
                            runner.bold = True
                            paragraph.add_run('                     CPF:')
                            runner = paragraph.add_run('' + cpfcli1 + ' \n')
                            runner.bold = True

                            runner = paragraph.add_run('                     ' + nomecli2 + ' \n')
                            runner.bold = True
                            paragraph.add_run('                     CPF:')
                            runner = paragraph.add_run('' + cpfcli2 + ' \n')
                            runner.bold = True



                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('ÁREA DO TERRENO – ' + areaobra + ' m² ')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)

                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('Descrição')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                'MOVIMENTO DE TERRA: Será realizada uma limpeza da superfície do terreno, tal como remoção da camada vegetal.')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'TIPO DE FUNDAÇÃO: Serão executadas brocas de concreto armado, posteriormente, sapatas na cabeça das brocas, as quais apoiarão as vigas baldrames, que por sua vez embasam a alvenaria. Todos os componentes da fundação serão executados e concretados “in loco”.')
                            paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'IMPERMEABILIZAÇÃO: Para evitar os fenômenos de capilaridade e percolação (umidade nas partes inferiores da alvenaria), todas as vigas baldrames e primeiras fiadas da alvenaria receberão tratamento contra a umidade proveniente do solo.')
                            paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ESTRUTURA: Será utilizadas estrutura convencional de concreto armado: pilares e vigas. As peças serão devidamente moldadas e concretadas “in loco”.')
                            paragraph.style = document.styles.add_style('style5.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ALVENARIA: Será executada alvenaria com blocos cerâmicos assentados com argamassa mista de cimento cal e areia. Serão executadas também vergas e contravergas nos vãos de janelas e portas.')
                            paragraph.style = document.styles.add_style('style5.5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'FORRO: Será executado em laje pré-fabricada de concreto armado em todas as dependências da edificação, com uso de vigas pré-fabricadas e lajotas cerâmicas.')
                            paragraph.style = document.styles.add_style('style5.6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'COBERTURA: A cobertura será composta por telhas Metálicas com estrutura de aço; o telhado apresentará inclinação média de 11%.')
                            paragraph.style = document.styles.add_style('style5.7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'REVESTIMENTO PAREDES: O revestimento das paredes de dependências consideradas molhadas será executado com revestimento cerâmico até o teto. Nas demais paredes, o revestimento será constituído de chapisco grosso e emboço com argamassa mista de cimento e areia.')
                            paragraph.style = document.styles.add_style('style5.8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PISOS E RODAPÉS: O revestimento do piso será de cerâmica em todas as dependências internas; externamente o piso será revestido de cimentado desempenado.')
                            paragraph.style = document.styles.add_style('style5.9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PORTAS: As portas, tanto internas quanto externas, serão de madeira e alumínio, todas com as devidas ferragens e em bom estado de conservação.')
                            paragraph.style = document.styles.add_style('style5.10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'JANELAS: Todas as janelas serão de alumínio, de correr ou de abrir, estas também, devidamente tratadas e envernizadas.')
                            paragraph.style = document.styles.add_style('style5.11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ÁGUAS PLUVIAIS: Serão coletadas e conduzidas à via pública. A captação será feita pelo telhado que por sua vez conduzirá estas águas até as calhas de onde irão para a via pública, por meio de tubos de PVC, as águas lançadas no quintal irão para via pública.')
                            paragraph.style = document.styles.add_style('style5.12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'INSTALAÇÕES ELÉTRICAS: Serão executadas de acordo com as Normas Técnicas. Serão utilizados fios e cabos de cobre, cujas secções atenderão as necessidades a que serão submetidos. Todas a instalações serão instaladas em perfeitas condições de uso. O medidor de entrada de energia será executado e instalado de acordo com as necessidades e exigências da concessionária fornecedora deste serviço.')
                            paragraph.style = document.styles.add_style('style5.13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'INSTALAÇÕES HIDRO-SANITÁRIAS: Estas instalações serão dimensionadas e executadas de acordo com as Normas Técnicas. Sendo que todos os aparelhos hidrossanitários estarão em funcionamento correto. Serão utilizadas e executadas caixas de inspeção e de gordura nos pontos necessários para boa manutenção e funcionamento destas instalações. Para condução de água (fria) potável, serão utilizados tubos de PVC marrom soldável. O sistema de esgoto e águas servidas tem seu lançamento para rede publica de esgoto.')
                            paragraph.style = document.styles.add_style('style5.14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PINTURA: A pintura será executada em látex em todas as paredes, tanto internas quando externas. As esquadrias e caixilharias em geral, como já descrito, serão devidamente impermeabilizadas e protegidas contra as intempéries.')
                            paragraph.style = document.styles.add_style('style5.15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'LIMPEZA: A obra será totalmente limpa de entulhos. A edificação será apta a ser habitada, obedecendo às condições mínimas de conforto, segurança e habitabilidade.')
                            paragraph.style = document.styles.add_style('style5.16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________                                    _____________________________________')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário:' + nomecli1 + '                                                Proprietário:' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli1 + '                                                                        CPF:' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________                                    _____________________________________')
                            paragraph.style = document.styles.add_style('style17.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário:' + nomecli3 + '                                                 ROGÉRIO ROCHA SOARES')
                            paragraph.style = document.styles.add_style('style19.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli3 + '                                                                         Engenheiro Civil')
                            paragraph.style = document.styles.add_style('style20.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '                                                                     CREA: 5070347192-SP')
                            paragraph.style = document.styles.add_style('style21.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '                                                                     ART' + artobra + '')
                            paragraph.style = document.styles.add_style('style22.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '/' + tipoobra + '/' + ano + '/Documentos/Memorial Descritivo ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '.docx')

                            # RRC sem lei
                        if self.cbox_reqslei.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(1)
                                section.bottom_margin = Cm(1)
                                section.left_margin = Cm(3)
                                section.right_margin = Cm(1)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(1.48), height=Cm(1.48))

                            paragraph = document.add_paragraph('Ao'
                                                               '\nExcelentíssimo Senhor Prefeito Municipal,')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '                    Venho respeitosamente à presença de Vossa Excelência requerer, por meio do representante legal que em conjunto este subscreve, que se digne em providenciar por meio do órgão competente o que segue:')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '(   ) PRED - Desdobro de lote 	(   ) PRED - Regularização de edificação')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            # add grid table

                            table = document.add_table(rows=3, cols=1, style='Table Grid')
                            table.left_margin = Cm(30.4)
                            row = table.rows[0]

                            tabela = 'Dados dos requerentes (titulares do lote ou da edificação)'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            tabela = '\nRazão social/nome: ' + nomecli1 + '' \
                                                                          '\nCNPJ/CPF nº: ' + cpfcli1 + '' \
                                                                                                        '\nE-mail*: ' + emailcli1 + '' \
                                                                                                                                    '\nTelefone para contato: ' + celularcli1 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\naaaaaaaaaaaaaaaaaa'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(3)
                            font = tabela_formatada.font
                            font.color.rgb = RGBColor(255, 255, 255)

                            tabela = '\nRazão social/nome: ' + nomecli2 + '' \
                                                                          '\nCNPJ/CPF nº: ' + cpfcli2 + '' \
                                                                                                        '\nE-mail*: ' + emailcli2 + '' \
                                                                                                                                    '\nTelefone para contato: ' + celularcli2 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\naaaaaaaaaaaaaaaaaa'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(3)
                            font = tabela_formatada.font
                            font.color.rgb = RGBColor(255, 255, 255)

                            tabela = '\nRazão social/nome: ' + nomecli3 + '' \
                                                                          '\nCNPJ/CPF nº: ' + cpfcli3 + '' \
                                                                                                        '\nE-mail*: ' + emailcli3 + '' \
                                                                                                                                    '\nTelefone para contato: ' + celularcli3 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\n*as notificações sobre este processo serão enviadas por e-mail. Favor atentar-se a isso no momento do preenchimento.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            row = table.rows[1]

                            tabela = 'Dados do imóvel:'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            tabela = '\nLote/Gleba/Quinhão nº: ' + loteobra + '' \
                                                                              '\nQuadra: ' + quadraobra + '' \
                                                                                                          '\nLoteamento: ' + bairroobra + '' \
                                                                                                                                          '\nInscrição Imobiliária: ' + inscimobobra + '' \
                                                                                                                                                                                       '\nEndereço: ' + endobra + '' \
                                                                                                                                                                                                                  '\nCEP: ' + cepcli1 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            row = table.rows[2]

                            tabela = 'Dados do Responsável Técnico pelo projeto'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            tabela = '\nNome completo: Rogério Rocha Soares' \
                                     '\nRegistro profissional: 5070374192 Órgão: CREA' \
                                     '\nEstá registrado no CPHO¹?  ( X ) sim     (   ) não' \
                                     '\nNº da Inscrição Mobiliária: 1036/18' \
                                     '\nE-mail²: rocha.soares@hotmail.com' \
                                     '\nTelefone para contato: (19)982009858' \
                                     '\n¹CPHO - Cadastro de Profissionais Habilitados junto aos órgãos da Prefeitura Municipal de Hortolândia.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\n²as notificações sobre este processo serão enviadas por e-mail. Favor atentar-se a isso no momento do preenchimento.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            paragraph = document.add_paragraph(
                                '\n(X) Declaro que os documentos, declarações e demais elementos submetidos na instrução deste requerimento são verdadeiros e que tenho ciência de que a falsidade de qualquer informação prestada acarreta automaticamente em crime de falsidade ideológica na forma do art. 299 do Código Penal Brasileiro.')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.line_spacing = Cm(0)
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph(
                                '(X) Declaro ter ciência de que, caso meu pedido não seja instruído nos termos que determina a legislação vigente, deverei regularizá-lo no prazo de 30 (trinta) dias corridos, sob pena de arquivamento e indeferimento deste processo.')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.line_spacing = Cm(0)
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph(
                                '(X) Declaro ter ciência do prazo de 180 (cento e oitenta) dias corridos, contados da entrega da planta aprovada, para o registro dos desdobros e das edificações junto ao Cartório de Registro de Imóveis competente.')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('        	Nestes termos,')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph('        	Peço Deferimento.')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'Hortolândia, ' + dia + ' de ' + mesescrito + ' de ' + ano + '. ')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n____________________________________________    ____________________________________________ '
                                '\nProprietário 1                                                                    Proprietário 2'
                                '\n\n\n____________________________________________    ____________________________________________ '
                                '\nProprietário 1                                                                    Responsável técnico')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '/' + tipoobra + '/' + ano + '/Documentos/Requerimento sem Lei ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '.docx')

                            # RRC com lei
                        if self.cbox_reqclei.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(16)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('DE ACORDO COM A LEI Nª3.491/2018')
                            paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(11)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('\n\nNós,')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            runner = paragraph.add_run(' ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + ' ')
                            runner.bold = True
                            paragraph.add_run(
                                'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

                            paragraph = document.add_paragraph(
                                '                                                                         Nestes Termos,\n'
                                '                                                                         Pede Deferimento.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n_________________________________                            _________________________________')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                '' + nomecli1 + '                                                           ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: ' + cpfcli1 + '                                                              CPF: ' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________ ')
                            paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                '' + nomecli3 + '')
                            paragraph.style = document.styles.add_style('style6.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: ' + cpfcli3 + '')
                            paragraph.style = document.styles.add_style('style7.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Dados Complementares:')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Dos Proprietários')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Nome:' + nomecli1 + '\n'
                                                                                    'Endereço: ' + endcli1 + ' N°' + numcli1 + '\n'
                                                                                                                               'Loteamento:' + bairrocli1 + '\n'
                                                                                                                                                            'CEP:' + cepcli1 + '\n'
                                                                                                                                                                               'Cidade/Estado:' + cidadecli1 + '-' + estadocli1 + '\n'
                                                                                                                                                                                                                                  'Telefone: ' + celularcli1 + '')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Nome:' + nomecli2 + '\n'
                                                                                    'Endereço: ' + endcli2 + ' N°' + numcli2 + '\n'
                                                                                                                               'Loteamento:' + bairrocli2 + '\n'
                                                                                                                                                            'CEP:' + cepcli2 + '\n'
                                                                                                                                                                               'Cidade/Estado:' + cidadecli2 + '-' + estadocli2 + '\n'
                                                                                                                                                                                                                                  'Telefone: ' + celularcli2 + '')
                            paragraph.style = document.styles.add_style('style10.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Nome:' + nomecli3 + '\n'
                                                                                    'Endereço: ' + endcli3 + ' N°' + numcli3 + '\n'
                                                                                                                               'Loteamento:' + bairrocli3 + '\n'
                                                                                                                                                            'CEP:' + cepcli3 + '\n'
                                                                                                                                                                               'Cidade/Estado:' + cidadecli3 + '-' + estadocli3 + '\n'
                                                                                                                                                                                                                                  'Telefone: ' + celularcli3 + '')
                            paragraph.style = document.styles.add_style('style10.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Da Obra')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Endereço: ' + endobra + ' nº ' + numobra +
                                                               '\nLOTE N° ' + loteobra + 'Loteamento:' + bairroobra + '\n'
                                                                                                                    'Quadra:' + quadraobra + '')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('\n\nDo Responsável Técnico')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                                               'CPF: 183.125.858-77\n'
                                                               'Celular: (19) 982009858\n'
                                                               'Inscrição SMPUGE: 1036/18\n'
                                                               'E-mail: rocha.soares@hotmail.com\n')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '/' + tipoobra + '/' + ano + '/Documentos/Requerimento com Lei ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '.docx')

                            # Procuração
                        if self.cbox_procuracao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('PROCURAÇÃO ')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                'Constitui procurador para a prática de atos perante a Prefeitura Municipal de Hortolândia\n')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(11.5)

                            runner = paragraph.add_run('I - OUTORGANTES:')
                            runner.bold = True

                            paragraph.add_run('\nSr.(a) ' + nomecli1 + ' CPF: ' + cpfcli1 + '')

                            paragraph.add_run('\nSr.(a) ' + nomecli2 + ' CPF: ' + cpfcli2 + '\n\n')

                            paragraph.add_run('\nSr.(a) ' + nomecli3 + ' CPF: ' + cpfcli3 + '\n\n')

                            runner = paragraph.add_run('II – OUTORGADO: ')
                            runner.bold = True

                            paragraph.add_run(
                                '\nSr. Rogério Rocha Soares, portador do CPF nº 183.125.858-77- endereço comercial: Rua Geraldo Denadai, n° 74 1º andar sala 03, Jardim da Paineiras na cidade de Hortolândia/SP.\n\n')

                            runner = paragraph.add_run('III – EXTENSÃO DOS PODERES: ')
                            runner.bold = True

                            paragraph.add_run('\nO OUTORGANTE é proprietário e legítimo dono do imóvel, '
                                              'Lote' + loteobra + 'da Quadra ' + quadraobra + ', localizado no endereço: '
                                                                                              '' + endobra + ' nº ' + numobra + ' Loteamento: ' + bairroobra + '.\n\n')

                            runner = paragraph.add_run('Específicos para a prática de Ato Determinado. ')
                            runner.bold = True

                            paragraph.add_run(
                                '\n(X) Retirar projeto aprovado e Alvará de construção, referente ao imóvel acima. '
                                '\nPor este instrumento particular de mandato e na melhor forma de direito, o OUTORGANTE acima qualificado, nomeia e constitui o PROCURADOR acima qualificado, a quem confere plenos poderes de representação perante a ')

                            runner = paragraph.add_run('PREFEITURA MUNICIPAL DE ' + cidadeobra + ',')
                            runner.bold = True

                            paragraph.add_run(
                                'especialmente para em seu nome e como se o próprio fosse praticar os atos especificados acima.')

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n_________________________________                            _________________________________')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'OUTORGANTE:                                                                                                OUTORGANTE:')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(9)
                            font.bold = True

                            paragraph = document.add_paragraph(
                                '' + nomecli1 + '                                                           ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: ' + cpfcli1 + '                                                              CPF: ' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style4.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'OUTORGANTE:')
                            paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(9)
                            font.bold = True

                            paragraph = document.add_paragraph(
                                '' + nomecli3 + '')
                            paragraph.style = document.styles.add_style('style6.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: ' + cpfcli3 + '')
                            paragraph.style = document.styles.add_style('style7.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('OUTORGADO:')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(9)
                            font.bold = True

                            paragraph = document.add_paragraph('Eng.º Civil: Rogério Rocha Soares')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(10)

                            paragraph = document.add_paragraph(
                                'CPF: 183.125.858-77')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(10)

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '/' + tipoobra + '/' + ano + '/Documentos/Procuração ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '.docx')

                            # Declaração
                        if self.cbox_declaracao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(3.49)
                                section.bottom_margin = Cm(1.1)
                                section.left_margin = Cm(2)
                                section.right_margin = Cm(2)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo_branco.jpg", width=Cm(3.5), height=Cm(2.65))

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/Logo_Hortolandia.png", width=Cm(4.9), height=Cm(2.65))

                            paragraph = document.add_paragraph('ANEXO I')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('DECLARAÇÃO')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('Em conformidade com o disposto no artigo 4º '
                                                               'da Lei Municipal nº 2.529, de 04 de abril de 2011,'
                                                               ' que “Dispõe sobre controle ambiental para utilização'
                                                               ' de produtos e subprodutos de madeira de origem nativa'
                                                               ' em obras e serviços de Engenharia Civil no Município'
                                                               ' de Hortolândia”, nós, ' + nomecli1 + ', (' + profissaocli1 + '),' + nomecli2 + ', (' + profissaocli2 + ') e ' + nomecli3 + ', (' + profissaocli3 + '),')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run(' Proprietários da obra ')
                            runner.bold = True

                            paragraph.add_run(
                                'localizada à ' + endobra + ', nº ' + numobra + ' Lote ' + loteobra + ', Quadra ' + quadraobra + ', Loteamento ' + bairroobra + ',' \
                                                                                                                                                                'cidade de Hortolândia-SP, DECLARAMOS estar ciente das disposições ' \
                                                                                                                                                                'constantes da Instrução Normativa nº 112, de 21 de agosto de 2006, ' \
                                                                                                                                                                'do Instituto Brasileiro do Meio Ambiente e dos Recursos Naturais ' \
                                                                                                                                                                'Renováveis - IBAMA, e me comprometendo a adquirir, para qualquer ' \
                                                                                                                                                                'serviço a ser realizado no referido imóvel, produtos e subprodutos ' \
                                                                                                                                                                'de madeira de origem não nativa ou nativa que tenha procedência ' \
                                                                                                                                                                'legal, decorrentes de desmatamento autorizado ou de manejo florestal' \
                                                                                                                                                                ' aprovado por órgão ambiental competente, integrante do Sistema ' \
                                                                                                                                                                'Nacional do Meio Ambiente – SISNAMA, com autorização de transporte ' \
                                                                                                                                                                'reconhecida pelo órgão ambiental competente, exigindo no ato da ' \
                                                                                                                                                                'compra que as empresas que comercializem madeiras, forneçam o DOF' \
                                                                                                                                                                ' (Documento de Origem Florestal), acompanhado de nota fiscal. ')

                            runner = paragraph.add_run('DECLARAMOS ')
                            runner.bold = True

                            paragraph.add_run(
                                'ainda, que quando da solicitação do Habite-se o proprietário deverá apresentar além dos documentos, declarações e comprovantes exigidos pelo poder Público, ')

                            runner = paragraph.add_run('um dos seguintes documentos: ')
                            runner.bold = True

                            paragraph = document.add_paragraph(
                                '       ●      Nota Fiscal constando o número do DOF, em caso de utilização de produtos ou subprodutos de madeira de ')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run('aaaaaaaaaaa')
                            font = runner.font
                            font.color.rgb = RGBColor(255, 255, 255)

                            paragraph.add_run('origem nativa;')

                            paragraph = document.add_paragraph(
                                '       ●      Declaração que fez a reutilização de madeira ou que utilizou madeira de reflorestamento;')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '       ●     Declaração que não utilizou madeira de origem nativa, por ter utilizado novas tecnologias ou produtos alternativos.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_____________________________________')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('PROPRIETÁRIO ' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_____________________________________')
                            paragraph.style = document.styles.add_style('style8.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('PROPRIETÁRIO ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style9.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_____________________________________')
                            paragraph.style = document.styles.add_style('style8.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('PROPRIETÁRIO ' + nomecli3 + '')
                            paragraph.style = document.styles.add_style('style9.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n\nEm conformidade com o disposto no artigo 4º da '
                                'Lei Municipal nº 2.529, de 04 de abril de 2011, '
                                'que “Dispõe sobre controle ambiental para utilização'
                                ' de produtos e subprodutos de madeira de origem nativa'
                                ' em obras e serviços de Engenharia Civil no Município '
                                'de Hortolândia”, eu, Rogério Rocha Soares, (Engenheiro Civil), '
                                'Autor do Projeto da obra localizada à '
                                'Rua ' + endobra + ', nº ' + numobra + ' Lote ' + loteobra + ', Quadra ' + quadraobra + ', Loteamento ' + bairroobra + ',cidade de Hortolândia-SP,'
                                                                                                                                                       ' DECLARO estar ciente das disposições constantes da Instrução Normativa'
                                                                                                                                                       ' nº 112, de 21 de agosto de 2006, do Instituto Brasileiro '
                                                                                                                                                       'do Meio Ambiente e dos Recursos Naturais Renováveis - IBAMA, e ')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run(
                                'me comprometendo a instruir meu cliente a adquirir produtos e subprodutos de madeira de origem não nativa ou nativa que tenha procedência legal, ')
                            runner.underline = True

                            paragraph.add_run(
                                'decorrentes de desmatamento autorizado ou de manejo florestal aprovado por órgão ambiental competente, integrante do Sistema Nacional do Meio Ambiente – SISNAMA, com autorização de transporte reconhecida pelo órgão ambiental competente, exigindo no ato da compra que as empresas que comercializem madeiras, forneçam o DOF (Documento de Origem Florestal), acompanhado de nota fiscal.')

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('AUTOR DO PROJETO')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            # runner_word.size = Pt(10)

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '/' + tipoobra + '/' + ano + '/Documentos/Declaração ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '.docx')

                            # ---------------------MEMORIAL DESCRITIVO PARA CONSTRUÇÃO---------------------------------------------------------------------------------------------
                        if self.cbox_memorialcontrucao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('MEMORIAL DESCRITIVO PARA CONSTRUÇÃO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                'Obra: REGULARIZAÇÃO E CONSTRUÇÃO RESIDENCIAL MULTIFAMILIAR – R2')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('(DE ACORDO COM A LEI 3.491/2018 - ANISTIA)')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('Local: ' + endobra + '- N° ' + numobra + '')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Lote: ' + loteobra + ' Quadra: ' + quadraobra + '')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Loteamento: ' + bairroobra + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Município: ' + cidadeobra + '/SP')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph(
                                'Proprietário(s): ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Responsável Técnico: Eng. Civil ROGÉRIO ROCHA SOARES')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph(
                                '01 – Preparação do terreno: Raspagem com moto-niveladora e os resíduos será retirado e depositado em local devidamente autorizado e legalizado pela Prefeitura Municipal.')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '02 – Fundações: A fundação do alicerçada em estacas manuais, blocos e vigas baldrame em concreto armado.')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '03 – Impermeabilização: Com Vedacit no concreto e na argamassa de embasamento. Será aplicado neutrol no baldrame antes do reaterro.')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '04 – Estrutura: Os pilares, vergas, contra-vergas, vigas de apoio e vigas de respaldo em concreto armado batido na própria obra.')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '05 - Alvenaria: As paredes externas e internas executadas com bloco cerâmico (barro) de oito furos (15 cm), assentados com argamassa de areia e cimento. A altura do pé direito será de 2,80m em toda a casa.')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('06 - Forro: O forro de laje pré-moldada de concreto.')
                            paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '07- Cobertura: O telhado foi executado com telhas de barro estrutura de madeira com inclinação de 26% na residência.')
                            paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '08 - Instalações hidráulicas: Instalado conforme normas da SABESP e NBR 7229/93,')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        a) - Água fria: Abastecida pela rede pública e armazenada em dois reservatórios elevados, com capacidade de 1.000 litros d’agua na residência;')
                            paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        b) - Esgotos sanitários: canalizados com tubos de PVC com diâmetro 100 mm para a rede pública de afastamento de esgoto.')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        c) – Águas pluviais: coletada por meio de sistema de calhas galvanizadas e desce por condutores e são canalizadas até a sarjeta por meio de tubos de PVC, por sob a calçada.')
                            paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '09 - Instalações elétricas: Instalado conforme normas da CPFL, composto de um ponto de luz para cada compartimento e tomadas em número suficiente para atender a demanda.')
                            paragraph.style = document.styles.add_style('style21', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '10 – Revestimento: Os banheiros da residência e da dependência em azulejo até a altura do forro;')
                            paragraph.style = document.styles.add_style('style22', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '11 - Pisos: Nivelado com argamassa de areia e cimento e revestimento cerâmico em todos os compartimentos.')
                            paragraph.style = document.styles.add_style('style23', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '12 - Esquadrias: As Janelas e portas externas são em alumínio e as portas internas de madeira, conforme as medidas indicadas no projeto.')
                            paragraph.style = document.styles.add_style('style24', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '13 – Pintura: As paredes de alvenaria, internas pintadas com tinta látex PVA; as paredes externas pintadas com tinta látex, acrílica, as janelas e portas serão pintadas com esmalte sintético.')
                            paragraph.style = document.styles.add_style('style25', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '14 – Limpeza Geral da Obra: Após a conclusão da obra, foi feita a retirada dos restos de materiais e entulho.')
                            paragraph.style = document.styles.add_style('style26', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style27', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n_________________________________                    _____________________________________')
                            paragraph.style = document.styles.add_style('style28', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário: ' + nomecli1 + '                               Proprietário: ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style29', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli1 + '                                                        CPF:' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style30', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style28.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário: ' + nomecli3 + '')
                            paragraph.style = document.styles.add_style('style29.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli3 + '')
                            paragraph.style = document.styles.add_style('style30.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style28.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'ROGÉRIO ROCHA SOARES')
                            paragraph.style = document.styles.add_style('style29.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Engenheiro Civil')
                            paragraph.style = document.styles.add_style('style30.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CREA: 5070347192-SP')
                            paragraph.style = document.styles.add_style('style31.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'SMPUGE: 1036/18')
                            paragraph.style = document.styles.add_style('style32.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '/' + tipoobra + '/' + ano + '/Documentos/Memorial Descritivo Para Construção ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '.docx')

                            # ---------------------Recibo---------------------------------------------------------------------------------------------
                        if self.cbox_recibo.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('RECIBO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            font.size = Pt(16)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('Eu ROGÉRIO ROCHA SOARES engenheiro civil CREA: '
                                                               '5070374192, recebi de ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + ', '
                                                                                                                        'parte do pagamento para aprovação de projeto '
                                                                                                                        'arquitetônico a quantia de R$ ' + valorparcobra + ',00 (' + num2words(
                                valorparcobra.replace(",", "."), lang='pt-br') +
                                                               '), de um total de '
                                                               'R$ ' + valorobra + ',00 (' + num2words(
                                (valorobra).replace(",", "."), lang='pt-br') + ').')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            font.size = Pt(12)

                            paragraph = document.add_paragraph(
                                '\n\n' + cidadeobra + ', ' + dia + ' / ' + mesescrito + ' / ' + ano + '.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            font.size = Pt(12)

                            # footer section
                            footer_section = document.sections[0]
                            footer = footer_section.footer

                            # footer text
                            footer_text = footer.paragraphs[0]
                            footer_text.text = "_______________________________________________________________________________________________" \
                                               "\nROGÉRIO ROCHA SOARES   TEL. : (19) 982009858" \
                                               "\nCREA: 5070347192" \
                                               "\nE-MAIL: rocha.soares@hotmail.com"
                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '/' + tipoobra + '/' + ano + '/Documentos/Recibo ' + nomecli1 + ', '+nomecli2+' e ' + nomecli3 + '.docx')

                        up.show()
                        pop.lbl_popup.setText("DOCUMENTOS CRIADOS")
                        pop.frame_popup.setStyleSheet("background-color: rgb(57, 173, 84);\n"
                                                      "border-radius:5px;")
                        pop.lbl_popup.setStyleSheet("color: rgb(35, 35, 35)")

                        self.txt_idobra.setText(None)
                        self.txt_idcli1.setText(None)
                        self.txt_idcli2.setText(None)
                        self.txt_idcli3.setText(None)

            elif self.cbox_4cli.isChecked() == True:
                if self.txt_idcli1.text() == "" or self.txt_idcli2.text() == "" or self.txt_idcli3.text() == "" or self.txt_idcli4.text() == "":
                    up.show()
                    pop.lbl_popup.setText("DIGITE O ID DO CLIENTE!")
                    pop.frame_popup.setStyleSheet("background-color: rgb(255, 11, 15);\n"
                                                  "border-radius:5px;")
                    pop.lbl_popup.setStyleSheet("color: rgb(200, 200, 255)")
                else:
                    idcli1 = self.txt_idcli1.text()

                    consultacli = 'SELECT * FROM cliente WHERE id=?'
                    cursor.execute(consultacli, (idcli1,))

                    dados_cli1 = cursor.fetchall()

                    nomecli1 = dados_cli1[0][1]
                    cpfcli1 = dados_cli1[0][2]
                    rgcli1 = dados_cli1[0][3]
                    endcli1 = dados_cli1[0][4]
                    bairrocli1 = dados_cli1[0][5]
                    numcli1 = dados_cli1[0][6]
                    cidadecli1 = dados_cli1[0][7]
                    estadocli1 = dados_cli1[0][8]
                    cepcli1 = dados_cli1[0][9]
                    nacionalidadecli1 = dados_cli1[0][10]
                    profissaocli1 = dados_cli1[0][11]
                    estadocivilcli1 = dados_cli1[0][12]
                    celularcli1 = dados_cli1[0][13]
                    emailcli1 = dados_cli1[0][14]

                    idcli2 = self.txt_idcli2.text()
                    cursor.execute(consultacli, (idcli2,))

                    dados_cli2 = cursor.fetchall()

                    nomecli2 = dados_cli2[0][1]
                    cpfcli2 = dados_cli2[0][2]
                    rgcli2 = dados_cli2[0][3]
                    endcli2 = dados_cli2[0][4]
                    bairrocli2 = dados_cli2[0][5]
                    numcli2 = dados_cli2[0][6]
                    cidadecli2 = dados_cli2[0][7]
                    estadocli2 = dados_cli2[0][8]
                    cepcli2 = dados_cli2[0][9]
                    nacionalidadecli2 = dados_cli2[0][10]
                    profissaocli2 = dados_cli2[0][11]
                    estadocivilcli2 = dados_cli2[0][12]
                    celularcli2 = dados_cli2[0][13]
                    emailcli2 = dados_cli2[0][14]

                    idcli3 = self.txt_idcli3.text()
                    cursor.execute(consultacli, (idcli3,))

                    dados_cli3 = cursor.fetchall()

                    nomecli3 = dados_cli3[0][1]
                    cpfcli3 = dados_cli3[0][2]
                    rgcli3 = dados_cli3[0][3]
                    endcli3 = dados_cli3[0][4]
                    bairrocli3 = dados_cli3[0][5]
                    numcli3 = dados_cli3[0][6]
                    cidadecli3 = dados_cli3[0][7]
                    estadocli3 = dados_cli3[0][8]
                    cepcli3 = dados_cli3[0][9]
                    nacionalidadecli3 = dados_cli3[0][10]
                    profissaocli3 = dados_cli3[0][11]
                    estadocivilcli3 = dados_cli3[0][12]
                    celularcli3 = dados_cli3[0][13]
                    emailcli3 = dados_cli3[0][14]

                    idcli4 = self.txt_idcli4.text()
                    cursor.execute(consultacli, (idcli4,))
                    dados_cli4 = cursor.fetchall()

                    nomecli4 = dados_cli4[0][1]
                    cpfcli4 = dados_cli4[0][2]
                    rgcli4 = dados_cli4[0][3]
                    endcli4 = dados_cli4[0][4]
                    bairrocli4 = dados_cli4[0][5]
                    numcli4 = dados_cli4[0][6]
                    cidadecli4 = dados_cli4[0][7]
                    estadocli4 = dados_cli4[0][8]
                    cepcli4 = dados_cli4[0][9]
                    nacionalidadecli4 = dados_cli4[0][10]
                    profissaocli4 = dados_cli4[0][11]
                    estadocivilcli4 = dados_cli4[0][12]
                    celularcli4 = dados_cli4[0][13]
                    emailcli4 = dados_cli4[0][14]

                    Path(
                        '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', ' + nomecli2 + ', '+ nomecli3+ ' e '+nomecli4+'/' + tipoobra + '/' + ano + '/Documentos''').mkdir(
                        parents=True, exist_ok=True)

                    if self.cbox_contrato.isChecked() == False and self.cbox_recibo.isChecked() == False \
                            and self.cbox_procuracao.isChecked() == False \
                            and self.cbox_reqslei.isChecked() == False and self.cbox_reqclei.isChecked() == False \
                            and self.cbox_memorial.isChecked() == False and self.cbox_memorialcontrucao.isChecked() == False:
                        up.show()
                        pop.lbl_popup.setText("ESCOLHA UM TIPO DE DOCUMENTO!")
                        pop.frame_popup.setStyleSheet("background-color: rgb(255, 11, 15);\n"
                                                      "border-radius:5px;")
                        pop.lbl_popup.setStyleSheet("color: rgb(200, 200, 255)")
                    else:
                        # ---------------------Contrato---------------------------------------------------------------------------------------------
                        if self.cbox_contrato.isChecked() == True:
                            document = Document()

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('CONTRATO DE PRESTAÇÃO DE SERVIÇOS TÉCNICOS')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.underline = True
                            font.color.rgb = RGBColor(0, 0, 255)
                            font.name = 'Book Antiqua'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                'Entre as partes abaixo identificadas e no final assinadas fica contratada os serviços de um imóvel mediante as seguintes clausulas e condições:')
                            paragraph.style = document.styles.add_style('style01', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.PARTES')
                            paragraph.style = document.styles.add_style('style0', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.1 CONTRATADO:')
                            paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run(' ROGÉRIO ROCHA SOARES, ').underline = True
                            paragraph.add_run(
                                'brasileiro, casado, Engenheiro civil, portador do RG n° 27.274.081-0 SSP/SP, inscrito no CPF n° 183.125.858-77 e CREA: 5070347192, residente e domiciliado na Rua Ricardo Mendes Horacy, n° 125, Jardim Nossa Senhora Auxiliadora na cidade de Hortolândia/SP. ').bold = False
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('1.1 CONTRATANTE: ')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run('' + nomecli1 + ', ').underline = True
                            paragraph.add_run(
                                '' + nacionalidadecli1 + ', ' + estadocivilcli1 + ', ' + profissaocli1 + ', '
                                                                                                         'portador(a) do RG n° ' + rgcli1 + ' SSP/SP, inscrito(a) no CPF n° ' + cpfcli1 + ''
                                                                                                                                                                                          ', residente e domiciliado(a) na ' + endcli1 + ', '
                                                                                                                                                                                                                                         'n° ' + numcli1 + ', ' + bairrocli1 + ' na cidade de ' + cidadecli1 + '/' + estadocli1 + '. ').bold = False

                            paragraph = document.add_paragraph('    CONTRATANTE: ')
                            paragraph.style = document.styles.add_style('style2.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run('' + nomecli2 + ', ').underline = True
                            paragraph.add_run(
                                '' + nacionalidadecli2 + ', ' + estadocivilcli2 + ', ' + profissaocli2 + ', '
                                                                                                         'portador(a) do RG n° ' + rgcli2 + ' SSP/SP, inscrito(a) no CPF n° ' + cpfcli2 + ''
                                                                                                                                                                                          ', residente e domiciliado(a) na ' + endcli2 + ', '
                                                                                                                                                                                                                                         'n° ' + numcli2 + ', ' + bairrocli2 + ' na cidade de ' + cidadecli2 + '/' + estadocli2 + '. ').bold = False

                            paragraph = document.add_paragraph('    CONTRATANTE: ')
                            paragraph.style = document.styles.add_style('style2.22', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run('' + nomecli3 + ', ').underline = True
                            paragraph.add_run(
                                '' + nacionalidadecli3 + ', ' + estadocivilcli3 + ', ' + profissaocli3 + ', '
                                                                                                         'portador(a) do RG n° ' + rgcli3 + ' SSP/SP, inscrito(a) no CPF n° ' + cpfcli3 + ''
                                                                                                                                                                                          ', residente e domiciliado(a) na ' + endcli3 + ', '
                                                                                                                                                                                                                                         'n° ' + numcli3 + ', ' + bairrocli3 + ' na cidade de ' + cidadecli3 + '/' + estadocli3 + '. ').bold = False

                            paragraph = document.add_paragraph('    CONTRATANTE: ')
                            paragraph.style = document.styles.add_style('style2.23', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.add_run('' + nomecli4 + ', ').underline = True
                            paragraph.add_run(
                                '' + nacionalidadecli4 + ', ' + estadocivilcli4 + ', ' + profissaocli4 + ', '
                                                                                                         'portador(a) do RG n° ' + rgcli4 + ' SSP/SP, inscrito(a) no CPF n° ' + cpfcli4 + ''
                                                                                                                                                                                          ', residente e domiciliado(a) na ' + endcli4 + ', '
                                                                                                                                                                                                                                         'n° ' + numcli4 + ', ' + bairrocli4 + ' na cidade de ' + cidadecli4 + '/' + estadocli4 + '. ').bold = False

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('2.0 OBJETO E FINALIDADES')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('A confecção de projeto arquitetônico e simplificado '
                                                               'residencial em “AUTOCAD”, conforme características do imóvel dos CONTRATANTES e '
                                                               'legislações pertinentes nos termos de Leis uso e ocupação de solo do município de '
                                                               + cidadeobra + ' e normas dos órgãos competentes e com aproveitamento para a aprovação de '
                                                                              'projeto até a liberação do Alvará. Para o imóvel: lote ' + loteobra + ', quadra ' + quadraobra + '; do loteamento '
                                                                                                                                                                                'denominado “' + bairroobra + '”, no município de ' + cidadeobra + '-SP.')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            font.italic = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '2.2 Atender e fornecer as informações necessárias para o bom andamento dos serviços;')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            paragraph = document.add_paragraph(
                                'a)	Documentos necessários;\nb)	Livre acesso ao local.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('3. VISITAS')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '3.1 Este contrato não prevê acompanhamento da execução da obra. \n\n'
                                '3.2 Este contrato prevê 2 (duas) visitas técnicas no local da obra.\n\n'
                                '3.3 As visitas devem ser agendadas com antecedência mínima de 2 (dois) dias e seguir a agenda de atividades do responsável técnico.\n\n'
                                '3.4 Caso houver interesse dos ')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            paragraph.add_run('CONTRATANTES').bold = True
                            paragraph.add_run('de mais visitas técnicas na obra, será cobrado a parte.\n'
                                              'Valor para cada visita técnica é de R$ ' + valorvisitaobra + ',00 (' + num2words(
                                valorvisitaobra.replace(",", "."), lang='pt-br') + ') hora.')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('4. PRAZO, PREÇO, FORMA DE PAGAMENTO')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '4.1 O prazo da presente prestação de serviço dá-se início após a assinatura deste contrato e se estendera conforme tramitação legal dos órgãos responsáveis para a liberação dos documentos essenciais com resultado final.\n\n'
                                '4.2 O valor deste contrato é de ')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            runner = paragraph.add_run(
                                "" + valorobra + ' (' + num2words((valorobra).replace(",", "."), lang='pt-br') + ')')
                            runner.bold = True
                            runner.underline = True

                            paragraph.add_run(', que o ')
                            runner = paragraph.add_run('CONTRATANTE ')
                            runner.bold = True
                            paragraph.add_run('se obriga a pagar ao  ')
                            runner = paragraph.add_run('CONTRATADO ')
                            runner.bold = True
                            paragraph.add_run(
                                'em ' + valorparcobra + ',00 (' + num2words(valorparcobra.replace(",", "."),
                                                                            lang='pt-br') + ') vezes mensais, com vencimento '
                                                                                            'todo o dia ' + (
                                datacontratoobra[:2]) + ' de cada mês, com início em '
                                + datacontratoobra + ', constituindo-se nenhuma '
                                                     'tolerância de qualquer recebimento depois do '
                                                     'prazo estipulado.\n\n 4.3 Ao ')
                            runner = paragraph.add_run('CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run(
                                'fica ciente que, o pagamento das ou qualquer outro encargo feito através da emissão de cheque e este, por qualquer motivo que seja recusado pela entidade financeira, será considerado nulo, ficando sem efeito de quitação anterior.\n\n4.4 Fica estipulado aos ')
                            runner = paragraph.add_run('CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run(
                                ' o pagamento de taxas, emissões de documentos, impressões e plotagens   referente as tramitações do processo bem como a exigências dos órgãos envolvidos.')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('5. MULTAS')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('5.1 ')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            runner = paragraph.add_run('MULTA DE MORA: ')
                            runner.bold = True
                            paragraph.add_run('Fica estipulada a multa de ')
                            runner = paragraph.add_run('10%')
                            runner.bold = True
                            paragraph.add_run(
                                '% (dez por cento) calculada sobre o valor do contrato devido à qual será devidamente cobrada juntamente com este pelo não pagamento no prazo previsto neste contrato. Os')
                            runner = paragraph.add_run(' CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run(
                                'fica ciente, ainda que a multa se refere simplesmente a mora, não impedindo a cobrança de outras multas devidas pela inflação ou inexecução de quaisquer das mais clausulas deste contrato e a sua cobrança em meses posteriores. Além dessa multa serão devidos juros e correção monetária, ambos calculados pelo máximo e na forma permitidos pela lei vigente, sempre que houver atraso no pagamento do contrato ou dos encargos.')

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('6. FORO COMPETENTE, DESAPROPRIAÇÂO E AÇOES.')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '6.1 O presente contrato reger-se-á pelo Código Civil Brasileiro e Legislação suplementar, não estabelecendo, em hipóteses alguma, vínculo empregatício entre os contratantes. \n\n'
                                '6.2 Fica eleito o foro')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            runner = paragraph.add_run(' HORTOLÂNDIA – SP')
                            runner.bold = True
                            paragraph.add_run(
                                ', com renúncia de outro por mais privilegio que seja, para quaisquer dúvidas decorrentes deste contrato, sejam quais forem os domicílios dos contratos, mesmo para ações posteriores ao termino deste contrato, inclusive para ações de cobrança, ainda que os ')
                            runner = paragraph.add_run('CONTRATANTES ')
                            runner.bold = True
                            paragraph.add_run('tenham se mudado ou já residam em outra comarca.\n\n'
                                              '6.3 Nos termos do deciso no processo 85.232/88, da Serasa, com força de Provimento, publicado no diário Oficial de 22/6/1995, Cad. I Parte I, se o poder Judiciário informar a ação á Serasa, as providencias pela baixa do nome nesse órgão ficam por conta do devedor, uma vez que não foi a locadora ou sua representante legal que fez a informação.')

                            paragraph = document.add_paragraph(
                                'E assim por estarem justos e contratados assinam o presente em 02 (duas) vias de igual teor e valor.')
                            paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ' / SP, ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '_________________________________          _____________________________________')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.italic = True
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'CONTRATADO:                                            CONTRATANTE:')
                            paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'ROGÉRIO ROCHA SOARES                      ' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '_________________________________          _____________________________________')
                            paragraph.style = document.styles.add_style('style17.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.italic = True
                            font.bold = True

                            paragraph = document.add_paragraph('CONTRATANTE:                                            CONTRATANTE:')
                            paragraph.style = document.styles.add_style('style18.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            paragraph = document.add_paragraph('' + nomecli2 + '                      ' + nomecli3 + '')
                            paragraph.style = document.styles.add_style('style19.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            enter = document.add_paragraph('')
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                '_________________________________')
                            paragraph.style = document.styles.add_style('style17.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.italic = True
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'CONTRATANTE:')
                            paragraph.style = document.styles.add_style('style18.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            paragraph = document.add_paragraph('' + nomecli4 + '')
                            paragraph.style = document.styles.add_style('style19.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            font.bold = True

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+ nomecli2 +', ' + nomecli3 + ' e ' + nomecli4 + '/' + tipoobra + '/' + ano + '/Documentos/Contrato ' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 + '.docx')

                        if self.cbox_memorial.isChecked() == True:
                            # Memorial Descritivo
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(1.27)
                                section.left_margin = Cm(1.27)
                                section.right_margin = Cm(1.27)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('MEMORIAL DESCRITIVO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('Obra: Regularização Residencial Unifamiliar (R1)\n'
                                                               'Local: ' + endobra + ', nº ' + numobra + ' Lote: ' + loteobra + ' – Quadra: ' + quadraobra + '\n'
                                                                                                                                                             'Loteamento: ' + bairroobra + ' -  ' + cidadeobra + ' - SP\n'
                                                                                                                                                                                                                 'Proprietário: ')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)
                            runner = paragraph.add_run('' + nomecli1 + ' \n')
                            runner.bold = True
                            paragraph.add_run('                     CPF:')
                            runner = paragraph.add_run('' + cpfcli1 + ' \n')
                            runner.bold = True

                            runner = paragraph.add_run('                     ' + nomecli2 + ' \n')
                            runner.bold = True
                            paragraph.add_run('                     CPF:')
                            runner = paragraph.add_run('' + cpfcli2 + ' \n')
                            runner.bold = True

                            runner = paragraph.add_run('                     ' + nomecli3 + ' \n')
                            runner.bold = True
                            paragraph.add_run('                     CPF:')
                            runner = paragraph.add_run('' + cpfcli3 + ' \n')
                            runner.bold = True

                            runner = paragraph.add_run('                     ' + nomecli4 + ' \n')
                            runner.bold = True
                            paragraph.add_run('                     CPF:')
                            runner = paragraph.add_run('' + cpfcli4 + ' \n')
                            runner.bold = True



                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('ÁREA DO TERRENO – ' + areaobra + ' m² ')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)

                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph('Descrição')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            enter = document.add_paragraph('')

                            paragraph = document.add_paragraph(
                                'MOVIMENTO DE TERRA: Será realizada uma limpeza da superfície do terreno, tal como remoção da camada vegetal.')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'TIPO DE FUNDAÇÃO: Serão executadas brocas de concreto armado, posteriormente, sapatas na cabeça das brocas, as quais apoiarão as vigas baldrames, que por sua vez embasam a alvenaria. Todos os componentes da fundação serão executados e concretados “in loco”.')
                            paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'IMPERMEABILIZAÇÃO: Para evitar os fenômenos de capilaridade e percolação (umidade nas partes inferiores da alvenaria), todas as vigas baldrames e primeiras fiadas da alvenaria receberão tratamento contra a umidade proveniente do solo.')
                            paragraph.style = document.styles.add_style('style5.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ESTRUTURA: Será utilizadas estrutura convencional de concreto armado: pilares e vigas. As peças serão devidamente moldadas e concretadas “in loco”.')
                            paragraph.style = document.styles.add_style('style5.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ALVENARIA: Será executada alvenaria com blocos cerâmicos assentados com argamassa mista de cimento cal e areia. Serão executadas também vergas e contravergas nos vãos de janelas e portas.')
                            paragraph.style = document.styles.add_style('style5.5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'FORRO: Será executado em laje pré-fabricada de concreto armado em todas as dependências da edificação, com uso de vigas pré-fabricadas e lajotas cerâmicas.')
                            paragraph.style = document.styles.add_style('style5.6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'COBERTURA: A cobertura será composta por telhas Metálicas com estrutura de aço; o telhado apresentará inclinação média de 11%.')
                            paragraph.style = document.styles.add_style('style5.7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'REVESTIMENTO PAREDES: O revestimento das paredes de dependências consideradas molhadas será executado com revestimento cerâmico até o teto. Nas demais paredes, o revestimento será constituído de chapisco grosso e emboço com argamassa mista de cimento e areia.')
                            paragraph.style = document.styles.add_style('style5.8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PISOS E RODAPÉS: O revestimento do piso será de cerâmica em todas as dependências internas; externamente o piso será revestido de cimentado desempenado.')
                            paragraph.style = document.styles.add_style('style5.9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PORTAS: As portas, tanto internas quanto externas, serão de madeira e alumínio, todas com as devidas ferragens e em bom estado de conservação.')
                            paragraph.style = document.styles.add_style('style5.10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'JANELAS: Todas as janelas serão de alumínio, de correr ou de abrir, estas também, devidamente tratadas e envernizadas.')
                            paragraph.style = document.styles.add_style('style5.11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'ÁGUAS PLUVIAIS: Serão coletadas e conduzidas à via pública. A captação será feita pelo telhado que por sua vez conduzirá estas águas até as calhas de onde irão para a via pública, por meio de tubos de PVC, as águas lançadas no quintal irão para via pública.')
                            paragraph.style = document.styles.add_style('style5.12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'INSTALAÇÕES ELÉTRICAS: Serão executadas de acordo com as Normas Técnicas. Serão utilizados fios e cabos de cobre, cujas secções atenderão as necessidades a que serão submetidos. Todas a instalações serão instaladas em perfeitas condições de uso. O medidor de entrada de energia será executado e instalado de acordo com as necessidades e exigências da concessionária fornecedora deste serviço.')
                            paragraph.style = document.styles.add_style('style5.13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'INSTALAÇÕES HIDRO-SANITÁRIAS: Estas instalações serão dimensionadas e executadas de acordo com as Normas Técnicas. Sendo que todos os aparelhos hidrossanitários estarão em funcionamento correto. Serão utilizadas e executadas caixas de inspeção e de gordura nos pontos necessários para boa manutenção e funcionamento destas instalações. Para condução de água (fria) potável, serão utilizados tubos de PVC marrom soldável. O sistema de esgoto e águas servidas tem seu lançamento para rede publica de esgoto.')
                            paragraph.style = document.styles.add_style('style5.14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'PINTURA: A pintura será executada em látex em todas as paredes, tanto internas quando externas. As esquadrias e caixilharias em geral, como já descrito, serão devidamente impermeabilizadas e protegidas contra as intempéries.')
                            paragraph.style = document.styles.add_style('style5.15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'LIMPEZA: A obra será totalmente limpa de entulhos. A edificação será apta a ser habitada, obedecendo às condições mínimas de conforto, segurança e habitabilidade.')
                            paragraph.style = document.styles.add_style('style5.16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________                                    _____________________________________')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário:' + nomecli1 + '                                                Proprietário:' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli1 + '                                                                        CPF:' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________                                    _____________________________________')
                            paragraph.style = document.styles.add_style('style17.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário:' + nomecli3 + '                                                 ' + nomecli4 + '')
                            paragraph.style = document.styles.add_style('style19.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli3 + '                                                                         ' + cpfcli4 + '')
                            paragraph.style = document.styles.add_style('style20.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_____________________________________''')
                            paragraph.style = document.styles.add_style('style17.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário: ROGÉRIO ROCHA SOARES')
                            paragraph.style = document.styles.add_style('style19.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: Engenheiro Civil')
                            paragraph.style = document.styles.add_style('style20.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CREA: 5070347192-SP')
                            paragraph.style = document.styles.add_style('style21.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'ART' + artobra + '')
                            paragraph.style = document.styles.add_style('style22.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 +'/' + tipoobra + '/' + ano + '/Documentos/Memorial Descritivo ' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 +'.docx')

                            # RRC sem lei
                        if self.cbox_reqslei.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(1)
                                section.bottom_margin = Cm(1)
                                section.left_margin = Cm(3)
                                section.right_margin = Cm(1)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(1.48), height=Cm(1.48))

                            paragraph = document.add_paragraph('Ao'
                                                               '\nExcelentíssimo Senhor Prefeito Municipal,')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '                    Venho respeitosamente à presença de Vossa Excelência requerer, por meio do representante legal que em conjunto este subscreve, que se digne em providenciar por meio do órgão competente o que segue:')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '(   ) PRED - Desdobro de lote 	(   ) PRED - Regularização de edificação')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            # add grid table

                            table = document.add_table(rows=3, cols=1, style='Table Grid')
                            table.left_margin = Cm(30.4)
                            row = table.rows[0]

                            tabela = 'Dados dos requerentes (titulares do lote ou da edificação)'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            tabela = '\nRazão social/nome: ' + nomecli1 + '' \
                                                                          '\nCNPJ/CPF nº: ' + cpfcli1 + '' \
                                                                                                        '\nE-mail*: ' + emailcli1 + '' \
                                                                                                                                    '\nTelefone para contato: ' + celularcli1 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\naaaaaaaaaaaaaaaaaa'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(3)
                            font = tabela_formatada.font
                            font.color.rgb = RGBColor(255, 255, 255)

                            tabela = '\nRazão social/nome: ' + nomecli2 + '' \
                                                                          '\nCNPJ/CPF nº: ' + cpfcli2 + '' \
                                                                                                        '\nE-mail*: ' + emailcli2 + '' \
                                                                                                                                    '\nTelefone para contato: ' + celularcli2 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\naaaaaaaaaaaaaaaaaa'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(3)
                            font = tabela_formatada.font
                            font.color.rgb = RGBColor(255, 255, 255)

                            tabela = '\nRazão social/nome: ' + nomecli3 + '' \
                                                                          '\nCNPJ/CPF nº: ' + cpfcli3 + '' \
                                                                                                        '\nE-mail*: ' + emailcli3 + '' \
                                                                                                                                    '\nTelefone para contato: ' + celularcli3 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\n*as notificações sobre este processo serão enviadas por e-mail. Favor atentar-se a isso no momento do preenchimento.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            tabela = '\naaaaaaaaaaaaaaaaaa'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(3)
                            font = tabela_formatada.font
                            font.color.rgb = RGBColor(255, 255, 255)

                            tabela = '\nRazão social/nome: ' + nomecli4 + '' \
                                                                          '\nCNPJ/CPF nº: ' + cpfcli4 + '' \
                                                                                                        '\nE-mail*: ' + emailcli4 + '' \
                                                                                                                                    '\nTelefone para contato: ' + celularcli4 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\n*as notificações sobre este processo serão enviadas por e-mail. Favor atentar-se a isso no momento do preenchimento.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            row = table.rows[1]

                            tabela = 'Dados do imóvel:'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            tabela = '\nLote/Gleba/Quinhão nº: ' + loteobra + '' \
                                                                              '\nQuadra: ' + quadraobra + '' \
                                                                                                          '\nLoteamento: ' + bairroobra + '' \
                                                                                                                                          '\nInscrição Imobiliária: ' + inscimobobra + '' \
                                                                                                                                                                                       '\nEndereço: ' + endobra + '' \
                                                                                                                                                                                                                  '\nCEP: ' + cepcli1 + ''
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            row = table.rows[2]

                            tabela = 'Dados do Responsável Técnico pelo projeto'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            tabela = '\nNome completo: Rogério Rocha Soares' \
                                     '\nRegistro profissional: 5070374192 Órgão: CREA' \
                                     '\nEstá registrado no CPHO¹?  ( X ) sim     (   ) não' \
                                     '\nNº da Inscrição Mobiliária: 1036/18' \
                                     '\nE-mail²: rocha.soares@hotmail.com' \
                                     '\nTelefone para contato: (19)982009858' \
                                     '\n¹CPHO - Cadastro de Profissionais Habilitados junto aos órgãos da Prefeitura Municipal de Hortolândia.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)

                            tabela = '\n²as notificações sobre este processo serão enviadas por e-mail. Favor atentar-se a isso no momento do preenchimento.'
                            tabela_formatada = row.cells[0].paragraphs[0].add_run(tabela)
                            tabela_formatada.font.name = 'Arial'
                            tabela_formatada.font.size = Pt(9)
                            tabela_formatada.bold = True

                            paragraph = document.add_paragraph(
                                '\n(X) Declaro que os documentos, declarações e demais elementos submetidos na instrução deste requerimento são verdadeiros e que tenho ciência de que a falsidade de qualquer informação prestada acarreta automaticamente em crime de falsidade ideológica na forma do art. 299 do Código Penal Brasileiro.')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.line_spacing = Cm(0)
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph(
                                '(X) Declaro ter ciência de que, caso meu pedido não seja instruído nos termos que determina a legislação vigente, deverei regularizá-lo no prazo de 30 (trinta) dias corridos, sob pena de arquivamento e indeferimento deste processo.')
                            font = paragraph.style.font
                            font.size = Pt(10)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.line_spacing = Cm(0)
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph(
                                '(X) Declaro ter ciência do prazo de 180 (cento e oitenta) dias corridos, contados da entrega da planta aprovada, para o registro dos desdobros e das edificações junto ao Cartório de Registro de Imóveis competente.')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('        	Nestes termos,')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.paragraph_format.space_after = Cm(0)

                            paragraph = document.add_paragraph('        	Peço Deferimento.')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                'Hortolândia, ' + dia + ' de ' + mesescrito + ' de ' + ano + '. ')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n____________________________________________    ____________________________________________ '
                                '\nProprietário 1                                                                    Proprietário 2'
                                '\n\n\n____________________________________________    ____________________________________________ '
                                '\nProprietário 1                                                                    Responsável técnico')
                            font = paragraph.style.font
                            font.size = Pt(9)
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 +'/' + tipoobra + '/' + ano + '/Documentos/Requerimento sem Lei ' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 + '.docx')

                            # RRC com lei
                        if self.cbox_reqclei.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('EXMO SR. PREFEITO DO MUNICÍPIO DE HORTOLÂNDIA,')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(16)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('DE ACORDO COM A LEI Nª3.491/2018')
                            paragraph.style = document.styles.add_style('style1', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(11)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('\n\nNós,')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            runner = paragraph.add_run(' ' + nomecli1 + ', '+nomecli2+', '+nomecli3+' e ' + nomecli4 + ' ')
                            runner.bold = True
                            paragraph.add_run(
                                'abaixo assinado vem mui respeitosamente, solicitar a aprovação do projeto para construção residencial familiar, no imóvel abaixo descrito, cuja documentação segue anexa.')

                            paragraph = document.add_paragraph(
                                '                                                                         Nestes Termos,\n'
                                '                                                                         Pede Deferimento.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.\n')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n_________________________________                            _________________________________')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                '' + nomecli1 + '                                                           ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: ' + cpfcli1 + '                                                              CPF: ' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n_________________________________                            _________________________________')
                            paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                '' + nomecli3 + '                                                           ' + nomecli4 + '')
                            paragraph.style = document.styles.add_style('style6.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: ' + cpfcli3 + '                                                              CPF: ' + cpfcli4 + '')
                            paragraph.style = document.styles.add_style('style7.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Dados Complementares:')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Dos Proprietários')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Nome:' + nomecli1 + '\n'
                                                                                    'Endereço: ' + endcli1 + ' N°' + numcli1 + '\n'
                                                                                                                               'Loteamento:' + bairrocli1 + '\n'
                                                                                                                                                            'CEP:' + cepcli1 + '\n'
                                                                                                                                                                               'Cidade/Estado:' + cidadecli1 + '-' + estadocli1 + '\n'
                                                                                                                                                                                                                                  'Telefone: ' + celularcli1 + '')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Nome:' + nomecli2 + '\n'
                                                                                    'Endereço: ' + endcli2 + ' N°' + numcli2 + '\n'
                                                                                                                               'Loteamento:' + bairrocli2 + '\n'
                                                                                                                                                            'CEP:' + cepcli2 + '\n'
                                                                                                                                                                               'Cidade/Estado:' + cidadecli2 + '-' + estadocli2 + '\n'
                                                                                                                                                                                                                                  'Telefone: ' + celularcli2 + '')
                            paragraph.style = document.styles.add_style('style10.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Nome:' + nomecli3 + '\n'
                                                                                    'Endereço: ' + endcli3 + ' N°' + numcli3 + '\n'
                                                                                                                               'Loteamento:' + bairrocli3 + '\n'
                                                                                                                                                            'CEP:' + cepcli3 + '\n'
                                                                                                                                                                               'Cidade/Estado:' + cidadecli3 + '-' + estadocli3 + '\n'
                                                                                                                                                                                                                                  'Telefone: ' + celularcli3 + '')
                            paragraph.style = document.styles.add_style('style10.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Nome:' + nomecli4 + '\n'
                                                                                    'Endereço: ' + endcli4 + ' N°' + numcli4 + '\n'
                                                                                                                               'Loteamento:' + bairrocli4 + '\n'
                                                                                                                                                            'CEP:' + cepcli4 + '\n'
                                                                                                                                                                               'Cidade/Estado:' + cidadecli4 + '-' + estadocli4 + '\n'
                                                                                                                                                                                                                                  'Telefone: ' + celularcli4 + '')
                            paragraph.style = document.styles.add_style('style10.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('Da Obra')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Endereço: ' + endobra + ' nº ' + numobra +
                                                               '\nLOTE N° ' + loteobra + 'Loteamento:' + bairroobra + '\n'
                                                                                                                    'Quadra:' + quadraobra + '')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph('\n\nDo Responsável Técnico')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('Nome: Rogério Rocha Soares\n'
                                                               'CPF: 183.125.858-77\n'
                                                               'Celular: (19) 982009858\n'
                                                               'Inscrição SMPUGE: 1036/18\n'
                                                               'E-mail: rocha.soares@hotmail.com\n')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+', '+nomecli3+' e ' + nomecli4 +'/' + tipoobra + '/' + ano + '/Documentos/Requerimento com Lei ' + nomecli1 + ', '+nomecli2+', '+nomecli3+' e ' + nomecli4 + '.docx')

                            # Procuração
                        if self.cbox_procuracao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('PROCURAÇÃO ')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(12)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                'Constitui procurador para a prática de atos perante a Prefeitura Municipal de Hortolândia\n')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(11.5)

                            runner = paragraph.add_run('I - OUTORGANTES:')
                            runner.bold = True

                            paragraph.add_run('\nSr.(a) ' + nomecli1 + ' CPF: ' + cpfcli1 + '')

                            paragraph.add_run('\nSr.(a) ' + nomecli2 + ' CPF: ' + cpfcli2 + '\n\n')

                            paragraph.add_run('\nSr.(a) ' + nomecli3 + ' CPF: ' + cpfcli3 + '\n\n')

                            paragraph.add_run('\nSr.(a) ' + nomecli4 + ' CPF: ' + cpfcli4 + '\n\n')

                            runner = paragraph.add_run('II – OUTORGADO: ')
                            runner.bold = True

                            paragraph.add_run(
                                '\nSr. Rogério Rocha Soares, portador do CPF nº 183.125.858-77- endereço comercial: Rua Geraldo Denadai, n° 74 1º andar sala 03, Jardim da Paineiras na cidade de Hortolândia/SP.\n\n')

                            runner = paragraph.add_run('III – EXTENSÃO DOS PODERES: ')
                            runner.bold = True

                            paragraph.add_run('\nO OUTORGANTE é proprietário e legítimo dono do imóvel, '
                                              'Lote' + loteobra + 'da Quadra ' + quadraobra + ', localizado no endereço: '
                                                                                              '' + endobra + ' nº ' + numobra + ' Loteamento: ' + bairroobra + '.\n\n')

                            runner = paragraph.add_run('Específicos para a prática de Ato Determinado. ')
                            runner.bold = True

                            paragraph.add_run(
                                '\n(X) Retirar projeto aprovado e Alvará de construção, referente ao imóvel acima. '
                                '\nPor este instrumento particular de mandato e na melhor forma de direito, o OUTORGANTE acima qualificado, nomeia e constitui o PROCURADOR acima qualificado, a quem confere plenos poderes de representação perante a ')

                            runner = paragraph.add_run('PREFEITURA MUNICIPAL DE ' + cidadeobra + ',')
                            runner.bold = True

                            paragraph.add_run(
                                'especialmente para em seu nome e como se o próprio fosse praticar os atos especificados acima.')

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n_________________________________                            _________________________________')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'OUTORGANTE:                                                                                                OUTORGANTE:')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(9)
                            font.bold = True

                            paragraph = document.add_paragraph(
                                '' + nomecli1 + '                                                           ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: ' + cpfcli1 + '                                                              CPF: ' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n_________________________________                            _________________________________')
                            paragraph.style = document.styles.add_style('style4.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph(
                                'OUTORGANTE:                                                                                                OUTORGANTE:')
                            paragraph.style = document.styles.add_style('style5.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(9)
                            font.bold = True

                            paragraph = document.add_paragraph(
                                '' + nomecli3 + '                                                           ' + nomecli4 + '')
                            paragraph.style = document.styles.add_style('style6.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF: ' + cpfcli3 + '                                                              CPF: ' + cpfcli4 + '')
                            paragraph.style = document.styles.add_style('style7.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.bold = True

                            paragraph = document.add_paragraph('OUTORGADO:')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(9)
                            font.bold = True

                            paragraph = document.add_paragraph('Eng.º Civil: Rogério Rocha Soares')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(10)

                            paragraph = document.add_paragraph(
                                'CPF: 183.125.858-77')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(10)

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+', '+nomecli3+' e ' + nomecli4 + '/' + tipoobra + '/' + ano + '/Documentos/Procuração ' + nomecli1 + ', '+nomecli2+', '+nomecli3+' e ' + nomecli4 + '.docx')

                            # Declaração
                        if self.cbox_declaracao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(3.49)
                                section.bottom_margin = Cm(1.1)
                                section.left_margin = Cm(2)
                                section.right_margin = Cm(2)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo_branco.jpg", width=Cm(3.5), height=Cm(2.65))

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/Logo_Hortolandia.png", width=Cm(4.9), height=Cm(2.65))

                            paragraph = document.add_paragraph('ANEXO I')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('DECLARAÇÃO')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('Em conformidade com o disposto no artigo 4º '
                                                               'da Lei Municipal nº 2.529, de 04 de abril de 2011,'
                                                               ' que “Dispõe sobre controle ambiental para utilização'
                                                               ' de produtos e subprodutos de madeira de origem nativa'
                                                               ' em obras e serviços de Engenharia Civil no Município'
                                                               ' de Hortolândia”, nós, ' + nomecli1 + ', (' + profissaocli1 + '),' + nomecli2 + ', (' + profissaocli2 + ') , ' + nomecli3 + ', (' + profissaocli3 + ') e ' + nomecli4 + ', (' + profissaocli4 + '),')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run(' Proprietários da obra ')
                            runner.bold = True

                            paragraph.add_run(
                                'localizada à ' + endobra + ', nº ' + numobra + ' Lote ' + loteobra + ', Quadra ' + quadraobra + ', Loteamento ' + bairroobra + ',' \
                                                                                                                                                                'cidade de Hortolândia-SP, DECLARAMOS estar ciente das disposições ' \
                                                                                                                                                                'constantes da Instrução Normativa nº 112, de 21 de agosto de 2006, ' \
                                                                                                                                                                'do Instituto Brasileiro do Meio Ambiente e dos Recursos Naturais ' \
                                                                                                                                                                'Renováveis - IBAMA, e me comprometendo a adquirir, para qualquer ' \
                                                                                                                                                                'serviço a ser realizado no referido imóvel, produtos e subprodutos ' \
                                                                                                                                                                'de madeira de origem não nativa ou nativa que tenha procedência ' \
                                                                                                                                                                'legal, decorrentes de desmatamento autorizado ou de manejo florestal' \
                                                                                                                                                                ' aprovado por órgão ambiental competente, integrante do Sistema ' \
                                                                                                                                                                'Nacional do Meio Ambiente – SISNAMA, com autorização de transporte ' \
                                                                                                                                                                'reconhecida pelo órgão ambiental competente, exigindo no ato da ' \
                                                                                                                                                                'compra que as empresas que comercializem madeiras, forneçam o DOF' \
                                                                                                                                                                ' (Documento de Origem Florestal), acompanhado de nota fiscal. ')

                            runner = paragraph.add_run('DECLARAMOS ')
                            runner.bold = True

                            paragraph.add_run(
                                'ainda, que quando da solicitação do Habite-se o proprietário deverá apresentar além dos documentos, declarações e comprovantes exigidos pelo poder Público, ')

                            runner = paragraph.add_run('um dos seguintes documentos: ')
                            runner.bold = True

                            paragraph = document.add_paragraph(
                                '       ●      Nota Fiscal constando o número do DOF, em caso de utilização de produtos ou subprodutos de madeira de ')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run('aaaaaaaaaaa')
                            font = runner.font
                            font.color.rgb = RGBColor(255, 255, 255)

                            paragraph.add_run('origem nativa;')

                            paragraph = document.add_paragraph(
                                '       ●      Declaração que fez a reutilização de madeira ou que utilizou madeira de reflorestamento;')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '       ●     Declaração que não utilizou madeira de origem nativa, por ter utilizado novas tecnologias ou produtos alternativos.')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_____________________________________')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('PROPRIETÁRIO ' + nomecli1 + '')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_____________________________________')
                            paragraph.style = document.styles.add_style('style8.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('PROPRIETÁRIO ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style9.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_____________________________________')
                            paragraph.style = document.styles.add_style('style8.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('PROPRIETÁRIO ' + nomecli3 + '')
                            paragraph.style = document.styles.add_style('style9.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_____________________________________')
                            paragraph.style = document.styles.add_style('style8.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('PROPRIETÁRIO ' + nomecli4 + '')
                            paragraph.style = document.styles.add_style('style9.4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n\nEm conformidade com o disposto no artigo 4º da '
                                'Lei Municipal nº 2.529, de 04 de abril de 2011, '
                                'que “Dispõe sobre controle ambiental para utilização'
                                ' de produtos e subprodutos de madeira de origem nativa'
                                ' em obras e serviços de Engenharia Civil no Município '
                                'de Hortolândia”, eu, Rogério Rocha Soares, (Engenheiro Civil), '
                                'Autor do Projeto da obra localizada à '
                                'Rua ' + endobra + ', nº ' + numobra + ' Lote ' + loteobra + ', Quadra ' + quadraobra + ', Loteamento ' + bairroobra + ',cidade de Hortolândia-SP,'
                                                                                                                                                       ' DECLARO estar ciente das disposições constantes da Instrução Normativa'
                                                                                                                                                       ' nº 112, de 21 de agosto de 2006, do Instituto Brasileiro '
                                                                                                                                                       'do Meio Ambiente e dos Recursos Naturais Renováveis - IBAMA, e ')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            runner = paragraph.add_run(
                                'me comprometendo a instruir meu cliente a adquirir produtos e subprodutos de madeira de origem não nativa ou nativa que tenha procedência legal, ')
                            runner.underline = True

                            paragraph.add_run(
                                'decorrentes de desmatamento autorizado ou de manejo florestal aprovado por órgão ambiental competente, integrante do Sistema Nacional do Meio Ambiente – SISNAMA, com autorização de transporte reconhecida pelo órgão ambiental competente, exigindo no ato da compra que as empresas que comercializem madeiras, forneçam o DOF (Documento de Origem Florestal), acompanhado de nota fiscal.')

                            paragraph = document.add_paragraph(
                                '\n' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph(
                                '\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            paragraph = document.add_paragraph('AUTOR DO PROJETO')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(10)
                            font.bold = True
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            # runner_word.size = Pt(10)

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+', '+nomecli3+' e ' + nomecli4 +'/' + tipoobra + '/' + ano + '/Documentos/Declaração ' + nomecli1 + ', '+nomecli2+', '+nomecli3+' e ' + nomecli4 +'.docx')

                            # ---------------------MEMORIAL DESCRITIVO PARA CONSTRUÇÃO---------------------------------------------------------------------------------------------
                        if self.cbox_memorialcontrucao.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('MEMORIAL DESCRITIVO PARA CONSTRUÇÃO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                'Obra: REGULARIZAÇÃO E CONSTRUÇÃO RESIDENCIAL MULTIFAMILIAR – R2')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('(DE ACORDO COM A LEI 3.491/2018 - ANISTIA)')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('Local: ' + endobra + '- N° ' + numobra + '')
                            paragraph.style = document.styles.add_style('style4', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Lote: ' + loteobra + ' Quadra: ' + quadraobra + '')
                            paragraph.style = document.styles.add_style('style5', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Loteamento: ' + bairroobra + '')
                            paragraph.style = document.styles.add_style('style6', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Município: ' + cidadeobra + '/SP')
                            paragraph.style = document.styles.add_style('style7', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph(
                                'Proprietário(s): ' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 + '')
                            paragraph.style = document.styles.add_style('style8', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph('Responsável Técnico: Eng. Civil ROGÉRIO ROCHA SOARES')
                            paragraph.style = document.styles.add_style('style9', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)

                            paragraph = document.add_paragraph(
                                '01 – Preparação do terreno: Raspagem com moto-niveladora e os resíduos será retirado e depositado em local devidamente autorizado e legalizado pela Prefeitura Municipal.')
                            paragraph.style = document.styles.add_style('style10', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '02 – Fundações: A fundação do alicerçada em estacas manuais, blocos e vigas baldrame em concreto armado.')
                            paragraph.style = document.styles.add_style('style11', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '03 – Impermeabilização: Com Vedacit no concreto e na argamassa de embasamento. Será aplicado neutrol no baldrame antes do reaterro.')
                            paragraph.style = document.styles.add_style('style12', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '04 – Estrutura: Os pilares, vergas, contra-vergas, vigas de apoio e vigas de respaldo em concreto armado batido na própria obra.')
                            paragraph.style = document.styles.add_style('style13', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '05 - Alvenaria: As paredes externas e internas executadas com bloco cerâmico (barro) de oito furos (15 cm), assentados com argamassa de areia e cimento. A altura do pé direito será de 2,80m em toda a casa.')
                            paragraph.style = document.styles.add_style('style14', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph('06 - Forro: O forro de laje pré-moldada de concreto.')
                            paragraph.style = document.styles.add_style('style15', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '07- Cobertura: O telhado foi executado com telhas de barro estrutura de madeira com inclinação de 26% na residência.')
                            paragraph.style = document.styles.add_style('style16', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '08 - Instalações hidráulicas: Instalado conforme normas da SABESP e NBR 7229/93,')
                            paragraph.style = document.styles.add_style('style17', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        a) - Água fria: Abastecida pela rede pública e armazenada em dois reservatórios elevados, com capacidade de 1.000 litros d’agua na residência;')
                            paragraph.style = document.styles.add_style('style18', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        b) - Esgotos sanitários: canalizados com tubos de PVC com diâmetro 100 mm para a rede pública de afastamento de esgoto.')
                            paragraph.style = document.styles.add_style('style19', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '        c) – Águas pluviais: coletada por meio de sistema de calhas galvanizadas e desce por condutores e são canalizadas até a sarjeta por meio de tubos de PVC, por sob a calçada.')
                            paragraph.style = document.styles.add_style('style20', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '09 - Instalações elétricas: Instalado conforme normas da CPFL, composto de um ponto de luz para cada compartimento e tomadas em número suficiente para atender a demanda.')
                            paragraph.style = document.styles.add_style('style21', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '10 – Revestimento: Os banheiros da residência e da dependência em azulejo até a altura do forro;')
                            paragraph.style = document.styles.add_style('style22', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '11 - Pisos: Nivelado com argamassa de areia e cimento e revestimento cerâmico em todos os compartimentos.')
                            paragraph.style = document.styles.add_style('style23', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '12 - Esquadrias: As Janelas e portas externas são em alumínio e as portas internas de madeira, conforme as medidas indicadas no projeto.')
                            paragraph.style = document.styles.add_style('style24', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '13 – Pintura: As paredes de alvenaria, internas pintadas com tinta látex PVA; as paredes externas pintadas com tinta látex, acrílica, as janelas e portas serão pintadas com esmalte sintético.')
                            paragraph.style = document.styles.add_style('style25', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '14 – Limpeza Geral da Obra: Após a conclusão da obra, foi feita a retirada dos restos de materiais e entulho.')
                            paragraph.style = document.styles.add_style('style26', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            paragraph = document.add_paragraph(
                                '' + cidadeobra + ', ' + dia + ' de ' + mesescrito + ' de ' + ano + '.')
                            paragraph.style = document.styles.add_style('style27', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'
                            font.size = Pt(9)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n_________________________________                    _____________________________________')
                            paragraph.style = document.styles.add_style('style28', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário: ' + nomecli1 + '                               Proprietário: ' + nomecli2 + '')
                            paragraph.style = document.styles.add_style('style29', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli1 + '                                                        CPF:' + cpfcli2 + '')
                            paragraph.style = document.styles.add_style('style30', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n_________________________________                    _____________________________________')
                            paragraph.style = document.styles.add_style('style28.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Proprietário: ' + nomecli3 + '                               Proprietário: ' + nomecli4 + '')
                            paragraph.style = document.styles.add_style('style29.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CPF:' + cpfcli3 + '                                                        CPF:' + cpfcli4 + '')
                            paragraph.style = document.styles.add_style('style30.2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                '\n\n\n\n\n\n\n_________________________________')
                            paragraph.style = document.styles.add_style('style28.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'ROGÉRIO ROCHA SOARES')
                            paragraph.style = document.styles.add_style('style29.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'Engenheiro Civil')
                            paragraph.style = document.styles.add_style('style30.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'CREA: 5070347192-SP')
                            paragraph.style = document.styles.add_style('style31.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            paragraph = document.add_paragraph(
                                'SMPUGE: 1036/18')
                            paragraph.style = document.styles.add_style('style32.3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Arial'

                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 + '/' + tipoobra + '/' + ano + '/Documentos/Memorial Descritivo Para Construção ' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 + '.docx')

                            # ---------------------Recibo---------------------------------------------------------------------------------------------
                        if self.cbox_recibo.isChecked() == True:
                            document = Document()

                            sections = document.sections
                            for section in sections:
                                section.top_margin = Cm(-4.5)
                                section.bottom_margin = Cm(2)
                                section.left_margin = Cm(2.5)
                                section.right_margin = Cm(1.75)

                            section = document.sections[0]

                            header = document.sections[0].header
                            logo = header.paragraphs[0]
                            logo_run = logo.add_run()
                            logo_run.add_picture("images/logo.png", width=Cm(2.65), height=Cm(2.65))

                            paragraph = document.add_paragraph('RECIBO')
                            paragraph.style = document.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            font.size = Pt(16)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            paragraph = document.add_paragraph('Eu ROGÉRIO ROCHA SOARES engenheiro civil CREA: '
                                                               '5070374192, recebi de ' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 +', '
                                                                                                                        'parte do pagamento para aprovação de projeto '
                                                                                                                        'arquitetônico a quantia de R$ ' + valorparcobra + ',00 (' + num2words(
                                valorparcobra.replace(",", "."), lang='pt-br') +
                                                               '), de um total de '
                                                               'R$ ' + valorobra + ',00 (' + num2words(
                                (valorobra).replace(",", "."), lang='pt-br') + ').')
                            paragraph.style = document.styles.add_style('style2', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.bold = True
                            font.name = 'Book Antiqua'
                            font.size = Pt(12)

                            paragraph = document.add_paragraph(
                                '\n\n' + cidadeobra + ', ' + dia + ' / ' + mesescrito + ' / ' + ano + '.')
                            paragraph.style = document.styles.add_style('style3', WD_STYLE_TYPE.PARAGRAPH)
                            font = paragraph.style.font
                            font.name = 'Book Antiqua'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            font.size = Pt(12)

                            # footer section
                            footer_section = document.sections[0]
                            footer = footer_section.footer

                            # footer text
                            footer_text = footer.paragraphs[0]
                            footer_text.text = "_______________________________________________________________________________________________" \
                                               "\nROGÉRIO ROCHA SOARES   TEL. : (19) 982009858" \
                                               "\nCREA: 5070347192" \
                                               "\nE-MAIL: rocha.soares@hotmail.com"
                            document.save(
                                '//ROGER2/Users/ROCHA/Documents/PROCESSO DE CLIENTES/' + cidadeobra + '/' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 +'/' + tipoobra + '/' + ano + '/Documentos/Recibo ' + nomecli1 + ', '+nomecli2+', ' + nomecli3 + ' e ' + nomecli4 +'.docx')

                        up.show()
                        pop.lbl_popup.setText("DOCUMENTOS CRIADOS")
                        pop.frame_popup.setStyleSheet("background-color: rgb(57, 173, 84);\n"
                                                       "border-radius:5px;")
                        pop.lbl_popup.setStyleSheet("color: rgb(35, 35, 35)")

                        self.txt_idobra.setText(None)
                        self.txt_idcli1.setText(None)
                        self.txt_idcli2.setText(None)
                        self.txt_idcli3.setText(None)

            else:
                up.show()
                pop.lbl_popup.setText("ESCOLHA A QUANTIDADE DE CLIENTES!")
                pop.frame_popup.setStyleSheet("background-color: rgb(255, 11, 15);\n"
                                              "border-radius:5px;")
                pop.lbl_popup.setStyleSheet("color: rgb(200, 200, 255)")

        else:
            up.show()
            pop.lbl_popup.setText("DIGITE O ID DA OBRA!")
            pop.frame_popup.setStyleSheet("background-color: rgb(255, 11, 15);\n"
                                           "border-radius:5px;")
            pop.lbl_popup.setStyleSheet("color: rgb(200, 200, 255)")

if __name__ == '__main__':
    qt = QApplication(sys.argv)
    widget = QtWidgets.QStackedWidget()
    up = QtWidgets.QStackedWidget()
    menu = Menu()
    pop = Pop()
    up.addWidget(pop)
    widget.addWidget(menu)
    #widget.setWindowIcon(QtGui.QIcon('Images/logo_preto.ico'))
    #widget.setFixedHeight(870)
    #widget.setFixedWidth(1039)
    up.setWindowFlags(QtCore.Qt.WindowType.FramelessWindowHint)
    up.setAttribute(QtCore.Qt.WidgetAttribute.WA_TranslucentBackground)
    widget.setWindowFlags(QtCore.Qt.WindowType.FramelessWindowHint)
    widget.setAttribute(QtCore.Qt.WidgetAttribute.WA_TranslucentBackground)
    widget.show()
    widget.setMinimumSize(QtCore.QSize(586, 716))
    #widget.showMaximized()
    sys.exit(qt.exec())
